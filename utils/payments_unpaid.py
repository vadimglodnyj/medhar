# -*- coding: utf-8 -*-
"""Неоплачені стаціонар/реабілітація/відпустка → відомість."""

from __future__ import annotations

import calendar
import json
import logging
import os
import re
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta
from typing import Iterable, Optional

import pandas as pd

from config import (
    COMBAT_DIAGNOSIS_PATTERNS_FILE,
    PAYMENTS_COMBAT_COLUMN_FALLBACK,
    PAYMENTS_DIR,
    PAYMENTS_FILE_RE,
    PAYMENTS_HISTORY_START,
    PAYMENTS_PREV_MONTH_CARRY_DAY,
    PAYMENTS_UNIT_FILTER,
)
import config as _cfg

logger = logging.getLogger(__name__)

JOURNAL_PAYABLE_CARE = {
    "inpatient": "stat",
    "rehab": "rehab",
    "vacation": "leave",
}

PAYABLE_TYPES = {
    "стаціонар": "stat",
    "реабілітація": "rehab",
    "відпустка": "leave",
}

MONTH_EN_TO_NUM = {
    "january": 1,
    "february": 2,
    "march": 3,
    "april": 4,
    "may": 5,
    "june": 6,
    "july": 7,
    "august": 8,
    "september": 9,
    "october": 10,
    "november": 11,
    "december": 12,
}

MONTH_UA = {
    1: "січень",
    2: "лютий",
    3: "березень",
    4: "квітень",
    5: "травень",
    6: "червень",
    7: "липень",
    8: "серпень",
    9: "вересень",
    10: "жовтень",
    11: "листопад",
    12: "грудень",
}

TYPE_UA = {
    "stat": "Стаціонар",
    "rehab": "Реабілітація",
    "leave": "Відпустка",
}


def history_start() -> date:
    return date.fromisoformat(PAYMENTS_HISTORY_START)


def normalize_pib(value: str) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    # уніфікація апострофів
    text = text.replace("’", "'").replace("`", "'").replace("ʼ", "'")
    return text.casefold()


def format_date_ua(d: Optional[date]) -> str:
    if not d:
        return ""
    return d.strftime("%d.%m.%Y")


def inclusive_days(start: date, end: date) -> int:
    if end < start:
        return 0
    return (end - start).days + 1


def month_bounds(year: int, month: int) -> tuple[date, date]:
    last = calendar.monthrange(year, month)[1]
    return date(year, month, 1), date(year, month, last)


def parse_any_date(value) -> Optional[date]:
    if value is None:
        return None
    try:
        if pd.isna(value):
            return None
    except (TypeError, ValueError):
        pass
    if isinstance(value, pd.Timestamp):
        return value.date()
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    text = str(value).strip()
    if not text or text.lower() in ("nan", "none", "nat", "-"):
        return None
    # лише рік
    if re.fullmatch(r"\d{4}", text):
        try:
            return date(int(text), 1, 1)
        except ValueError:
            return None
    for fmt in ("%d.%m.%Y", "%Y-%m-%d", "%d/%m/%Y", "%d.%m.%y", "%Y.%m.%d"):
        try:
            return datetime.strptime(text[:10], fmt).date()
        except ValueError:
            continue
    # українські місяці / вільний текст — взяти першу дату
    m = re.search(r"(\d{1,2})[./](\d{1,2})[./](\d{2,4})", text)
    if m:
        d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if y < 100:
            y += 2000
        try:
            return date(y, mo, d)
        except ValueError:
            return None
    try:
        ts = pd.to_datetime(text, dayfirst=True, errors="coerce")
        if pd.notna(ts):
            return ts.date()
    except Exception:
        pass
    return None


def parse_date_list(value) -> list[date]:
    """Клітинка може містити кілька дат через \\n."""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return []
    if isinstance(value, (datetime, date, pd.Timestamp)):
        d = parse_any_date(value)
        return [d] if d else []
    text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    parts = re.split(r"[\n;]+", text)
    out: list[date] = []
    for part in parts:
        d = parse_any_date(part.strip())
        if d:
            out.append(d)
    return out


def zip_periods(starts, ends) -> list[tuple[date, date]]:
    start_list = parse_date_list(starts)
    end_list = parse_date_list(ends)
    n = min(len(start_list), len(end_list))
    periods = []
    for i in range(n):
        a, b = start_list[i], end_list[i]
        if b < a:
            a, b = b, a
        periods.append((a, b))
    return periods


def clip_interval(
    start: date, end: date, lo: date, hi: date
) -> Optional[tuple[date, date]]:
    a = max(start, lo)
    b = min(end, hi)
    if b < a:
        return None
    return a, b


def subtract_intervals(
    base: tuple[date, date], paid: Iterable[tuple[date, date]]
) -> list[tuple[date, date]]:
    """Повертає неоплачені залишки base після віднімання paid."""
    remaining = [base]
    for p_start, p_end in paid:
        next_rem: list[tuple[date, date]] = []
        for r_start, r_end in remaining:
            if p_end < r_start or p_start > r_end:
                next_rem.append((r_start, r_end))
                continue
            if p_start > r_start:
                next_rem.append((r_start, p_start - timedelta(days=1)))
            if p_end < r_end:
                next_rem.append((p_end + timedelta(days=1), r_end))
        remaining = [(a, b) for a, b in next_rem if b >= a]
    return remaining


# --- combat diagnosis patterns ---

_patterns_cache: Optional[list[tuple[str, str, re.Pattern]]] = None


def load_combat_patterns(
    *, force_reload: bool = False
) -> list[tuple[str, str, re.Pattern]]:
    global _patterns_cache
    if _patterns_cache is not None and not force_reload:
        return _patterns_cache
    path = COMBAT_DIAGNOSIS_PATTERNS_FILE
    items: list[tuple[str, str, re.Pattern]] = []
    try:
        with open(path, "r", encoding="utf-8") as fh:
            data = json.load(fh)
        for raw in data.get("patterns") or []:
            pid = str(raw.get("id") or "").strip() or "pat"
            label = str(raw.get("label") or pid).strip()
            regex = str(raw.get("regex") or "").strip()
            if not regex:
                continue
            try:
                items.append((pid, label, re.compile(regex)))
            except re.error as e:
                logger.warning("Невалідний regex '%s': %s", regex, e)
    except FileNotFoundError:
        logger.warning("Немає %s — бойовий фільтр порожній", path)
    except Exception as e:
        logger.warning("Не вдалося прочитати патерни діагнозів: %s", e)
    _patterns_cache = items
    return items


def match_combat_diagnosis(text: str) -> Optional[dict]:
    blob = str(text or "")
    if not blob.strip():
        return None
    for pid, label, cre in load_combat_patterns():
        m = cre.search(blob)
        if m:
            start = max(0, m.start() - 40)
            end = min(len(blob), m.end() + 40)
            snippet = blob[start:end].replace("\n", " ").strip()
            return {
                "pattern_id": pid,
                "pattern_label": label,
                "snippet": snippet,
            }
    return None


# --- payments files ---


@dataclass
class PaymentFileInfo:
    path: str
    filename: str
    month: int
    year: int
    label: str


def list_payment_files() -> list[PaymentFileInfo]:
    os.makedirs(PAYMENTS_DIR, exist_ok=True)
    out: list[PaymentFileInfo] = []
    for name in sorted(os.listdir(PAYMENTS_DIR)):
        m = PAYMENTS_FILE_RE.match(name)
        if not m:
            continue
        month_en = m.group(1).lower()
        year = int(m.group(2))
        month = MONTH_EN_TO_NUM[month_en]
        out.append(
            PaymentFileInfo(
                path=os.path.join(PAYMENTS_DIR, name),
                filename=name,
                month=month,
                year=year,
                label=f"{MONTH_UA[month]} {year}",
            )
        )
    out.sort(key=lambda x: (x.year, x.month))
    return out


def _find_vedomost_sheet(xls: pd.ExcelFile) -> str:
    for name in xls.sheet_names:
        if "відом" in str(name).casefold() or "vedom" in str(name).casefold():
            return name
    return xls.sheet_names[0]


def load_paid_intervals_by_pib() -> dict[str, dict[str, list[tuple[date, date]]]]:
    """
    {pib_key: {"stat": [(s,e),...], "rehab": [...], "leave": [...]}}
    з усіх файлів data/payments/*.xlsx (лист Відомість).
    """
    result: dict[str, dict[str, list[tuple[date, date]]]] = {}
    for info in list_payment_files():
        try:
            xls = pd.ExcelFile(info.path)
            sheet = _find_vedomost_sheet(xls)
            raw = pd.read_excel(info.path, sheet_name=sheet, header=None)
        except Exception as e:
            logger.warning("Не вдалося прочитати %s: %s", info.filename, e)
            continue
        # дані з рядка 3 (після 3 рядків заголовка)
        for i in range(3, len(raw)):
            row = raw.iloc[i]
            pib = str(row.iloc[4] if len(row) > 4 else "").strip()
            if not pib or pib.lower() in ("nan", "none"):
                continue
            # пропуск рядків-номерів заголовка
            if pib.isdigit():
                continue
            key = normalize_pib(pib)
            bucket = result.setdefault(
                key, {"stat": [], "rehab": [], "leave": [], "pib_display": pib}
            )
            if len(row) > 8:
                for a, b in zip_periods(row.iloc[7], row.iloc[8]):
                    bucket["stat"].append((a, b))
            if len(row) > 10:
                for a, b in zip_periods(row.iloc[9], row.iloc[10]):
                    bucket["rehab"].append((a, b))
            if len(row) > 12:
                for a, b in zip_periods(row.iloc[11], row.iloc[12]):
                    bucket["leave"].append((a, b))
    return result


def normalize_treatment_type(value: str) -> Optional[str]:
    text = re.sub(r"\s+", " ", str(value or "")).strip().casefold()
    if not text:
        return None
    for needle, code in PAYABLE_TYPES.items():
        if needle in text:
            return code
    return None


@dataclass
class UnpaidSegment:
    pib: str
    pib_key: str
    unit: str
    position: str
    rank: str
    injury_date: Optional[date]
    care_type: str  # stat|rehab|leave
    start: date
    end: date
    days: int
    pattern_label: str
    diagnosis_snippet: str
    diagnosis_full: str
    missing_discharge: bool = False


@dataclass
class VedomostRow:
    n: int
    unit: str
    position: str
    rank: str
    pib: str
    injury_date: str
    days: int
    days_text: str
    stat_start: str
    stat_end: str
    rehab_start: str
    rehab_end: str
    leave_start: str
    leave_end: str
    segments: list[UnpaidSegment] = field(default_factory=list)


def row_marked_combat(row) -> bool:
    """Колонка Excel «Бойова/ небойова» (лише якщо явно «Бойова», не «Не Бойова»)."""
    col = re.sub(
        r"\s+",
        " ",
        str(
            row.get("Бойова/ небойова")
            or row.get("Бойова / небойова")
            or ""
        ),
    ).strip()
    if not col:
        return False
    low = col.casefold()
    if low.startswith("не ") or low.startswith("не\u00a0"):
        return False
    return low.startswith("бойова")


def resolve_combat_match(row, diag: str) -> Optional[dict]:
    matched = match_combat_diagnosis(diag)
    if matched:
        return matched
    if PAYMENTS_COMBAT_COLUMN_FALLBACK and row_marked_combat(row):
        return {
            "pattern_id": "combat_column",
            "pattern_label": "колонка «Бойова»",
            "snippet": str(
                row.get("Бойова/ небойова")
                or row.get("Бойова / небойова")
                or "Бойова"
            ).strip(),
        }
    return None


def unit_matches_filter(unit: str, unit_filter: str = None) -> bool:
    """True якщо підрозділ належить до 2 БОП (або фільтр вимкнено)."""
    needle = (unit_filter if unit_filter is not None else PAYMENTS_UNIT_FILTER) or ""
    needle = re.sub(r"\s+", " ", needle).strip()
    if not needle:
        return True
    text = re.sub(r"\s+", " ", str(unit or "")).strip()
    # «2 БОП» / «2БОП», без «12 БОП»
    esc = re.escape(needle).replace(r"\ ", r"\s*")
    return bool(re.search(rf"(?<!\d){esc}", text, flags=re.IGNORECASE))


def _row_pib(row) -> str:
    sur = str(row.get("Прізвище") or "").strip()
    nam = str(row.get("Ім'я") or "").strip()
    pat = str(row.get("По батькові") or "").strip()
    if sur or nam or pat:
        return re.sub(r"\s+", " ", f"{sur} {nam} {pat}").strip()
    return str(row.get("ПІБ") or "").strip()


def payments_treatments_path(year: int) -> str:
    """Той самий Excel, що на сторінці «Бази даних» (desktop: AppData / Dropbox)."""
    return os.path.join(_cfg.EXCEL_DATA_DIR, f"treatments_{int(year)}.xlsx")


def load_payments_treatments_df(year: int) -> pd.DataFrame:
    """
    База лікувань для відомості оплат — лише treatments_{year}.xlsx
    (як у ручній відомості, без злиття 2024/2025).
    """
    path = payments_treatments_path(year)
    if not os.path.isfile(path):
        return pd.DataFrame()
    df = pd.read_excel(path)
    df.columns = df.columns.str.strip()
    for col in ["Прізвище", "Ім'я", "По батькові"]:
        if col in df.columns:
            df[col] = df[col].fillna("").astype(str)
    if all(c in df.columns for c in ["Прізвище", "Ім'я", "По батькові"]):
        df["ПІБ"] = (
            df["Прізвище"].astype(str).str.strip()
            + " "
            + df["Ім'я"].astype(str).str.strip()
            + " "
            + df["По батькові"].astype(str).str.strip()
        ).str.replace(r"\s+", " ", regex=True).str.strip()
    return df


def collect_journal_payment_candidates(
    *,
    default_end: Optional[date] = None,
) -> list[dict]:
    """
    Епізоди з амбулаторного журналу: є поранення + стаціонар/реаб/відпустка.
    Без здогадок по тексті діагнозу.
    """
    from utils import patient_cards_db as cards_db

    hs = history_start()
    out: list[dict] = []
    for row in cards_db.list_payment_journal_episodes():
        care = JOURNAL_PAYABLE_CARE.get(_normalize_care(row.get("care_type")))
        if not care:
            continue
        start = parse_any_date(row.get("leave_start") or "") or parse_any_date(
            row.get("visit_date") or ""
        )
        end = parse_any_date(row.get("leave_end") or "")
        missing_discharge = False
        if not start:
            continue
        if not end:
            if default_end is None or start > default_end:
                continue
            end = default_end
            missing_discharge = True
        if end < start:
            continue
        clipped = clip_interval(start, end, hs, date(9999, 12, 31))
        if not clipped:
            continue
        start, end = clipped
        pib = (row.get("pib") or "").strip()
        if not pib:
            continue
        injury = parse_any_date(row.get("injury_date") or "")
        injury_title = (row.get("injury_title") or "").strip() or "Поранення"
        basis = injury_title
        if injury:
            basis = f"{injury_title} {injury.strftime('%d.%m.%Y')}"
        unit = cards_db.format_wa_unit_line(
            row.get("unit_short") or "",
            canonical_unit=(row.get("rank_unit") or ""),
        )
        diag = re.sub(r"\s+", " ", str(row.get("diagnosis") or "")).strip()
        out.append(
            {
                "pib": pib,
                "pib_key": normalize_pib(pib),
                "unit": unit,
                "position": str(row.get("position") or "").strip(),
                "rank": str(row.get("rank") or "").strip(),
                "injury_date": injury,
                "care_type": care,
                "start": start,
                "end": end,
                "pattern_label": basis,
                "diagnosis_snippet": diag[:160],
                "diagnosis_full": diag[:500],
                "missing_discharge": missing_discharge,
            }
        )
    return out


def _normalize_care(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip().casefold()


def vedomost_display_parts(
    u_start: date,
    u_end: date,
    year: int,
    month: int,
) -> list[tuple[date, date, int]]:
    """
    Частини неоплаченого інтервалу для відомості за місяць.
    Дні — лише в межах місяця відомості; «хвіст» попереднього місяця — окремо.
    Кінець може виходити за місяць (напр. відпустка до 06.08), але дні = лише липень.
    """
    month_lo, month_hi = month_bounds(year, month)
    if u_end < month_lo:
        return [(u_start, u_end, inclusive_days(u_start, u_end))]
    if u_start > month_hi:
        return []
    out: list[tuple[date, date, int]] = []
    if u_start < month_lo:
        prev_end = min(u_end, month_lo - timedelta(days=1))
        if prev_end >= u_start:
            out.append((u_start, prev_end, inclusive_days(u_start, prev_end)))
    if u_end >= month_lo:
        m_start = max(u_start, month_lo)
        m_end_clip = min(u_end, month_hi)
        if m_end_clip >= m_start:
            out.append((m_start, u_end, inclusive_days(m_start, m_end_clip)))
    return out


def segment_for_vedomost_month(
    seg_start: date,
    seg_end: date,
    year: int,
    month: int,
) -> bool:
    """
    Чи включати неоплачений сегмент у відомість за місяць.
    — перетинає обраний місяць, або
    — «хвіст» попереднього місяця (закінчення не раніше PAYMENTS_PREV_MONTH_CARRY_DAY).
    """
    month_lo, month_hi = month_bounds(year, month)
    if seg_start > month_hi:
        return False
    if seg_end >= month_lo and seg_start <= month_hi:
        return True
    carry_day = max(1, int(PAYMENTS_PREV_MONTH_CARRY_DAY or 18))
    if month == 1:
        carry_from = date(year - 1, 12, carry_day)
        prev_hi = date(year - 1, 12, 31)
    else:
        carry_from = date(year, month - 1, carry_day)
        prev_hi = date(year, month - 1, calendar.monthrange(year, month - 1)[1])
    return seg_end >= carry_from and seg_end <= prev_hi


def find_unpaid_segments(
    *,
    year: int,
    month: int,
) -> list[UnpaidSegment]:
    """
    Неоплачені залишки з журналу для відомості за місяць.
    Дні рахуються лише в межах обраного місяця (або повністю для червневого «хвоста»).
    """
    paid_map = load_paid_intervals_by_pib()
    _, month_hi = month_bounds(year, month)
    candidates = collect_journal_payment_candidates(default_end=month_hi)
    segments: list[UnpaidSegment] = []
    for c in candidates:
        paid_list = paid_map.get(c["pib_key"], {}).get(c["care_type"], [])
        unpaid = subtract_intervals((c["start"], c["end"]), paid_list)
        for u_start, u_end in unpaid:
            if not segment_for_vedomost_month(u_start, u_end, year, month):
                continue
            for disp_start, disp_end, disp_days in vedomost_display_parts(
                u_start, u_end, year, month
            ):
                segments.append(
                    UnpaidSegment(
                        pib=c["pib"],
                        pib_key=c["pib_key"],
                        unit=c["unit"],
                        position=c["position"],
                        rank=c["rank"],
                        injury_date=c["injury_date"],
                        care_type=c["care_type"],
                        start=disp_start,
                        end=disp_end,
                        days=disp_days,
                        pattern_label=c["pattern_label"],
                        diagnosis_snippet=c["diagnosis_snippet"],
                        diagnosis_full=c["diagnosis_full"],
                        missing_discharge=c.get("missing_discharge", False),
                    )
                )
    segments.sort(key=lambda s: (s.pib_key, s.care_type, s.start))
    return segments


def _join_dates(dates: list[date]) -> str:
    return "\n".join(format_date_ua(d) for d in dates)


def aggregate_vedomost_rows(segments: list[UnpaidSegment]) -> list[VedomostRow]:
    """Один рядок відомості на ПІБ; дати через \\n."""
    by_pib: dict[str, list[UnpaidSegment]] = {}
    for seg in segments:
        by_pib.setdefault(seg.pib_key, []).append(seg)

    rows: list[VedomostRow] = []
    for idx, key in enumerate(sorted(by_pib.keys()), start=1):
        segs = by_pib[key]
        first = segs[0]
        meta = max(segs, key=lambda s: s.end)
        injury_src = min(segs, key=lambda s: s.start)
        injury = injury_src.injury_date

        def pack(care: str) -> tuple[str, str, list[int]]:
            items = sorted(
                [s for s in segs if s.care_type == care], key=lambda s: s.start
            )
            starts = [s.start for s in items]
            ends = [s.end for s in items]
            days = [s.days for s in items]
            return _join_dates(starts), _join_dates(ends), days

        st_s, st_e, st_d = pack("stat")
        rh_s, rh_e, rh_d = pack("rehab")
        lv_s, lv_e, lv_d = pack("leave")
        all_days = st_d + rh_d + lv_d
        total = sum(all_days)
        days_text = "\n".join(str(d) for d in all_days) if all_days else "0"

        rows.append(
            VedomostRow(
                n=idx,
                unit=meta.unit or first.unit,
                position=meta.position or first.position,
                rank=meta.rank or first.rank,
                pib=meta.pib or first.pib,
                injury_date=format_date_ua(injury) if injury else "",
                days=total,
                days_text=days_text if len(all_days) > 1 else str(total),
                stat_start=st_s,
                stat_end=st_e,
                rehab_start=rh_s,
                rehab_end=rh_e,
                leave_start=lv_s,
                leave_end=lv_e,
                segments=segs,
            )
        )
    return rows


def unique_matched_diagnoses(segments: list[UnpaidSegment], limit: int = 80) -> list[dict]:
    seen = set()
    out = []
    for s in segments:
        key = (s.pattern_label, s.diagnosis_full[:120])
        if key in seen:
            continue
        seen.add(key)
        out.append(
            {
                "pattern": s.pattern_label,
                "diagnosis": s.diagnosis_full[:200],
            }
        )
        if len(out) >= limit:
            break
    return out


def build_docx_context(rows: list[VedomostRow], year: int, month: int) -> dict:
    return {
        "month": MONTH_UA.get(month, str(month)),
        "year": str(year),
        "rows": [
            {
                "n": r.n,
                "unit": r.unit,
                "position": r.position,
                "rank": r.rank,
                "pib": r.pib,
                "injury_date": r.injury_date,
                "days": r.days_text,
                "stat_start": r.stat_start,
                "stat_end": r.stat_end,
                "rehab_start": r.rehab_start,
                "rehab_end": r.rehab_end,
                "leave_start": r.leave_start,
                "leave_end": r.leave_end,
            }
            for r in rows
        ],
    }


def _set_cell_text(cell, text: str, *, size_pt: int = 8) -> None:
    from docx.shared import Pt
    from docx.oxml.ns import qn

    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text or ""))
    run.font.size = Pt(size_pt)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")


def render_payments_docx(
    template_path: str,
    output_path: str,
    rows: list[VedomostRow],
    *,
    year: int,
    month: int,
) -> str:
    """
    Підставляє month/year через docxtpl, потім додає рядки таблиці через python-docx.
    """
    from docxtpl import DocxTemplate
    from docx import Document

    tpl = DocxTemplate(template_path)
    tpl.render(
        {
            "month": MONTH_UA.get(month, str(month)),
            "year": str(year),
        }
    )
    tpl.save(output_path)

    doc = Document(output_path)
    if not doc.tables:
        raise RuntimeError("У шаблоні відомості немає таблиці")
    table = doc.tables[0]
    for r in rows:
        cells = table.add_row().cells
        values = [
            str(r.n),
            r.unit,
            r.position,
            r.rank,
            r.pib,
            r.injury_date,
            r.days_text,
            r.stat_start,
            r.stat_end,
            r.rehab_start,
            r.rehab_end,
            r.leave_start,
            r.leave_end,
        ]
        for i, val in enumerate(values):
            if i < len(cells):
                _set_cell_text(cells[i], val)
    doc.save(output_path)
    return output_path
