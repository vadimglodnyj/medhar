import os
from docx import Document

def describe_paragraph(p):
    fmt = p.paragraph_format
    indent_left = getattr(fmt, 'left_indent', None)
    indent_right = getattr(fmt, 'right_indent', None)
    first_line = getattr(fmt, 'first_line_indent', None)
    space_before = getattr(fmt, 'space_before', None)
    space_after = getattr(fmt, 'space_after', None)
    line_spacing = getattr(fmt, 'line_spacing', None)
    keep_together = getattr(fmt, 'keep_together', None)
    keep_with_next = getattr(fmt, 'keep_with_next', None)
    outline_level = getattr(fmt, 'outline_level', None)

    list_info = None
    # Detect numbering/bullets if present
    try:
        if p._p.pPr is not None and p._p.pPr.numPr is not None:
            ilvl = p._p.pPr.numPr.ilvl.val if p._p.pPr.numPr.ilvl is not None else None
            numId = p._p.pPr.numPr.numId.val if p._p.pPr.numPr.numId is not None else None
            list_info = {"numId": numId, "ilvl": ilvl}
    except Exception:
        list_info = None

    def twips(val):
        try:
            return None if val is None else int(val)
        except Exception:
            try:
                # python-docx uses Length type; it can be converted to twips via .twips
                return getattr(val, 'twips', None)
            except Exception:
                return None

    return {
        "text": p.text,
        "style": p.style.name if p.style is not None else None,
        "left_indent_twips": twips(indent_left),
        "right_indent_twips": twips(indent_right),
        "first_line_indent_twips": twips(first_line),
        "space_before_twips": twips(space_before),
        "space_after_twips": twips(space_after),
        "line_spacing": line_spacing,
        "keep_together": keep_together,
        "keep_with_next": keep_with_next,
        "outline_level": outline_level,
        "list": list_info,
    }

def inspect_docx(path):
    if not os.path.exists(path):
        print(f"File not found: {path}")
        return
    doc = Document(path)
    print(f"Paragraphs: {len(doc.paragraphs)}")
    for idx, p in enumerate(doc.paragraphs):
        info = describe_paragraph(p)
        print(f"--- Paragraph {idx+1} ---")
        print(f"Text: {info['text']}")
        # Leading spaces count for precise visual indent
        leading_spaces = len(info['text']) - len(info['text'].lstrip(' '))
        print(f"Leading spaces: {leading_spaces}")
        print(f"Style: {info['style']}")
        print(f"Indents (twips): left={info['left_indent_twips']}, first_line={info['first_line_indent_twips']}, right={info['right_indent_twips']}")
        print(f"Spacing (twips): before={info['space_before_twips']}, after={info['space_after_twips']}")
        print(f"Line spacing: {info['line_spacing']}")
        print(f"List: {info['list']}")

if __name__ == "__main__":
    inspect_docx(os.path.join(os.path.dirname(os.path.dirname(__file__)), 'output', 'medchar_output.docx'))


