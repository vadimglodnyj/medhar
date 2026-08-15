#!/usr/bin/env python3
"""
Головний Flask додаток для генерації медичних документів
"""

import os
import sys
import shutil
import tempfile
import threading
import json
import time
import sqlite3
import pandas as pd
from flask import Flask, render_template, request, send_file, make_response, flash, jsonify, redirect, url_for
from werkzeug.utils import secure_filename

from docxtpl import DocxTemplate
from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
import io
import re
import base64
import logging
from datetime import datetime

# Додаємо поточну папку до шляху
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# Імпортуємо конфігурацію
from config import *

from utils.circumstances_parser import parse_circumstances
from utils.ukrainian_pib_genitive import (
    build_pib_rodovyi_for_document,
    format_nominative_pib_display,
    nominative_pib_to_genitive_line,
)
from utils import patient_cards_db as cards_db

app = Flask(__name__)
app.secret_key = SECRET_KEY

# Порожня папка даних — щоб користувач міг одразу покласти Excel
try:
    os.makedirs(DATA_DIR, exist_ok=True)
except OSError:
    pass

# Налаштування логування
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Глобальна змінна для кешування даних
treatments_cache = None
cache_timestamp = None
# Сигнатура файлів для інвалідації кешу після зміни Excel
treatments_cache_file_signature = None
_treatments_load_lock = threading.Lock()
treatments_load_in_progress = False
treatments_last_load_error = None
SERVICE_SIGNATORIES_FILE = os.path.join(DATA_DIR, "service_signatories.json")
VLK_SIGNATORIES_FILE = os.path.join(DATA_DIR, "vlk_signatories.json")
BASE_PERSONNEL_FILE = os.path.join(DATA_DIR, "base.xlsx")
RKCH_PERSONNEL_FILE = os.path.join(DATA_DIR, "rkch.xlsx")

# Кеш base.xlsx: DataFrame + індекс ПІБ → дані (для швидкого автокомпліту)
base_personnel_cache = None
base_personnel_index = None
base_cache_file_signature = None
_base_load_lock = threading.Lock()

# Кеш rkch.xlsx (розпорядження командира частини)
rkch_personnel_index = None
rkch_cache_file_signature = None
_rkch_load_lock = threading.Lock()

def _excel_cell_str(value, default=""):
    """Текст з комірки Excel/pandas без 'nan' у рядку."""
    if value is None:
        return default
    try:
        if pd.isna(value):
            return default
    except (TypeError, ValueError):
        pass
    s = str(value).strip()
    if not s or s.lower() == "nan" or s == "<na>":
        return default
    return s


def load_service_signatories():
    """Повертає збережені поля підписантів для службової характеристики."""
    defaults = {
        "pidpysant_1_zvannya": "",
        "pidpysant_1_pib": "",
        "pidpysant_2_zvannya": "",
        "pidpysant_2_pib": "",
    }
    try:
        if not os.path.isfile(SERVICE_SIGNATORIES_FILE):
            return defaults
        with open(SERVICE_SIGNATORIES_FILE, "r", encoding="utf-8") as f:
            raw = json.load(f)
        if not isinstance(raw, dict):
            return defaults
        for k in defaults.keys():
            defaults[k] = _normalize_spaces(raw.get(k, ""))
        return defaults
    except Exception as e:
        logger.warning("Не вдалося зчитати service_signatories.json: %s", e)
        return defaults


def save_service_signatories(data: dict):
    """Зберігає поля підписантів (атомарний запис)."""
    os.makedirs(DATA_DIR, exist_ok=True)
    payload = {
        "pidpysant_1_zvannya": _normalize_spaces(data.get("pidpysant_1_zvannya", "")),
        "pidpysant_1_pib": _normalize_spaces(data.get("pidpysant_1_pib", "")),
        "pidpysant_2_zvannya": _normalize_spaces(data.get("pidpysant_2_zvannya", "")),
        "pidpysant_2_pib": _normalize_spaces(data.get("pidpysant_2_pib", "")),
        "updated_at": datetime.now().isoformat(timespec="seconds"),
    }
    tmp_path = SERVICE_SIGNATORIES_FILE + ".tmp"
    with open(tmp_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    os.replace(tmp_path, SERVICE_SIGNATORIES_FILE)


def _normalize_signatory_option(item) -> dict:
    if not isinstance(item, dict):
        return {"zvannya": "", "pib": ""}
    return {
        "zvannya": _normalize_spaces(item.get("zvannya", "")),
        "pib": _normalize_spaces(item.get("pib", "")),
    }


def _default_vlk_signatory_options():
    return {
        "kombrig_options": [
            {"zvannya": "полковник", "pib": "Євстахій МОСПАН"},
        ],
        "kombat_options": [
            {"zvannya": "старший лейтенант", "pib": "Кирило МЕЛЕШКО"},
            {"zvannya": "капітан", "pib": "Максим БОРИСЕНКО"},
            {"zvannya": "капітан", "pib": "Олександр ПОРТЯНКО"},
            {"zvannya": "капітан", "pib": "Віталій КОВТУН"},
        ],
    }


def _dedupe_signatory_options(items):
    out = []
    seen = set()
    for raw in items or []:
        opt = _normalize_signatory_option(raw)
        if not opt["zvannya"] or not opt["pib"]:
            continue
        key = (opt["zvannya"].lower(), opt["pib"].lower())
        if key in seen:
            continue
        seen.add(key)
        out.append(opt)
    return out


def load_vlk_signatories():
    """Повертає збережені поля підписантів для рапорту ВЛК + списки опцій."""
    defaults = {
        "tvo_kombrig_zvannya": "",
        "tvo_kombrig_pib": "",
        "tvo_kombat_zvannya": "",
        "tvo_kombat_pib": "",
        "kombrig_options": [],
        "kombat_options": [],
    }
    seeded = _default_vlk_signatory_options()
    try:
        if not os.path.isfile(VLK_SIGNATORIES_FILE):
            defaults.update(seeded)
            defaults["tvo_kombrig_zvannya"] = seeded["kombrig_options"][0]["zvannya"]
            defaults["tvo_kombrig_pib"] = seeded["kombrig_options"][0]["pib"]
            defaults["tvo_kombat_zvannya"] = seeded["kombat_options"][0]["zvannya"]
            defaults["tvo_kombat_pib"] = seeded["kombat_options"][0]["pib"]
            return defaults
        with open(VLK_SIGNATORIES_FILE, "r", encoding="utf-8") as f:
            raw = json.load(f)
        if not isinstance(raw, dict):
            defaults.update(seeded)
            return defaults
        for k in ("tvo_kombrig_zvannya", "tvo_kombrig_pib", "tvo_kombat_zvannya", "tvo_kombat_pib"):
            defaults[k] = _normalize_spaces(raw.get(k, ""))
        defaults["kombrig_options"] = _dedupe_signatory_options(
            raw.get("kombrig_options") or seeded["kombrig_options"]
        )
        defaults["kombat_options"] = _dedupe_signatory_options(
            raw.get("kombat_options") or seeded["kombat_options"]
        )
        if not defaults["kombrig_options"]:
            defaults["kombrig_options"] = list(seeded["kombrig_options"])
        if not defaults["kombat_options"]:
            defaults["kombat_options"] = list(seeded["kombat_options"])
        return defaults
    except Exception as e:
        logger.warning("Не вдалося зчитати vlk_signatories.json: %s", e)
        defaults.update(seeded)
        return defaults


def save_vlk_signatories(data: dict):
    """Зберігає поля підписантів ВЛК та списки опцій (атомарний запис)."""
    os.makedirs(DATA_DIR, exist_ok=True)
    current = load_vlk_signatories()
    kombrig_options = data.get("kombrig_options")
    kombat_options = data.get("kombat_options")
    if kombrig_options is None:
        kombrig_options = current.get("kombrig_options", [])
    if kombat_options is None:
        kombat_options = current.get("kombat_options", [])
    payload = {
        "tvo_kombrig_zvannya": _normalize_spaces(data.get("tvo_kombrig_zvannya", "")),
        "tvo_kombrig_pib": _normalize_spaces(data.get("tvo_kombrig_pib", "")),
        "tvo_kombat_zvannya": _normalize_spaces(data.get("tvo_kombat_zvannya", "")),
        "tvo_kombat_pib": _normalize_spaces(data.get("tvo_kombat_pib", "")),
        "kombrig_options": _dedupe_signatory_options(kombrig_options),
        "kombat_options": _dedupe_signatory_options(kombat_options),
        "updated_at": datetime.now().isoformat(timespec="seconds"),
    }
    tmp_path = VLK_SIGNATORIES_FILE + ".tmp"
    with open(tmp_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    os.replace(tmp_path, VLK_SIGNATORIES_FILE)


def _row_full_name(row) -> str:
    """ПІБ для відображення та пошуку: спочатку колонка ПІБ, інакше з трьох частин."""
    pib = _excel_cell_str(row.get("ПІБ"))
    if pib:
        return pib
    sur = _excel_cell_str(row.get("Прізвище"))
    first = _excel_cell_str(row.get("Ім'я"))
    pat = _excel_cell_str(row.get("По батькові"))
    return " ".join(p for p in (sur, first, pat) if p).strip()


def list_treatments_year_files_sorted():
    """Усі `treatments_YYYY.xlsx` у DATA_DIR, відсортовані за роком зростання."""
    if not os.path.isdir(DATA_DIR):
        return []
    found = []
    for name in os.listdir(DATA_DIR):
        m = TREATMENTS_YEAR_FILE_RE.match(name)
        if not m:
            continue
        year = int(m.group(1))
        path = os.path.join(DATA_DIR, name)
        if os.path.isfile(path):
            found.append((year, path))
    found.sort(key=lambda x: x[0])
    return found


def treatments_path_for_year(year: int) -> str:
    """Шлях до файлу бази за конкретний рік (для завантаження / відображення)."""
    return os.path.join(DATA_DIR, f"treatments_{int(year)}.xlsx")


def _treatments_excel_file_signature():
    """Час модифікації всіх релевантних Excel (для скидання кешу)."""
    paths = []
    if os.path.exists(TREATMENTS_FINAL_FILE):
        paths.append(TREATMENTS_FINAL_FILE)
    for _, p in list_treatments_year_files_sorted():
        paths.append(p)
    return tuple((p, os.path.getmtime(p)) for p in paths)


def _validate_treatments_upload_dataframe(df):
    """Перевірка структури завантаженого файлу перед заміною на диску."""
    if df is None or len(df) < 1:
        return False, "Файл порожній або не містить рядків даних"
    df = df.copy()
    df.columns = df.columns.astype(str).str.strip()
    has_triple = all(c in df.columns for c in ["Прізвище", "Ім'я", "По батькові"])
    has_pib = "ПІБ" in df.columns
    if not has_triple and not has_pib:
        return (
            False,
            "Потрібні колонки «Прізвище», «Ім'я», «По батькові» або колонка «ПІБ»",
        )
    return True, None


def _invalidate_treatments_cache_unlocked():
    """Скинути кеш після зміни файлів (викликати під _treatments_load_lock)."""
    global treatments_cache, cache_timestamp, treatments_cache_file_signature
    treatments_cache = None
    cache_timestamp = None
    treatments_cache_file_signature = None


def load_treatments_data():
    """Завантажує дані з Excel: опційний архів + 2024 + 2025 (об'єднано) + 2026 (поточний рік)."""
    global treatments_cache, cache_timestamp, treatments_cache_file_signature
    global treatments_load_in_progress, treatments_last_load_error

    sig = _treatments_excel_file_signature()
    now = datetime.now()

    if treatments_cache is not None and cache_timestamp is not None:
        cache_fresh = (now - cache_timestamp).total_seconds() < 300  # 5 хвилин
        sig_ok = treatments_cache_file_signature == sig
        if cache_fresh and sig_ok:
            return treatments_cache

    with _treatments_load_lock:
        sig = _treatments_excel_file_signature()
        now = datetime.now()
        if treatments_cache is not None and cache_timestamp is not None:
            cache_fresh = (now - cache_timestamp).total_seconds() < 300
            sig_ok = treatments_cache_file_signature == sig
            if cache_fresh and sig_ok:
                return treatments_cache

        treatments_load_in_progress = True
        treatments_last_load_error = None
        try:
            return _load_treatments_data_unlocked()
        except Exception as e:
            treatments_last_load_error = str(e)
            raise
        finally:
            treatments_load_in_progress = False


def _load_treatments_data_unlocked():
    """Внутрішнє завантаження Excel (викликати лише під _treatments_load_lock)."""
    global treatments_cache, cache_timestamp, treatments_cache_file_signature

    try:
        year_files = list_treatments_year_files_sorted()
        has_final = os.path.exists(TREATMENTS_FINAL_FILE)
        if not year_files and not has_final:
            raise FileNotFoundError(
                f"У папці {DATA_DIR!r} немає файлів treatments_YYYY.xlsx і немає treatments_final.xlsx"
            )

        logger.info(
            "Завантаження Excel: %s + %s",
            "treatments_final.xlsx" if has_final else "(без архіву)",
            ", ".join(f"treatments_{y}.xlsx" for y, _ in year_files) if year_files else "(немає річних файлів)",
        )

        frames = []

        if has_final:
            df_final = pd.read_excel(TREATMENTS_FINAL_FILE)
            df_final.columns = df_final.columns.str.strip()
            logger.info(f"Архів treatments_final.xlsx: {len(df_final)} записів")
            frames.append(df_final)

        for year, path in year_files:
            df_y = pd.read_excel(path)
            df_y.columns = df_y.columns.str.strip()
            logger.info(f"treatments_{year}.xlsx: {len(df_y)} записів")
            frames.append(df_y)

        if not frames:
            raise ValueError("Немає жодного кадру даних для об'єднання")

        treatments_df = pd.concat(frames, ignore_index=True)

        logger.info("Видаляємо дублікати (пріоритет у пізніших джерелів за роком файлу)...")
        initial_count = len(treatments_df)

        treatments_df['ПІБ_чисте'] = (
            treatments_df['Прізвище'].fillna('').astype(str) + ' ' +
            treatments_df['Ім\'я'].fillna('').astype(str) + ' ' +
            treatments_df['По батькові'].fillna('').astype(str)
        ).str.replace(r'\s+', ' ', regex=True).str.strip().str.lower()

        if 'Дата надходження в поточний Л/З' in treatments_df.columns:
            treatments_df['Дата надходження в поточний Л/З'] = pd.to_datetime(
                treatments_df['Дата надходження в поточний Л/З'], errors='coerce'
            )
            treatments_df = treatments_df.drop_duplicates(
                subset=['ПІБ_чисте', 'Дата надходження в поточний Л/З'],
                keep='last',
            )
        else:
            treatments_df = treatments_df.drop_duplicates(subset=['ПІБ_чисте'], keep='last')

        logger.info(f"Видалено дублікатів: {initial_count - len(treatments_df)}; записів після: {len(treatments_df)}")
        
        # Обробка дат
        date_columns = ['Дата надходження в поточний Л/З', 'Дата виписки', 'Дата народження', 'Дата первинної госпіталізації', 'Дата виписки з поточного Л/З']
        for col in date_columns:
            if col in treatments_df.columns:
                treatments_df[col] = pd.to_datetime(treatments_df[col], errors='coerce')
        
        # Обробка ПІБ (якщо ще не створено)
        for col in ['Прізвище', 'Ім\'я', 'По батькові']:
            if col in treatments_df.columns:
                treatments_df[col] = treatments_df[col].fillna('').astype(str)
                treatments_df[col] = treatments_df[col].replace(
                    to_replace=r'(?i)^(nan|<na>|none)$', value='', regex=True
                )

        # Завжди збираємо ПІБ з трьох полів, якщо вони є — так не залишаються порожні NaN з колонки «ПІБ» у Excel
        if all(c in treatments_df.columns for c in ['Прізвище', 'Ім\'я', 'По батькові']):
            treatments_df['ПІБ'] = (
                treatments_df['Прізвище'].astype(str).str.strip() + ' ' +
                treatments_df['Ім\'я'].astype(str).str.strip() + ' ' +
                treatments_df['По батькові'].astype(str).str.strip()
            ).str.replace(r'\s+', ' ', regex=True).str.strip()
        elif 'ПІБ' in treatments_df.columns:
            treatments_df['ПІБ'] = treatments_df['ПІБ'].map(lambda x: _excel_cell_str(x))

        # Після нормалізації ПІБ оновлюємо ключ пошуку
        treatments_df['ПІБ_чисте'] = treatments_df['ПІБ'].str.replace(r'\s+', ' ', regex=True).str.strip().str.lower()
        
        # Оновлюємо кеш (сигнатура файлів — зміни на диску скидають кеш)
        treatments_cache = treatments_df
        cache_timestamp = datetime.now()
        treatments_cache_file_signature = _treatments_excel_file_signature()

        logger.info(f"Дані успішно завантажено. Записів: {len(treatments_df)}")
        return treatments_df
        
    except Exception as e:
        logger.error(f"Помилка при завантаженні даних: {e}")
        raise

def format_rank_genitive(rank):
    """Конвертує звання в родовому відмінку для використання в шапці"""
    if not rank or not rank.strip():
        return ""
    
    rank_base, has_medical_service_suffix = split_medical_service_rank(rank)
    rank_lower = rank_base.lower().strip()
    
    # Словник для перетворення звань в родовому відмінку
    genitive_ranks = {
        'солдат': 'солдата',
        'старший солдат': 'старшого солдата',
        'молодший сержант': 'молодшого сержанта',
        'сержант': 'сержанта',
        'старший сержант': 'старшого сержанта',
        'головний сержант': 'головного сержанта',
        'штаб-сержант': 'штаб-сержанта',
        'майстер-сержант': 'майстер-сержанта',
        'старший майстер-сержант': 'старшого майстер-сержанта',
        'головний майстер-сержант': 'головного майстер-сержанта',
        'молодший лейтенант': 'молодшого лейтенанта',
        'лейтенант': 'лейтенанта',
        'старший лейтенант': 'старшого лейтенанта',
        'капітан': 'капітана',
        'майор': 'майора',
        'підполковник': 'підполковника',
        'полковник': 'полковника',
    }
    
    # Перевіряємо точне співпадіння
    if rank_lower in genitive_ranks:
        formatted_rank = genitive_ranks[rank_lower]
        return ensure_medical_service_rank(formatted_rank) if has_medical_service_suffix else formatted_rank
    
    # Якщо не знайдено точне співпадіння, повертаємо оригінал
    return ensure_medical_service_rank(rank_base) if has_medical_service_suffix else rank_base


def format_rank_dative(rank):
    """Конвертує звання в давальному відмінку для службової характеристики."""
    if not rank or not rank.strip():
        return ""

    rank_base, has_medical_service_suffix = split_medical_service_rank(rank)
    rank_lower = rank_base.lower().strip()
    dative_ranks = {
        'солдат': 'солдату',
        'старший солдат': 'старшому солдату',
        'молодший сержант': 'молодшому сержанту',
        'сержант': 'сержанту',
        'старший сержант': 'старшому сержанту',
        'головний сержант': 'головному сержанту',
        'штаб-сержант': 'штаб-сержанту',
        'майстер-сержант': 'майстер-сержанту',
        'старший майстер-сержант': 'старшому майстер-сержанту',
        'головний майстер-сержант': 'головному майстер-сержанту',
        'молодший лейтенант': 'молодшому лейтенанту',
        'лейтенант': 'лейтенанту',
        'старший лейтенант': 'старшому лейтенанту',
        'капітан': 'капітану',
        'майор': 'майору',
        'підполковник': 'підполковнику',
        'полковник': 'полковнику',
    }
    formatted_rank = dative_ranks.get(rank_lower, rank_base)
    return ensure_medical_service_rank(formatted_rank) if has_medical_service_suffix else formatted_rank


def split_medical_service_rank(rank):
    """Повертає базове звання та ознаку наявності приставки медичної служби."""
    normalized_rank = _normalize_spaces(rank)
    if not normalized_rank:
        return "", False

    has_medical_service_suffix = bool(re.search(r'\s+м[\\/]+с$', normalized_rank, flags=re.IGNORECASE))
    if not has_medical_service_suffix:
        return normalized_rank, False

    base_rank = re.sub(r'\s+м[\\/]+с$', '', normalized_rank, flags=re.IGNORECASE).strip()
    return base_rank, True


def ensure_medical_service_rank(rank):
    """Нормалізує звання з приставкою медичної служби до формату `м/с`."""
    base_rank, has_medical_service_suffix = split_medical_service_rank(rank)
    if not base_rank:
        return ""
    if has_medical_service_suffix:
        return f"{base_rank} м/с"
    return f"{base_rank} м/с"


def _normalize_spaces(text: str) -> str:
    return re.sub(r'\s+', ' ', (text or '')).strip()


def _loader_preview_delay():
    """Пауза перед віддачею файлу (для перевірки лоадера)."""
    delay = float(globals().get("LOADER_PREVIEW_DELAY_SEC", 0) or 0)
    if delay > 0:
        time.sleep(delay)


def _safe_download_stem(text: str, fallback: str = "document") -> str:
    """Ім'я файлу для download: лишає кирилицю, прибирає небезпечні символи."""
    cleaned = _normalize_spaces(text).replace(" ", "_")
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "", cleaned)
    cleaned = cleaned.strip(" ._")
    return cleaned or fallback


def _ensure_final_period(text: str) -> str:
    t = _normalize_spaces(text).rstrip()
    if not t:
        return ""
    t = t.rstrip(' ,;:')
    if t.endswith('.'):
        return t
    return t + '.'


def _title_first_letter(text: str) -> str:
    t = _normalize_spaces(text)
    if not t:
        return ""
    return t[0].upper() + t[1:]


def _format_komisariat_and_date(raw_value: str) -> str:
    """
    Нормалізація рядка "комісаріат + дата":
    - обов'язково містить дату дд.мм.рррр;
    - додає кому перед датою;
    - завершує словом "року.".
    """
    value = _normalize_spaces(raw_value)
    m = re.search(r'(\d{2}\.\d{2}\.\d{4})', value)
    if not m:
        return ""

    date_value = m.group(1)
    if not validate_date_format(date_value):
        return ""

    before = _normalize_spaces(value[:m.start()])
    if not before:
        return ""
    before = before.rstrip(' ,.;:')
    return f"{before}, {date_value} року."


def _format_komisariat_parts(komisariat: str, data_pryzovu: str) -> str:
    """Збирає рядок комісаріату з окремих полів форми."""
    komisariat = _normalize_spaces(komisariat).rstrip(' ,.;:')
    data_pryzovu = _normalize_spaces(data_pryzovu)
    if not komisariat or not data_pryzovu:
        return ""
    if not validate_date_format(data_pryzovu):
        return ""
    return f"{komisariat}, {data_pryzovu} року."


def _format_osvita_parts(
    osvita_type: str,
    navchalnyy_zaklad: str,
    misto: str,
    rik: str,
) -> str:
    """
    Збирає рядок освіти для шаблону, напр.:
    «середня, ЗОШ №276 м. Баку, у 1996 році.»
    """
    osvita_type = _normalize_spaces(osvita_type).lower()
    zaklad = _normalize_spaces(navchalnyy_zaklad)
    misto = _normalize_spaces(misto)
    rik = _normalize_spaces(rik)
    if not osvita_type or not zaklad or not misto or not rik:
        return ""
    if not re.fullmatch(r'\d{4}', rik):
        return ""
    year_now = datetime.now().year
    year_int = int(rik)
    if year_int < 1950 or year_int > year_now:
        return ""
    misto_part = misto if misto.lower().startswith("м.") or misto.lower().startswith("м ") else f"м. {misto}"
    line = f"{osvita_type}, {zaklad} {misto_part}, у {rik} році"
    return _ensure_final_period(_title_first_letter(line))


def _split_pib(full_name: str):
    parts = _normalize_spaces(full_name).split(' ')
    if len(parts) < 3:
        return "", "", ""
    return parts[0], parts[1], " ".join(parts[2:])


def _format_initials_name(prizvyshche: str, imya: str, po_batkovi: str) -> str:
    if not prizvyshche:
        return ""
    i = (imya[:1] + ".") if imya else ""
    p = (po_batkovi[:1] + ".") if po_batkovi else ""
    return _normalize_spaces(f"{prizvyshche} {i}{p}")


def _surname_first_capital(text: str) -> str:
    """Формат прізвища: лише перша літера велика."""
    t = _normalize_spaces(text)
    if not t:
        return ""
    t = t.lower()
    return t[:1].upper() + t[1:]


def _surname_to_dative(surname_nominative: str) -> str:
    """
    Наближене перетворення прізвища з називного у давальний відмінок.
    Використовується для службової характеристики (формат ініціалів).
    """
    s = _normalize_spaces(surname_nominative).lower()
    if not s:
        return ""

    # Глодний -> Глодному, Синій -> Синьому
    if s.endswith("ий"):
        return _surname_first_capital(s[:-2] + "ому")
    if s.endswith("ій"):
        return _surname_first_capital(s[:-2] + "ьому")

    # Ковальчук -> Ковальчуку, Шевченко -> Шевченку, Мелешко -> Мелешку, Іванов -> Іванову
    if s.endswith("ко"):
        return _surname_first_capital(s[:-1] + "у")
    if s.endswith(("ов", "ев", "єв", "ин", "ін", "їн", "ук", "юк", "чук", "чак", "ак", "як", "ич", "ець", "єць", "ань", "ень")):
        return _surname_first_capital(s + "у")

    # Для жіночих/м'яких форм на -а/-я -> -і
    if s.endswith("а"):
        return _surname_first_capital(s[:-1] + "і")
    if s.endswith("я"):
        return _surname_first_capital(s[:-1] + "і")

    # Базовий варіант для більшості чоловічих прізвищ на приголосний
    return _surname_first_capital(s + "у")


def _first_name_to_dative(name_nominative: str) -> str:
    """Наближене перетворення імені у давальний відмінок."""
    n = _normalize_spaces(name_nominative).lower()
    if not n:
        return ""
    if n.endswith("ій"):
        return _surname_first_capital(n[:-2] + "ію")
    if n.endswith("й"):
        return _surname_first_capital(n[:-1] + "ю")
    if n.endswith("а"):
        return _surname_first_capital(n[:-1] + "і")
    if n.endswith("я"):
        return _surname_first_capital(n[:-1] + "і")
    if n.endswith("о"):
        return _surname_first_capital(n[:-1] + "у")
    if n.endswith("р"):
        return _surname_first_capital(n + "у")
    return _surname_first_capital(n + "у")


def _pib_to_dative(full_name: str) -> str:
    """Наближене перетворення ПІБ у давальний відмінок."""
    parts = _normalize_spaces(full_name).split(" ")
    if not parts:
        return ""
    if len(parts) == 1:
        return _surname_to_dative(parts[0])
    if len(parts) == 2:
        first, surname = parts
        return _normalize_spaces(f"{_first_name_to_dative(first)} {_surname_to_dative(surname)}")
    first, surname, patronymic = parts[0], parts[1], " ".join(parts[2:])
    pat = patronymic
    if patronymic.lower().endswith("ич"):
        pat = patronymic + "у"
    elif patronymic.lower().endswith("на"):
        pat = patronymic[:-1] + "і"
    return _normalize_spaces(f"{_first_name_to_dative(first)} {_surname_to_dative(surname)} {pat}")


def _uppercase_last_word(full_name: str) -> str:
    """Робить останнє слово (прізвище) у ПІБ ВЕЛИКИМИ літерами."""
    parts = _normalize_spaces(full_name).split(" ")
    if not parts:
        return ""
    parts[-1] = parts[-1].upper()
    return _normalize_spaces(" ".join(parts))


def _split_rank_and_name(full_value: str):
    """Розбиває рядок виду 'підполковник Олександр РЯСНИЙ' на (звання, ПІБ)."""
    text = _normalize_spaces(full_value)
    if not text:
        return "", ""
    known_ranks = [
        "головний майстер-сержант",
        "старший майстер-сержант",
        "майстер-сержант",
        "головний сержант",
        "старший сержант",
        "молодший сержант",
        "старший лейтенант",
        "молодший лейтенант",
        "старший солдат",
        "штаб-сержант",
        "підполковник",
        "полковник",
        "лейтенант",
        "капітан",
        "сержант",
        "солдат",
        "майор",
    ]
    lower = text.lower()
    for rank in known_ranks:
        if lower.startswith(rank + " "):
            return text[:len(rank)], text[len(rank):].strip()
    parts = text.split(" ", 1)
    return (parts[0], parts[1]) if len(parts) > 1 else (parts[0], "")


def _rank_short(rank_value: str) -> str:
    """Скорочення звання для блоку НМС."""
    rank_base, _ = split_medical_service_rank(rank_value)
    short_map = {
        "полковник": "п-к",
        "підполковник": "п/п-к",
        "майор": "м-р",
        "капітан": "к-н",
        "старший лейтенант": "ст. л-т",
        "лейтенант": "л-т",
        "молодший лейтенант": "мол. л-т",
        "старший сержант": "ст. с-т",
        "сержант": "с-т",
        "молодший сержант": "мол. с-т",
        "солдат": "с-т",
    }
    return short_map.get(rank_base.lower().strip(), rank_base)


def _pib_short_with_initials(full_name: str) -> str:
    """Формат 'О. РЯСНИЙ' з повного імені/ПІБ."""
    parts = _normalize_spaces(full_name).split(" ")
    if not parts:
        return ""
    if len(parts) == 1:
        return parts[0].upper()
    first_name = parts[0]
    surname = parts[-1].upper()
    return _normalize_spaces(f"{first_name[:1].upper()}. {surname}")


def _extract_position_from_row(row) -> str:
    for col in (
        "Займана посада",
        "Посада",
        "Військова посада",
        "Штатна посада",
    ):
        val = _excel_cell_str(row.get(col))
        if val:
            return val
    return ""


def _to_genitive_unit_phrase(phrase: str) -> str:
    """
    Груба нормалізація назви підрозділу в родовий відмінок для шаблонного тексту посади.
    Приклади:
    - "1-ий взвод оперативного призначення" -> "1-го взводу оперативного призначення"
    - "Розвідувальний взвод спеціального призначення" -> "розвідувального взводу спеціального призначення"
    """
    text = _normalize_spaces(phrase)
    if not text:
        return ""

    words = text.split(" ")
    adjective_idx = 0
    noun_idx = 1 if len(words) > 1 else 0

    # 1-ий / 1-ша / 1-ше -> 1-го / 1-ої
    ordinal_re = re.compile(r'^(\d+)\s*-\s*([а-яіїєґa-z]+)$', flags=re.IGNORECASE)
    m = ordinal_re.match(words[0]) if words else None
    if m:
        # Варіанти:
        # - "1-ий взвод ..." (числівник + іменник)
        # - "1-ше розвідувальне відділення ..." (числівник + прикметник + іменник)
        second_lower = words[1].lower() if len(words) > 1 else ""
        second_looks_like_adj = second_lower.endswith(("ий", "ій", "е", "є"))
        if len(words) > 2 and second_looks_like_adj:
            adjective_idx = 1
            noun_idx = 2
        else:
            adjective_idx = -1
            noun_idx = 1 if len(words) > 1 else 0

        noun_for_gender = words[noun_idx].lower() if len(words) > noun_idx else ""
        number = m.group(1)
        fem_nouns = ("рота", "група", "команда", "батарея")
        ordinal_suffix = "ої" if noun_for_gender.startswith(fem_nouns) else "го"
        words[0] = f"{number}-{ordinal_suffix}"

    # Нормалізація прикметника
    if 0 <= adjective_idx < len(words):
        adj_lower = words[adjective_idx].lower()
        if adj_lower.endswith("ий"):
            words[adjective_idx] = words[adjective_idx][:-2] + "ого"
        elif adj_lower.endswith("ій"):
            words[adjective_idx] = words[adjective_idx][:-2] + "ього"
        elif adj_lower.endswith("а"):
            words[adjective_idx] = words[adjective_idx][:-1] + "ої"
        elif adj_lower.endswith("я"):
            words[adjective_idx] = words[adjective_idx][:-1] + "ьої"
        elif adj_lower.endswith("е"):
            words[adjective_idx] = words[adjective_idx][:-1] + "ого"
        elif adj_lower.endswith("є"):
            words[adjective_idx] = words[adjective_idx][:-1] + "ього"

    # Нормалізація іменника (друге слово у фразі)
    if 0 <= noun_idx < len(words):
        noun = words[noun_idx]
        noun_map = {
            "взвод": "взводу",
            "батальйон": "батальйону",
            "полк": "полку",
            "дивізіон": "дивізіону",
            "пункт": "пункту",
            "відділення": "відділення",
            "рота": "роти",
            "група": "групи",
            "команда": "команди",
            "служба": "служби",
        }
        noun_lower = noun.lower()
        words[noun_idx] = noun_map.get(noun_lower, noun)

    out = " ".join(words)
    return out[:1].lower() + out[1:] if out else ""


def _extract_position_from_base_row(row) -> str:
    """
    Формує шаблонний текст посади з base.xlsx:
    I (Посада) + F (Підрозділ 4) + E (Підрозділ 3) + D (Підрозділ 2, якщо є).
    """
    try:
        posada_i = _excel_cell_str(row.iloc[8])
        pidrozdil_f = _excel_cell_str(row.iloc[5])
        pidrozdil_e = _excel_cell_str(row.iloc[4])
        pidrozdil_d = _excel_cell_str(row.iloc[3])
    except Exception:
        return ""

    if not posada_i:
        return ""

    chunks = [posada_i]
    for part in (pidrozdil_f, pidrozdil_e, pidrozdil_d):
        norm = _to_genitive_unit_phrase(part)
        if norm:
            chunks.append(norm)
    return _normalize_spaces(" ".join(chunks))


def _extract_rank_from_base_row(row) -> str:
    """Звання з base.xlsx: колонка M (фактичне), запасний варіант — L (по штату)."""
    try:
        rank_l = _excel_cell_str(row.iloc[11])
        rank_m = _excel_cell_str(row.iloc[12])
    except Exception:
        return ""
    return rank_m or rank_l


def _empty_base_person() -> dict:
    return {
        "zvanie": "",
        "zaimana_posada": "",
        "prizvyshche": "",
        "imya": "",
        "po_batkovi": "",
        "birth_date": "",
        "unit_short": "",
        "service_category": "",
    }


def _format_base_birth_date(value) -> str:
    """Дата народження з base.xlsx → дд.мм.рррр."""
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    if hasattr(value, "strftime"):
        try:
            return value.strftime("%d.%m.%Y")
        except Exception:
            pass
    raw = _excel_cell_str(value)
    if not raw:
        return ""
    parsed = cards_db.parse_ui_date(raw)
    if parsed:
        return cards_db.format_ui_date(parsed)
    # ISO-подібне «1983-02-20 00:00:00»
    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})", raw)
    if m:
        return f"{m.group(3)}.{m.group(2)}.{m.group(1)}"
    return raw


def _base_excel_file_signature():
    """mtime + size base.xlsx для інвалідації кешу."""
    if not os.path.isfile(BASE_PERSONNEL_FILE):
        return None
    try:
        st = os.stat(BASE_PERSONNEL_FILE)
        return (BASE_PERSONNEL_FILE, st.st_mtime, st.st_size)
    except OSError:
        return None


def _rkch_excel_file_signature():
    """mtime + size rkch.xlsx для інвалідації кешу."""
    if not os.path.isfile(RKCH_PERSONNEL_FILE):
        return None
    try:
        st = os.stat(RKCH_PERSONNEL_FILE)
        return (RKCH_PERSONNEL_FILE, st.st_mtime, st.st_size)
    except OSError:
        return None


def _format_mtime_display(path):
    """Дата/час модифікації файлу у форматі дд.мм.рррр гг:хх."""
    if not path or not os.path.isfile(path):
        return None
    try:
        return datetime.fromtimestamp(os.path.getmtime(path)).strftime("%d.%m.%Y %H:%M")
    except OSError:
        return None


def _mtime_iso(path):
    if not path or not os.path.isfile(path):
        return None
    try:
        return datetime.fromtimestamp(os.path.getmtime(path)).isoformat(timespec="seconds")
    except OSError:
        return None


def _invalidate_base_cache_unlocked():
    """Скинути кеш base.xlsx після заміни файлу (викликати під _base_load_lock)."""
    global base_personnel_cache, base_personnel_index, base_cache_file_signature
    base_personnel_cache = None
    base_personnel_index = None
    base_cache_file_signature = None


def _validate_base_upload_dataframe(df, reference_df=None):
    """
    Перевірка структури нового base.xlsx перед заміною.
    Якщо є поточний файл — кількість і назви колонок мають збігатися.
    """
    if df is None or len(df) < 1:
        return False, "Файл порожній або не містить рядків даних"
    if len(df.columns) < 16:
        return (
            False,
            "У файлі має бути щонайменше 16 колонок (A–P), як у поточному base.xlsx",
        )

    new_cols = [str(c).strip() for c in df.columns]
    if reference_df is not None and len(reference_df.columns) > 0:
        ref_cols = [str(c).strip() for c in reference_df.columns]
        if len(new_cols) != len(ref_cols):
            return (
                False,
                f"Кількість колонок не збігається з поточним base.xlsx "
                f"(було {len(ref_cols)}, у новому файлі {len(new_cols)})",
            )
        if new_cols != ref_cols:
            mismatches = [
                f"«{a}» → «{b}»"
                for a, b in zip(ref_cols, new_cols)
                if a != b
            ][:5]
            detail = "; ".join(mismatches) if mismatches else "відмінність у назвах"
            return (
                False,
                "Назви колонок не збігаються з поточним base.xlsx. "
                f"Перші відмінності: {detail}",
            )

    index = _build_base_personnel_index(df)
    if len(index) < 1:
        return (
            False,
            "Не знайдено жодного ПІБ у колонках N–P (прізвище, ім'я, по батькові). "
            "Перевірте, що структура файлу така сама, як у попереднього base.xlsx",
        )
    return True, None


def _build_base_personnel_index(base_df) -> dict:
    """Побудова dict[ПІБ_чисте] → дані персони з рядка base.xlsx."""
    index = {}
    if base_df is None or len(base_df.columns) < 16:
        return index
    for _, row in base_df.iterrows():
        sur = _excel_cell_str(row.iloc[13])
        first = _excel_cell_str(row.iloc[14])
        pat = _excel_cell_str(row.iloc[15])
        row_pib = _normalize_spaces(f"{sur} {first} {pat}").lower()
        if not row_pib:
            continue
        birth_raw = row["Дата народження"] if "Дата народження" in row.index else ""
        unit_short = ""
        if "Підрозділ (скорочено)" in row.index:
            unit_short = _excel_cell_str(row["Підрозділ (скорочено)"])
        service_raw = ""
        if "Вид служби" in row.index:
            service_raw = _excel_cell_str(row["Вид служби"])
        elif "вид служби (скорочено)" in row.index:
            service_raw = _excel_cell_str(row["вид служби (скорочено)"])
        index[row_pib] = {
            "zvanie": _extract_rank_from_base_row(row),
            "zaimana_posada": _extract_position_from_base_row(row),
            "prizvyshche": sur,
            "imya": first,
            "po_batkovi": pat,
            "birth_date": _format_base_birth_date(birth_raw),
            "unit_short": unit_short,
            "service_category": cards_db.normalize_service_category(service_raw),
        }
    return index


def load_base_personnel_data():
    """
    Завантажує base.xlsx у пам'ять і будує індекс за ПІБ.
    Перечитує файл лише коли змінилась signature (mtime/size).
    """
    global base_personnel_cache, base_personnel_index, base_cache_file_signature

    sig = _base_excel_file_signature()
    if (
        base_personnel_index is not None
        and base_personnel_cache is not None
        and base_cache_file_signature == sig
        and sig is not None
    ):
        return base_personnel_cache

    with _base_load_lock:
        sig = _base_excel_file_signature()
        if (
            base_personnel_index is not None
            and base_personnel_cache is not None
            and base_cache_file_signature == sig
            and sig is not None
        ):
            return base_personnel_cache

        if sig is None:
            base_personnel_cache = None
            base_personnel_index = {}
            base_cache_file_signature = None
            return None

        try:
            base_df = pd.read_excel(BASE_PERSONNEL_FILE)
        except Exception as e:
            logger.warning("Не вдалося прочитати base.xlsx: %s", e)
            base_personnel_cache = None
            base_personnel_index = {}
            base_cache_file_signature = None
            return None

        index = _build_base_personnel_index(base_df)
        base_personnel_cache = base_df
        base_personnel_index = index
        base_cache_file_signature = sig
        logger.info("Кеш base.xlsx оновлено: %s записів у індексі", len(index))
        return base_personnel_cache


def load_base_personnel_index() -> dict:
    """Повертає індекс ПІБ → дані з base.xlsx (з кешем)."""
    load_base_personnel_data()
    return base_personnel_index if base_personnel_index is not None else {}


def _invalidate_rkch_cache_unlocked():
    global rkch_personnel_index, rkch_cache_file_signature
    rkch_personnel_index = None
    rkch_cache_file_signature = None


def _build_rkch_personnel_index(rkch_df) -> dict:
    """dict[ПІБ_чисте] → дані з rkch.xlsx (останній рядок за ПІБ)."""
    index = {}
    if rkch_df is None or getattr(rkch_df, "empty", True):
        return index
    df = rkch_df.copy()
    df.columns = [str(c).strip() for c in df.columns]
    for _, row in df.iterrows():
        sur = _excel_cell_str(row.get("Прізвище"))
        first = _excel_cell_str(row.get("Ім'я"))
        pat = _excel_cell_str(
            row.get("Ім'я по батькові")
            or row.get("По батькові")
            or ""
        )
        full = _normalize_spaces(f"{sur} {first} {pat}")
        if not full:
            full = _excel_cell_str(row.get("ПІБ"))
        if not full:
            continue
        if not (sur and first):
            sur2, first2, pat2 = _split_pib(full)
            sur = sur or sur2
            first = first or first2
            pat = pat or pat2
            full = _normalize_spaces(f"{sur} {first} {pat}") or full
        key = re.sub(r"\s+", " ", full).strip().casefold()
        unit_short = _excel_cell_str(row.get("Підрозділ (скорочено)"))
        vidryv = _excel_cell_str(
            row.get("Відрив") or row.get("Відрив (скорочено)") or ""
        )
        index[key] = {
            "pib": full,
            "zvanie": _excel_cell_str(row.get("Військове звання")),
            "zaimana_posada": "У розпорядженні командира військової частини",
            "prizvyshche": sur,
            "imya": first,
            "po_batkovi": pat,
            "unit_short": unit_short,
            "vidryv": vidryv,
            "service_category": cards_db.normalize_service_category(
                _excel_cell_str(row.get("Вид служби"))
            ),
            "source": "rkch",
        }
    return index


def load_rkch_personnel_index() -> dict:
    """
    Індекс людей у розпорядженні командира частини (data/rkch.xlsx).
    Перечитує файл лише при зміні signature.
    """
    global rkch_personnel_index, rkch_cache_file_signature
    sig = _rkch_excel_file_signature()
    if (
        rkch_personnel_index is not None
        and rkch_cache_file_signature == sig
    ):
        return rkch_personnel_index
    with _rkch_load_lock:
        sig = _rkch_excel_file_signature()
        if (
            rkch_personnel_index is not None
            and rkch_cache_file_signature == sig
        ):
            return rkch_personnel_index
        if not sig:
            rkch_personnel_index = {}
            rkch_cache_file_signature = None
            return rkch_personnel_index
        try:
            df = pd.read_excel(RKCH_PERSONNEL_FILE, sheet_name=0)
            rkch_personnel_index = _build_rkch_personnel_index(df)
            rkch_cache_file_signature = sig
            logger.info(
                "Завантажено rkch.xlsx: %s записів ПІБ",
                len(rkch_personnel_index),
            )
        except Exception as e:
            logger.warning("Не вдалося прочитати rkch.xlsx: %s", e)
            rkch_personnel_index = {}
            rkch_cache_file_signature = None
        return rkch_personnel_index


def _search_rkch_results(q_clean: str, limit: int = 10) -> list:
    """Автокомпліт за rkch.xlsx: спочатку prefix, потім contains."""
    index = load_rkch_personnel_index()
    if not index or not q_clean:
        return []
    prefix_hits = []
    contains_hits = []
    for pib_clean, person in index.items():
        if not pib_clean:
            continue
        if pib_clean.startswith(q_clean):
            prefix_hits.append(person)
        elif q_clean in pib_clean:
            contains_hits.append(person)
    matched = (prefix_hits + contains_hits)[:limit]
    results = []
    for person in matched:
        full_name = person.get("pib") or _normalize_spaces(
            f"{person.get('prizvyshche', '')} {person.get('imya', '')} {person.get('po_batkovi', '')}"
        )
        if not full_name:
            continue
        rank = person.get("zvanie", "")
        unit = person.get("unit_short", "")
        position = person.get("zaimana_posada", "")
        if unit:
            position = f"{position} ({unit})" if position else unit
        label = full_name
        if rank:
            label = f"{full_name} ({rank})"
        label = f"{label} · РКЧ"
        results.append({
            "label": label,
            "value": full_name,
            "rank": rank,
            "position": position,
            "prizvyshche": person.get("prizvyshche", ""),
            "imya": person.get("imya", ""),
            "po_batkovi": person.get("po_batkovi", ""),
            "birth_date": "",
            "unit_short": unit,
            "source": "rkch",
        })
    return results


def _search_treatments_results(q_clean: str, limit: int = 10) -> list:
    """Автокомпліт за treatments_*.xlsx (злита база лікувань)."""
    treatments_df = load_treatments_data()
    if (
        treatments_df is None
        or "ПІБ_чисте" not in treatments_df.columns
        or "ПІБ" not in treatments_df.columns
    ):
        return []
    matched = treatments_df[
        treatments_df["ПІБ_чисте"].str.contains(q_clean, na=False)
    ]
    if matched.empty:
        return []
    unique_patients = matched.drop_duplicates(
        subset=["ПІБ_чисте"], keep="first"
    ).head(limit)
    results = []
    for _, row in unique_patients.iterrows():
        full_name = _row_full_name(row)
        if not full_name:
            continue
        rank = _excel_cell_str(row.get("Військове звання"))
        position = _extract_position_from_row(row)
        birth_date = row.get("Дата народження")
        sur = _excel_cell_str(row.get("Прізвище"))
        first = _excel_cell_str(row.get("Ім'я"))
        pat = _excel_cell_str(row.get("По батькові"))
        if not (sur and first and pat):
            sur, first, pat = _split_pib(full_name)
        birth_date_str = ""
        try:
            if pd.notna(birth_date):
                birth_date_str = (
                    birth_date.strftime("%d.%m.%Y")
                    if hasattr(birth_date, "strftime")
                    else _excel_cell_str(birth_date)
                )
        except Exception:
            birth_date_str = ""
        label = full_name
        if rank:
            label = f"{full_name} ({rank})"
        results.append({
            "label": label,
            "value": full_name,
            "rank": rank,
            "position": position,
            "prizvyshche": sur,
            "imya": first,
            "po_batkovi": pat,
            "birth_date": birth_date_str,
            "source": "treatments",
        })
    return results


def _lookup_person_in_base_excel(pib_nazivnyi: str) -> dict:
    """
    Пошук персональних даних за ПІБ у data/base.xlsx (через кешований індекс).
    Повертає словник: zvanie, zaimana_posada, prizvyshche, imya, po_batkovi,
    birth_date, unit_short, service_category.
    """
    empty = _empty_base_person()
    if not pib_nazivnyi:
        return empty
    target_clean = re.sub(r'\s+', ' ', pib_nazivnyi).strip().lower()
    if not target_clean:
        return empty
    person = load_base_personnel_index().get(target_clean)
    if not person:
        return empty
    merged = dict(empty)
    merged.update(person)
    return merged


def _lookup_person_in_treatments_excel(pib_nazivnyi: str) -> dict:
    """
    Телефон / категорія / дата народження / діагноз з Бази лікувань.
    Бере останній запис за ПІБ.
    """
    empty = {
        "phone": "",
        "birth_date": "",
        "service_category": "",
        "diagnosis": "",
        "kategoriia": "",
    }
    target = _normalize_spaces(pib_nazivnyi).casefold()
    if not target:
        return empty
    try:
        treatments_df = load_treatments_data()
    except Exception as e:
        logger.warning("База лікувань недоступна для картки: %s", e)
        return empty
    if treatments_df is None or "ПІБ_чисте" not in treatments_df.columns:
        return empty
    matched = treatments_df[treatments_df["ПІБ_чисте"] == target]
    if matched.empty:
        return empty
    row = matched.iloc[-1]
    phone = ""
    if "Номер телефону" in matched.columns:
        phone = _excel_cell_str(row.get("Номер телефону"))
        phone = re.sub(r"[\r\n]+", "; ", phone)
        phone = _normalize_spaces(phone).strip(" ;")
    birth_raw = row.get("Дата народження") if "Дата народження" in matched.columns else ""
    kategoriia = ""
    if "Категорія" in matched.columns:
        kategoriia = _excel_cell_str(row.get("Категорія"))
    diagnosis = ""
    for col in ("Заключний діагноз", "Попередній діагноз"):
        if col in matched.columns:
            diagnosis = _excel_cell_str(row.get(col))
            if diagnosis:
                break
    return {
        "phone": phone,
        "birth_date": _format_base_birth_date(birth_raw),
        "service_category": cards_db.normalize_service_category(kategoriia),
        "kategoriia": kategoriia,
        "diagnosis": diagnosis,
    }


def _sync_patient_from_base(patient_id: int, pib: str = "", *, overwrite: bool = False) -> bool:
    """Підтягує звання/посаду/підрозділ/дату народження/категорію з base.xlsx у картку."""
    patient = cards_db.get_patient(patient_id)
    if not patient:
        return False
    lookup_pib = pib or patient.get("pib") or ""
    person = _lookup_person_in_base_excel(lookup_pib)
    changed = False
    if any(person.get(k) for k in ("zvanie", "zaimana_posada", "unit_short", "birth_date", "service_category")):
        changed = cards_db.apply_base_fields_to_patient(patient_id, person, overwrite=overwrite) or changed
    return changed


def _sync_patient_from_treatments(patient_id: int, pib: str = "", *, overwrite: bool = False) -> bool:
    """Підтягує телефон / категорію / ДН з Бази лікувань."""
    patient = cards_db.get_patient(patient_id)
    if not patient:
        return False
    lookup_pib = pib or patient.get("pib") or ""
    person = _lookup_person_in_treatments_excel(lookup_pib)
    if not any(person.get(k) for k in ("phone", "birth_date", "service_category")):
        return False
    return cards_db.apply_treatments_fields_to_patient(patient_id, person, overwrite=overwrite)


def _sync_patient_card(patient_id: int, pib: str = "", *, overwrite: bool = False) -> bool:
    """Повний sync картки: кадрова база + база лікувань."""
    changed = _sync_patient_from_base(patient_id, pib, overwrite=overwrite)
    changed = _sync_patient_from_treatments(patient_id, pib, overwrite=overwrite) or changed
    return changed


def _ensure_patients_from_base(*, enrich_treatments: bool = True) -> dict:
    """
    Створює картки на всіх в/с з base.xlsx, яких ще немає.
    Повертає статистику: created, updated, total_base, total_patients.
    """
    index = load_base_personnel_index()
    created = 0
    updated = 0
    for pib_key, person in index.items():
        sur = person.get("prizvyshche") or ""
        first = person.get("imya") or ""
        pat = person.get("po_batkovi") or ""
        pib = _normalize_spaces(f"{sur} {first} {pat}")
        if not pib:
            # ключ уже нормалізований lower — відновити з person частин
            continue
        existing_id = cards_db.find_patient_id_by_pib(pib)
        if existing_id:
            pid = existing_id
        else:
            rank_unit = ", ".join(
                x for x in (person.get("zvanie") or "", person.get("zaimana_posada") or "") if x
            )
            pid = cards_db.get_or_create_patient(pib, rank_unit)
            created += 1
        if _sync_patient_from_base(pid, pib, overwrite=False):
            updated += 1
        if enrich_treatments and _sync_patient_from_treatments(pid, pib, overwrite=False):
            updated += 1
    return {
        "created": created,
        "updated": updated,
        "total_base": len(index),
        "total_patients": cards_db.count_patients(),
    }


def _enrich_patient_for_whatsapp(patient: dict) -> dict:
    """Доповнює картку категорією з base.xlsx для WhatsApp."""
    if not patient:
        return patient
    out = dict(patient)
    if cards_db.normalize_service_category(out.get("service_category") or ""):
        return out
    person = _lookup_person_in_base_excel(out.get("pib") or "")
    cat = cards_db.normalize_service_category(person.get("service_category") or "")
    if cat:
        out["service_category"] = cat
    return out


def _patient_for_whatsapp(patient_id: int = None, pib: str = "", rank_unit: str = "") -> dict:
    """Гарантує картку пацієнта + sync з Excel, повертає dict для WA."""
    pid = int(patient_id) if patient_id else None
    if not pid and pib:
        pid = cards_db.get_or_create_patient(pib, rank_unit)
        _sync_patient_card(pid, pib)
    patient = cards_db.get_patient(pid) if pid else None
    if not patient and pib:
        patient = {"pib": pib, "unit_short": "", "birth_date": "", "phone": ""}
    return _enrich_patient_for_whatsapp(patient)


def _whatsapp_for_visit(visit: dict, *, event: str = None) -> str:
    """Текст WhatsApp для звернення журналу."""
    if not visit:
        return ""
    resolved_event = event if event is not None else cards_db.wa_event_for_visit(visit)
    patient = None
    tid = visit.get("treatment_id")
    if tid:
        try:
            t = cards_db.get_treatment(int(tid))
            if t and t.get("patient_id"):
                patient = cards_db.get_patient(int(t["patient_id"]))
        except Exception:
            patient = None
    if not patient:
        patient = _patient_for_whatsapp(
            pib=visit.get("pib") or "",
            rank_unit=visit.get("rank_unit") or "",
        )
    else:
        # підхопити звання з журналу, якщо в картці порожньо
        patient = dict(patient)
        if not _normalize_spaces(patient.get("rank") or ""):
            if visit.get("rank_unit"):
                patient["rank_unit"] = visit.get("rank_unit") or patient.get("rank_unit") or ""
    patient = _enrich_patient_for_whatsapp(patient)
    prev_lpz = ""
    if visit.get("treatment_id"):
        try:
            tid_i = int(visit.get("treatment_id"))
            vid_i = int(visit.get("id") or 0)
            for prev in reversed(cards_db.list_visits_for_treatment(tid_i)):
                if int(prev.get("id") or 0) < vid_i:
                    prev_lpz = prev.get("lpz") or ""
                    break
        except (TypeError, ValueError):
            prev_lpz = ""

    action = cards_db.wa_action_phrase(
        resolved_event,
        care_type=visit.get("care_type") or "",
        lpz=visit.get("lpz") or "",
        date=visit.get("date") or "",
        leave_start=visit.get("leave_start") or "",
        leave_end=visit.get("leave_end") or "",
        specialist=visit.get("referral_to") or "",
        from_lpz=visit.get("from_lpz") or prev_lpz,
    )
    canonical = _parse_outpatient_unit(visit.get("rank_unit") or "")
    return cards_db.build_whatsapp_message(
        patient,
        action=action,
        diagnosis=visit.get("diagnosis") or "",
        recommendations=visit.get("exam_result") or visit.get("notes") or "",
        canonical_unit=canonical,
    )


def _lookup_position_in_base_excel(pib_nazivnyi: str) -> str:
    """Пошук посади за ПІБ у data/base.xlsx."""
    return _lookup_person_in_base_excel(pib_nazivnyi).get("zaimana_posada", "")


def _iter_doc_paragraphs_for_replace(doc: DocxDocument):
    """Ітерує всі абзаци документа, включно з таблицями."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _extract_vlk_highlight_samples(template_path: str):
    """
    Повертає список жовтих фрагментів з шаблону у порядку згори-вниз.
    Кожна послідовність підсвічених run-ів вважається одним динамічним полем.
    """
    if not os.path.isfile(template_path):
        return []
    doc = DocxDocument(template_path)
    samples = []
    for p in _iter_doc_paragraphs_for_replace(doc):
        seq = []
        for run in p.runs:
            is_yellow = run.font.highlight_color == WD_COLOR_INDEX.YELLOW
            if is_yellow:
                seq.append(run.text or "")
            elif seq:
                text = _normalize_spaces("".join(seq))
                if text:
                    samples.append(text)
                seq = []
        if seq:
            text = _normalize_spaces("".join(seq))
            if text:
                samples.append(text)
    return samples


def _replace_vlk_highlights(doc: DocxDocument, values):
    """
    Замінює жовті фрагменти у документі значеннями з `values` за порядком.
    """
    idx = 0
    for p in _iter_doc_paragraphs_for_replace(doc):
        seq_runs = []
        for run in p.runs:
            is_yellow = run.font.highlight_color == WD_COLOR_INDEX.YELLOW
            if is_yellow:
                seq_runs.append(run)
                continue
            if seq_runs:
                replacement = values[idx] if idx < len(values) else ""
                seq_runs[0].text = replacement
                for extra in seq_runs[1:]:
                    extra.text = ""
                idx += 1
                seq_runs = []
        if seq_runs:
            replacement = values[idx] if idx < len(values) else ""
            seq_runs[0].text = replacement
            for extra in seq_runs[1:]:
                extra.text = ""
            idx += 1

def validate_date_format(date_string):
    """Валідує формат дати дд.мм.рррр"""
    if not date_string:
        return True
    
    pattern = r'^\d{2}\.\d{2}\.\d{4}$'
    if not re.match(pattern, date_string):
        return False
    
    try:
        day, month, year = map(int, date_string.split('.'))
        if month < 1 or month > 12 or day < 1 or day > 31:
            return False
        if year < 1900 or year > datetime.now().year:
            return False
        # Перевіряємо чи дата існує
        datetime(year, month, day)
        return True
    except ValueError:
        return False

def format_treatment_history(person_treatments, hide_diagnosis):
    """Форматує історію лікування"""
    if person_treatments.empty:
        default_line = "За час проходження військової служби не знаходився на стаціонарному або амбулаторному лікуванні у закладах Міністерства охорони здоров'я України та медичних територіальних об'єднань Міністерства внутрішніх справ України."
        return ["\t" + default_line]

    history_lines = []
    # Хронологічне сортування з явним парсингом дат (dayfirst) і NaT в кінець
    df_sorted = person_treatments.reset_index().rename(columns={'index': 'orig_idx'}).copy()
    df_sorted['__start_dt'] = pd.to_datetime(df_sorted['Дата надходження в поточний Л/З'], errors='coerce', dayfirst=True)
    df_sorted['__end_dt'] = pd.to_datetime(df_sorted['Дата виписки'], errors='coerce', dayfirst=True)

    # Пріоритет типів лікування для впорядкування подій одного дня
    priority_map = {
        'стабілізаційний пункт': 0,
        'стаціонар': 1,
        'стаціонарне': 1,
        'денний стаціонар': 2,
        'амбулаторно': 3,
        'амбулаторне': 3,
        'реабілітація': 4,
        'влк': 5,
        'відпустка': 6,
    }
    df_sorted['__type_prio'] = df_sorted['Вид лікування'].astype(str).str.lower().map(priority_map).fillna(9).astype(int)

    df_sorted = df_sorted.sort_values(
        by=['__start_dt', '__type_prio', '__end_dt', 'orig_idx'],
        kind='mergesort',
        na_position='last'
    )
    try:
        logger.info("Порядок лікувань після сортування (перші 50):")
        for i, rec in enumerate(df_sorted.head(50).to_dict('records')):
            logger.info(
                f"{i+1}) start_raw={rec.get('Дата надходження в поточний Л/З')} | "
                f"end_raw={rec.get('Дата виписки')} | "
                f"start_dt={rec.get('__start_dt')} | end_dt={rec.get('__end_dt')} | "
                f"type={rec.get('Вид лікування')} | prio={rec.get('__type_prio')} | place={rec.get('Місце госпіталізації')}"
            )
    except Exception as _e:
        logger.warning(f"Не вдалося вивести діагностичний порядок сортування: {_e}")

    for index, row in df_sorted.iterrows():
        start_date_obj = row['Дата надходження в поточний Л/З']
        end_date_obj = row['Дата виписки']
        
        # Безпечно обробляємо дати
        try:
            start_date = start_date_obj.strftime('%d.%m.%Y') if pd.notna(start_date_obj) else "[дата не вказана]"
        except:
            start_date = "[дата не вказана]"
            
        try:
            end_date = end_date_obj.strftime('%d.%m.%Y') if pd.notna(end_date_obj) else "по теперішній час"
        except:
            end_date = "по теперішній час"

        treatment_type = str(row.get('Вид лікування', '')).lower()
        place = row['Місце госпіталізації']
        diagnosis = row['Попередній діагноз']
        vlk_conclusion = row['Заключення ВЛК']
        line = ""

        if treatment_type == 'стабілізаційний пункт':
            circumstances_text = row.get('Обставини отримання поранення/ травмування', '') or ''
            info = parse_circumstances(circumstances_text)
            injury_date = info.get('injury_date') or start_date
            location = info.get('location')
            factor = info.get('factor')

            # Формуємо потрібну конструкцію
            part1 = f"{injury_date} під час виконання бойового завдання"
            if location:
                part1 += f" в районі н. п. {location}"
            if factor:
                part1 += f" отримав поранення внаслідок {factor}"
            else:
                part1 += " отримав поранення"

            part2 = f"{start_date} евакуйований для надання першої медичної допомоги в {place}"

            base_text = f"{part1}. {part2}."
            if hide_diagnosis:
                line = base_text
            else:
                # Додаємо крапку після діагнозу, якщо її немає
                diagnosis_text = diagnosis.strip() if diagnosis else ""
                if diagnosis_text and not diagnosis_text.endswith('.'):
                    diagnosis_text += "."
                line = base_text + f" Діагноз: {diagnosis_text}"
        
        elif treatment_type in ['стаціонар', 'стаціонарне']:
            base_text = f"З {start_date} по {end_date} перебував на стаціонарному лікуванні в {place}"
            if hide_diagnosis:
                line = base_text + "."
            else:
                # Додаємо крапку після діагнозу, якщо її немає
                diagnosis_text = diagnosis.strip() if diagnosis else ""
                if diagnosis_text and not diagnosis_text.endswith('.'):
                    diagnosis_text += "."
                line = base_text + f". Діагноз: {diagnosis_text}"
        
        elif treatment_type in ['амбулаторно', 'амбулаторне']:
            base_text = f"З {start_date} по {end_date} перебував на амбулаторному лікуванні в {place}"
            if hide_diagnosis:
                line = base_text + "."
            else:
                # Додаємо крапку після діагнозу, якщо її немає
                diagnosis_text = diagnosis.strip() if diagnosis else ""
                if diagnosis_text and not diagnosis_text.endswith('.'):
                    diagnosis_text += "."
                line = base_text + f". Діагноз: {diagnosis_text}"

        elif treatment_type == 'реабілітація':
            base_text = f"З {start_date} по {end_date} проходив реабілітаційне лікування в {place}"
            if hide_diagnosis:
                line = base_text + "."
            else:
                # Додаємо крапку після діагнозу, якщо її немає
                diagnosis_text = diagnosis.strip() if diagnosis else ""
                if diagnosis_text and not diagnosis_text.endswith('.'):
                    diagnosis_text += "."
                line = base_text + f". Діагноз: {diagnosis_text}"

        elif treatment_type == 'денний стаціонар':
            base_text = f"З {start_date} по {end_date} перебував на денному стаціонарі в {place}"
            if hide_diagnosis:
                line = base_text + "."
            else:
                # Додаємо крапку після діагнозу, якщо її немає
                diagnosis_text = diagnosis.strip() if diagnosis else ""
                if diagnosis_text and not diagnosis_text.endswith('.'):
                    diagnosis_text += "."
                line = base_text + f". Діагноз: {diagnosis_text}"

        elif treatment_type == 'лазарет':
            base_text = f"З {start_date} по {end_date} перебував на лікуванні в лазареті {place}"
            if hide_diagnosis:
                line = base_text + "."
            else:
                # Додаємо крапку після діагнозу, якщо її немає
                diagnosis_text = diagnosis.strip() if diagnosis else ""
                if diagnosis_text and not diagnosis_text.endswith('.'):
                    diagnosis_text += "."
                line = base_text + f". Діагноз: {diagnosis_text}"

        elif treatment_type == 'лікування за кордоном':
            base_text = f"З {start_date} по {end_date} проходив лікування за кордоном в {place}"
            if hide_diagnosis:
                line = base_text + "."
            else:
                # Додаємо крапку після діагнозу, якщо її немає
                diagnosis_text = diagnosis.strip() if diagnosis else ""
                if diagnosis_text and not diagnosis_text.endswith('.'):
                    diagnosis_text += "."
                line = base_text + f". Діагноз: {diagnosis_text}"
            
        elif treatment_type == 'влк':
            line = f"З {start_date} по {end_date} проходив військово-лікарську комісію (ВЛК) в {place}."
            
        elif treatment_type == 'відпустка':
            conclusion = vlk_conclusion if pd.notna(vlk_conclusion) and vlk_conclusion else "перебував у відпустці за станом здоров'я за рішенням ВЛК"
            line = f"З {start_date} по {end_date} {conclusion}."
        
        if line: 
            history_lines.append(line)

    if not history_lines:
        default_line = "За час проходження військової служби не знаходився на стаціонарному або амбулаторному лікуванні у закладах Міністерства охорони здоров'я України та медичних територіальних об'єднань Міністерства внутрішніх справ України."
        return ["\t" + default_line]

    # Повертаємо список рядків; таб додаємо на початку для візуального відступу
    return ["\t" + line for line in history_lines]


@app.route('/api/treatments_ready', methods=['GET'])
def treatments_ready():
    """Статус прогріву кешу Excel (для прелоадера на сторінці форми)."""
    return jsonify({
        "ready": treatments_cache is not None,
        "loading": treatments_load_in_progress and treatments_cache is None,
        "error": treatments_last_load_error,
    })


@app.route('/api/treatments_sources', methods=['GET'])
def treatments_sources():
    """Список підключених файлів treatments_YYYY.xlsx (без повних шляхів на диску)."""
    items = []
    latest_mtime = None
    latest_path = None
    for year, path in list_treatments_year_files_sorted():
        try:
            mtime = os.path.getmtime(path)
        except OSError:
            mtime = None
        items.append({
            'year': year,
            'filename': os.path.basename(path),
            'mtime': mtime,
            'updated_at': _mtime_iso(path),
            'updated_display': _format_mtime_display(path),
        })
        if mtime is not None and (latest_mtime is None or mtime > latest_mtime):
            latest_mtime = mtime
            latest_path = path
    if os.path.isfile(TREATMENTS_FINAL_FILE):
        try:
            final_mtime = os.path.getmtime(TREATMENTS_FINAL_FILE)
        except OSError:
            final_mtime = None
        if final_mtime is not None and (latest_mtime is None or final_mtime > latest_mtime):
            latest_mtime = final_mtime
            latest_path = TREATMENTS_FINAL_FILE
    cy = datetime.now().year
    dest = treatments_path_for_year(cy)
    return jsonify({
        'year_files': items,
        'has_final': os.path.isfile(TREATMENTS_FINAL_FILE),
        'calendar_year': cy,
        'current_year_target': cy,
        'current_year_filename': os.path.basename(dest),
        'current_year_exists': os.path.isfile(dest),
        'last_updated_at': _mtime_iso(latest_path) if latest_path else None,
        'last_updated_display': _format_mtime_display(latest_path) if latest_path else None,
        'last_updated_filename': os.path.basename(latest_path) if latest_path else None,
    })


@app.route('/api/treatments_upload', methods=['POST'])
def treatments_upload():
    """
    Заміна / створення data/treatments_YYYY.xlsx: перевірка, атомарний запис, скидання кешу.
    Старий файл того ж року перезаписується.
    """
    if 'file' not in request.files:
        return jsonify({'ok': False, 'error': 'Файл не передано (поле file)'}), 400
    up = request.files['file']
    if not up or not up.filename:
        return jsonify({'ok': False, 'error': 'Оберіть файл .xlsx'}), 400
    if not up.filename.lower().endswith('.xlsx'):
        return jsonify({'ok': False, 'error': 'Допускається лише розширення .xlsx'}), 400

    year = request.form.get('year', type=int)
    if year is None:
        year = datetime.now().year
    if year < 1990 or year > 2100:
        return jsonify({'ok': False, 'error': 'Некоректний рік (1990–2100)'}), 400

    os.makedirs(DATA_DIR, exist_ok=True)
    os.makedirs(TEMP_DIR, exist_ok=True)
    tmp_path = None
    try:
        fd, tmp_path = tempfile.mkstemp(suffix='.xlsx', dir=TEMP_DIR)
        os.close(fd)
        up.save(tmp_path)
        size = os.path.getsize(tmp_path)
        if size > TREATMENTS_UPLOAD_MAX_BYTES:
            return (
                jsonify({
                    'ok': False,
                    'error': f'Файл завеликий (макс. {TREATMENTS_UPLOAD_MAX_BYTES // (1024 * 1024)} МБ)',
                }),
                413,
            )
        try:
            df = pd.read_excel(tmp_path)
        except Exception as e:
            return jsonify({'ok': False, 'error': f'Не вдалося прочитати Excel: {e}'}), 400
        ok, err_msg = _validate_treatments_upload_dataframe(df)
        if not ok:
            return jsonify({'ok': False, 'error': err_msg}), 400

        dest = treatments_path_for_year(year)
        staging = os.path.join(DATA_DIR, f'.treatments_{year}.upload.tmp.xlsx')
        with _treatments_load_lock:
            shutil.copyfile(tmp_path, staging)
            os.replace(staging, dest)
            _invalidate_treatments_cache_unlocked()
        logger.info(
            "Оновлено %s з файлу %s (%s рядків)",
            os.path.basename(dest),
            secure_filename(up.filename),
            len(df),
        )
        return jsonify({
            'ok': True,
            'year': year,
            'filename': os.path.basename(dest),
            'rows': int(len(df)),
            'updated_at': _mtime_iso(dest),
            'updated_display': _format_mtime_display(dest),
        })
    finally:
        if tmp_path and os.path.isfile(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


@app.route('/api/base_info', methods=['GET'])
def base_info():
    """Актуальність кадрової бази data/base.xlsx."""
    exists = os.path.isfile(BASE_PERSONNEL_FILE)
    rows = 0
    indexed = 0
    if exists:
        try:
            df = load_base_personnel_data()
            if df is not None:
                rows = int(len(df))
            indexed = len(load_base_personnel_index())
        except Exception as e:
            logger.warning("base_info: не вдалося прочитати base.xlsx: %s", e)
    return jsonify({
        'ok': True,
        'exists': exists,
        'filename': 'base.xlsx',
        'rows': rows,
        'indexed': indexed,
        'updated_at': _mtime_iso(BASE_PERSONNEL_FILE) if exists else None,
        'updated_display': _format_mtime_display(BASE_PERSONNEL_FILE) if exists else None,
    })


@app.route('/api/base_upload', methods=['POST'])
def base_upload():
    """
    Заміна data/base.xlsx: валідація структури відносно поточного файлу,
    атомарний запис, скидання кешу. Без створення base_1/base_2 тощо.
    """
    if 'file' not in request.files:
        return jsonify({'ok': False, 'error': 'Файл не передано (поле file)'}), 400
    up = request.files['file']
    if not up or not up.filename:
        return jsonify({'ok': False, 'error': 'Оберіть файл .xlsx'}), 400
    if not up.filename.lower().endswith('.xlsx'):
        return jsonify({'ok': False, 'error': 'Допускається лише розширення .xlsx'}), 400

    os.makedirs(DATA_DIR, exist_ok=True)
    os.makedirs(TEMP_DIR, exist_ok=True)
    tmp_path = None
    try:
        fd, tmp_path = tempfile.mkstemp(suffix='.xlsx', dir=TEMP_DIR)
        os.close(fd)
        up.save(tmp_path)
        size = os.path.getsize(tmp_path)
        if size > BASE_UPLOAD_MAX_BYTES:
            return (
                jsonify({
                    'ok': False,
                    'error': f'Файл завеликий (макс. {BASE_UPLOAD_MAX_BYTES // (1024 * 1024)} МБ)',
                }),
                413,
            )
        try:
            df = pd.read_excel(tmp_path)
        except Exception as e:
            return jsonify({'ok': False, 'error': f'Не вдалося прочитати Excel: {e}'}), 400

        reference_df = None
        if os.path.isfile(BASE_PERSONNEL_FILE):
            try:
                reference_df = pd.read_excel(BASE_PERSONNEL_FILE)
            except Exception as e:
                logger.warning("Не вдалося прочитати поточний base.xlsx для порівняння: %s", e)
                reference_df = None

        ok, err_msg = _validate_base_upload_dataframe(df, reference_df)
        if not ok:
            return jsonify({'ok': False, 'error': err_msg}), 400

        staging = os.path.join(DATA_DIR, '.base.upload.tmp.xlsx')
        with _base_load_lock:
            shutil.copyfile(tmp_path, staging)
            os.replace(staging, BASE_PERSONNEL_FILE)
            _invalidate_base_cache_unlocked()
        # Прогріти кеш одразу після заміни
        load_base_personnel_data()
        indexed = len(load_base_personnel_index())
        logger.info(
            "Оновлено base.xlsx з файлу %s (%s рядків, %s у індексі)",
            secure_filename(up.filename),
            len(df),
            indexed,
        )
        return jsonify({
            'ok': True,
            'filename': 'base.xlsx',
            'rows': int(len(df)),
            'indexed': indexed,
            'updated_at': _mtime_iso(BASE_PERSONNEL_FILE),
            'updated_display': _format_mtime_display(BASE_PERSONNEL_FILE),
        })
    finally:
        if tmp_path and os.path.isfile(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


@app.route('/api/stats', methods=['GET'])
def get_stats():
    """API endpoint для отримання статистики бази даних"""
    try:
        treatments_df = load_treatments_data()
        latest_path = None
        latest_mtime = None
        for _, path in list_treatments_year_files_sorted():
            try:
                mtime = os.path.getmtime(path)
            except OSError:
                continue
            if latest_mtime is None or mtime > latest_mtime:
                latest_mtime = mtime
                latest_path = path
        if os.path.isfile(TREATMENTS_FINAL_FILE):
            try:
                final_mtime = os.path.getmtime(TREATMENTS_FINAL_FILE)
            except OSError:
                final_mtime = None
            if final_mtime is not None and (latest_mtime is None or final_mtime > latest_mtime):
                latest_path = TREATMENTS_FINAL_FILE
        stats = {
            "total_records": int(len(treatments_df)),
            "unique_patients": int(treatments_df["ПІБ_чисте"].nunique()) if "ПІБ_чисте" in treatments_df.columns else 0,
            "last_updated_at": _mtime_iso(latest_path) if latest_path else None,
            "last_updated_display": _format_mtime_display(latest_path) if latest_path else None,
            "last_updated_filename": os.path.basename(latest_path) if latest_path else None,
        }
        return jsonify(stats)
        
    except Exception as e:
        logger.error(f"Помилка при отриманні статистики: {e}")
        return jsonify({'error': str(e)})

@app.route('/api/pib_genitive', methods=['GET'])
def api_pib_genitive():
    """Автоматичний родовий відмінок з називного (для підказки у формі)."""
    q = request.args.get('q', '').strip()
    if len(q) < 2:
        return jsonify({'genitive': ''})
    try:
        gen = nominative_pib_to_genitive_line(q)
        return jsonify({'genitive': gen or ''})
    except Exception as e:
        logger.warning("api_pib_genitive: %s", e)
        return jsonify({'genitive': '', 'error': str(e)})


@app.route('/api/service_signatories', methods=['GET', 'POST'])
def api_service_signatories():
    """Зчитування/збереження підписантів службової характеристики."""
    if request.method == 'GET':
        return jsonify(load_service_signatories())

    data = request.get_json(silent=True) or {}
    required = (
        "pidpysant_1_zvannya",
        "pidpysant_1_pib",
        "pidpysant_2_zvannya",
        "pidpysant_2_pib",
    )
    payload = {k: _normalize_spaces(data.get(k, "")) for k in required}
    if not all(payload.values()):
        return jsonify({
            'ok': False,
            'error': "Заповніть усі поля підписантів перед збереженням",
        }), 400
    try:
        save_service_signatories(payload)
        return jsonify({'ok': True, **payload})
    except Exception as e:
        logger.error("Помилка збереження підписантів: %s", e)
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/vlk_signatories', methods=['GET', 'POST'])
def api_vlk_signatories():
    """Зчитування/збереження підписантів для рапорту ВЛК."""
    if request.method == 'GET':
        return jsonify(load_vlk_signatories())

    data = request.get_json(silent=True) or {}
    action = _normalize_spaces(data.get("action", "save_selection")).lower() or "save_selection"
    current = load_vlk_signatories()

    if action == "add_option":
        role = _normalize_spaces(data.get("role", "")).lower()
        zvannya = _normalize_spaces(data.get("zvannya", ""))
        pib = _normalize_spaces(data.get("pib", ""))
        if role not in ("kombrig", "kombat"):
            return jsonify({"ok": False, "error": "Невідома роль підписанта"}), 400
        if not zvannya or not pib:
            return jsonify({"ok": False, "error": "Вкажіть звання та ПІБ"}), 400
        # Прізвище великими для узгодженості з документами
        parts = pib.split(" ")
        if len(parts) >= 2:
            parts[-1] = parts[-1].upper()
            pib = _normalize_spaces(" ".join(parts))
        key = "kombrig_options" if role == "kombrig" else "kombat_options"
        options = list(current.get(key) or [])
        options.append({"zvannya": zvannya, "pib": pib})
        payload = {
            "tvo_kombrig_zvannya": current.get("tvo_kombrig_zvannya", ""),
            "tvo_kombrig_pib": current.get("tvo_kombrig_pib", ""),
            "tvo_kombat_zvannya": current.get("tvo_kombat_zvannya", ""),
            "tvo_kombat_pib": current.get("tvo_kombat_pib", ""),
            "kombrig_options": current.get("kombrig_options", []),
            "kombat_options": current.get("kombat_options", []),
        }
        payload[key] = options
        if role == "kombrig":
            payload["tvo_kombrig_zvannya"] = zvannya
            payload["tvo_kombrig_pib"] = pib
        else:
            payload["tvo_kombat_zvannya"] = zvannya
            payload["tvo_kombat_pib"] = pib
        try:
            save_vlk_signatories(payload)
            return jsonify({"ok": True, **load_vlk_signatories()})
        except Exception as e:
            logger.error("Помилка додавання підписанта ВЛК: %s", e)
            return jsonify({"ok": False, "error": str(e)}), 500

    required = (
        "tvo_kombrig_zvannya",
        "tvo_kombrig_pib",
        "tvo_kombat_zvannya",
        "tvo_kombat_pib",
    )
    payload = {k: _normalize_spaces(data.get(k, "")) for k in required}
    if not all(payload.values()):
        return jsonify({
            'ok': False,
            'error': "Заповніть усі поля підписантів перед збереженням",
        }), 400
    payload["kombrig_options"] = current.get("kombrig_options", [])
    payload["kombat_options"] = current.get("kombat_options", [])
    try:
        save_vlk_signatories(payload)
        return jsonify({'ok': True, **load_vlk_signatories()})
    except Exception as e:
        logger.error("Помилка збереження VLK-підписантів: %s", e)
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/lpz_list', methods=['GET'])
def api_lpz_list():
    """Список ЛПЗ для автокомпліту рапортів."""
    path = os.path.join(DATA_DIR, "lpz_list.json")
    try:
        if not os.path.isfile(path):
            return jsonify([])
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if not isinstance(data, list):
            return jsonify([])
        items = []
        for x in data:
            s = _normalize_spaces(str(x)) if x is not None else ""
            if s:
                items.append(s)
        return jsonify(items)
    except Exception as e:
        logger.warning("Не вдалося зчитати lpz_list.json: %s", e)
        return jsonify([])


@app.route('/api/likar_specializations', methods=['GET'])
def api_likar_specializations():
    """Список спеціалізацій лікаря (родовий відмінок) для рапортів."""
    path = os.path.join(DATA_DIR, "likar_specializations.json")
    try:
        if not os.path.isfile(path):
            return jsonify([])
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if not isinstance(data, list):
            return jsonify([])
        items = []
        for x in data:
            s = _normalize_spaces(str(x)) if x is not None else ""
            if s:
                items.append(s)
        return jsonify(items)
    except Exception as e:
        logger.warning("Не вдалося зчитати likar_specializations.json: %s", e)
        return jsonify([])


@app.route('/api/search_pib', methods=['GET'])
def search_pib():
    """API endpoint для пошуку ПІБ"""
    try:
        query = request.args.get('q', '').strip()
        context_mode = request.args.get('context', '').strip().lower()
        if len(query) < 2:
            return jsonify({'results': []})

        q_clean = re.sub(r"\s+", " ", query).strip().lower()
        if not q_clean:
            return jsonify({'results': []})

        # Рапорти / службова характеристика — лише base.xlsx (кешований індекс).
        if context_mode == "service":
            index = load_base_personnel_index()
            prefix_hits = []
            contains_hits = []
            for pib_clean, person in index.items():
                if not pib_clean:
                    continue
                if pib_clean.startswith(q_clean):
                    prefix_hits.append(person)
                elif q_clean in pib_clean:
                    contains_hits.append(person)
            matched_people = (prefix_hits + contains_hits)[:10]
            results = []
            for person in matched_people:
                sur = person.get("prizvyshche", "")
                first = person.get("imya", "")
                pat = person.get("po_batkovi", "")
                full_name = _normalize_spaces(f"{sur} {first} {pat}")
                if not full_name:
                    continue
                rank = person.get("zvanie", "")
                position = person.get("zaimana_posada", "")
                label = f"{full_name} ({rank})" if rank else full_name
                results.append({
                    "label": label,
                    "value": full_name,
                    "rank": rank,
                    "position": position,
                    "prizvyshche": sur,
                    "imya": first,
                    "po_batkovi": pat,
                    "birth_date": "",
                })
            return jsonify({'results': results})

        # Амбулаторний журнал / нове звернення: спочатку rkch.xlsx, інакше treatments.
        if context_mode in ("outpatient", "rkch"):
            results = _search_rkch_results(q_clean, limit=10)
            if not results:
                results = _search_treatments_results(q_clean, limit=10)
            return jsonify({'results': results})

        # Медична характеристика — пошук у treatments.
        return jsonify({'results': _search_treatments_results(q_clean, limit=10)})

    except Exception as e:
        logger.error(f"Помилка при пошуку ПІБ: {e}")
        return jsonify({'results': [], 'error': str(e)})

def _welcome_template_context():
    """Контекст вітальної сторінки: чи є Excel у data/."""
    has_files = bool(list_treatments_year_files_sorted()) or os.path.isfile(
        TREATMENTS_FINAL_FILE
    )
    return {
        "data_ready": has_files,
        "data_dir": os.path.abspath(DATA_DIR),
    }


@app.route('/', methods=['GET'])
def index():
    """Вітальна сторінка з інструкцією щодо data/ та встановлення."""
    return render_template('welcome.html', **_welcome_template_context())


@app.route('/databases', methods=['GET'])
def databases():
    """Статистика та оновлення Excel-баз (treatments / base)."""
    return render_template('databases.html')


@app.route('/medical-characteristic', methods=['GET', 'POST'])
def medical_characteristic():
    """Генерація медичної характеристики"""
    def _wants_ajax():
        accept = (request.headers.get("Accept") or "").lower()
        return (
            request.headers.get("X-Requested-With") == "XMLHttpRequest"
            or "application/json" in accept
        )

    def _ajax_error(message: str, status: int = 400):
        if _wants_ajax():
            return jsonify({"ok": False, "error": message}), status
        flash(message, "error")
        return render_template('medical_characteristic.html')

    if request.method == 'POST':
        pib_nazivnyi = request.form.get('pib_nazivnyi', '').strip()
        pib_rodovyi_input = request.form.get('pib_rodovyi', '').strip()
        hide_diagnosis_flag = request.form.get('no_diagnosis')
        enlistment_date = request.form.get('enlistment_date', '').strip()
        enlistment_date_custom = request.form.get('enlistment_date_custom', '').strip()
        observation_end = request.form.get('observation_end', '').strip()
        observation_end_custom = request.form.get('observation_end_custom', '').strip()
        signatory = request.form.get('signatory', '').strip()
        birth_date = request.form.get('birth_date', '').strip()

        # Обробка дати призову
        if enlistment_date == "custom":
            if not enlistment_date_custom:
                return _ajax_error("Вкажіть конкретну дату зарахування")
            if not validate_date_format(enlistment_date_custom):
                return _ajax_error("Невірний формат дати зарахування. Використовуйте формат дд.мм.рррр")
            final_enlistment_date = f"з {enlistment_date_custom} року"
        elif enlistment_date == "з моменту призову":
            final_enlistment_date = "з моменту призову"
        else:
            # Для конкретних дат додаємо префікс "з" і слово "року"
            final_enlistment_date = f"з {enlistment_date} року"

        # Обробка дати завершення нагляду
        if observation_end == "custom":
            if not observation_end_custom:
                return _ajax_error("Вкажіть конкретну дату завершення нагляду")
            if not validate_date_format(observation_end_custom):
                return _ajax_error("Невірний формат дати завершення нагляду. Використовуйте формат дд.мм.рррр")
            final_observation_end = f"по {observation_end_custom} року"
        elif observation_end == "по теперішній час":
            final_observation_end = "по теперішній час"
        else:
            final_observation_end = ""

        # Обробка підписанта (розділяємо на окремі поля для правильного вирівнювання в Word)
        if signatory == "acting_chief":
            signatory_position = "Т. в. о. начальника медичної служби"
            signatory_department = "секції тилу"
            signatory_rank = "капітан м/с"
            signatory_name = "Ірина ГОНЧАРОВА"
        elif signatory == "chief":
            signatory_position = "Начальник медичної служби"
            signatory_department = "секції тилу"
            signatory_rank = "майор м/с"
            signatory_name = "Юлія ЮРЧАК"
        elif signatory == "company_commander":
            signatory_position = "Командир медичної роти"
            signatory_department = None
            signatory_rank = "капітан м/с"
            signatory_name = "Євгеній МАЖАЄВ"
        else:
            signatory_position = None
            signatory_department = None
            signatory_rank = None
            signatory_name = None

        # Валідація обов'язкових полів
        if not pib_nazivnyi:
            return _ajax_error("ПІБ (в називному відмінку) є обов'язковим полем")
        if not final_enlistment_date:
            return _ajax_error("Дата зарахування є обов'язковим полем")
        if not signatory:
            return _ajax_error("Оберіть підписанта")
        if birth_date and not validate_date_format(birth_date):
            return _ajax_error("Невірний формат дати народження. Використовуйте формат дд.мм.рррр")

        try:
            treatments_df = load_treatments_data()
        except Exception as e:
            logger.error(f"Помилка при завантаженні даних: {e}")
            return _ajax_error(f"Помилка при завантаженні даних: {e}", 500)

        pib_nazivnyi_clean = re.sub(r'\s+', ' ', pib_nazivnyi).strip().lower()
        soldier_records = treatments_df[treatments_df['ПІБ_чисте'] == pib_nazivnyi_clean]

        context = {}
        data_source = "manual"

        if not soldier_records.empty:
            data_source = "treatments"
            first_record = soldier_records.iloc[0]
            kategoriia = first_record['Категорія']
            birth_date_obj = first_record['Дата народження']
            zvanie, _ = split_medical_service_rank(first_record['Військове звання'])
            try:
                birth_date_str = birth_date_obj.strftime('%d.%m.%Y') if pd.notna(birth_date_obj) else "[дата не вказана]"
            except Exception:
                birth_date_str = "[дата не вказана]"
            context = {
                'zvanie': zvanie,
                'sluzhba_type': 'за контрактом' if 'контр' in str(kategoriia).lower() else 'під час мобілізації',
                'birth_date': birth_date_str,
                'treatment_history': format_treatment_history(soldier_records, hide_diagnosis_flag),
            }
        else:
            context = {
                'zvanie': split_medical_service_rank(request.form.get('zvanie'))[0],
                'sluzhba_type': request.form.get('sluzhba_type'),
                'birth_date': request.form.get('birth_date'),
                'treatment_history': ["\t" + "За час проходження військової служби не знаходився на стаціонарному або амбулаторному лікуванні у закладах Міністерства охорони здоров'я України та медичних територіальних об'єднань Міністерства внутрішніх справ України."],
            }

        pib_nazivnyi_display = format_nominative_pib_display(pib_nazivnyi)
        context['pib_nazivnyi'] = pib_nazivnyi_display
        context['pib_rodovyi'] = build_pib_rodovyi_for_document(
            pib_nazivnyi_display, pib_rodovyi_input
        )
        context['enlistment_date'] = final_enlistment_date
        context['observation_end'] = final_observation_end
        context['signatory_position'] = signatory_position
        context['signatory_rank'] = signatory_rank
        context['signatory_name'] = signatory_name

        # Додаємо signatory_department тільки якщо воно не None (тобто не для Мажаєва)
        if signatory_department is not None:
            context['signatory_department'] = signatory_department
            context['signatory_department_with_break'] = signatory_department + "\n"
        else:
            context['signatory_department_with_break'] = ""

        # Створюємо поле для звання та імені з табуляцією для вирівнювання по краях
        context['signatory_rank_and_name'] = f"{signatory_rank}\t\t\t\t\t{signatory_name}"

        # Додаємо звання в родовому відмінку для шапки
        context['zvanie_genitive'] = format_rank_genitive(context.get('zvanie', ''))
        context['data_stvorennya'] = datetime.now().strftime("%d.%m.%Y")

        # Вибір шаблону залежно від підписанта
        if signatory == "company_commander":
            # Використовуємо окремий шаблон для Мажаєва
            template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates', 'medical_characteristic_mazhaev_template.docx')
        else:
            # Стандартний шаблон для Гончарової та Юрчак
            template_path = MEDICAL_CHARACTERISTIC_TEMPLATE

        # Рендеримо DOCX із маркерами розділення абзаців для історії лікування
        doc = DocxTemplate(template_path)
        # Готуємо плейсхолдер з маркером для безпечної пост-обробки
        context_for_tpl = dict(context)
        history_parts = [line.lstrip('\t') for line in context.get('treatment_history', [])]
        joined_with_marker = "[[PARA_SPLIT]]".join(history_parts)
        context_for_tpl['treatment_history'] = joined_with_marker
        doc.render(context_for_tpl)

        # Відкриваємо результат та розгортаємо маркер у окремі абзаци (безпечний спосіб)
        temp_stream = io.BytesIO()
        doc.save(temp_stream)
        temp_stream.seek(0)
        rendered = DocxDocument(temp_stream)

        # Обробка маркерів розбиття на абзаци (тільки для історії лікування)
        split_marker = "[[PARA_SPLIT]]"
        for p in list(rendered.paragraphs):
            if split_marker in p.text:
                parts = p.text.split(split_marker)

                # Запам'ятовуємо стиль і зразок шрифту з плейсхолдера
                placeholder_style = p.style
                sample_font = p.runs[0].font if p.runs else None

                def apply_formatting(paragraph: DocxParagraph):
                    paragraph.style = placeholder_style
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    fmt = paragraph.paragraph_format
                    fmt.space_before = Pt(0)
                    fmt.space_after = Pt(0)

                # Вставляємо кожен пункт перед плейсхолдером у прямому порядку
                for part in parts:
                    new_p = p.insert_paragraph_before()
                    apply_formatting(new_p)
                    run = new_p.add_run()
                    if sample_font is not None:
                        if sample_font.name is not None:
                            run.font.name = sample_font.name
                        if sample_font.size is not None:
                            run.font.size = sample_font.size
                        run.font.bold = sample_font.bold
                        run.font.italic = sample_font.italic
                        run.font.underline = sample_font.underline
                    run.add_tab()
                    run.add_text(part)

                # Видаляємо оригінальний абзац з маркером
                p._element.getparent().remove(p._element)

        file_stream = io.BytesIO()
        rendered.save(file_stream)
        file_stream.seek(0)
        file_bytes = file_stream.getvalue()
        download_name = f'Медична_характеристика_{pib_nazivnyi_display.replace(" ", "_")}.docx'

        _loader_preview_delay()

        if _wants_ajax():
            form_birth = context.get("birth_date") or ""
            if form_birth == "[дата не вказана]":
                form_birth = ""
            return jsonify({
                "ok": True,
                "source": data_source,
                "resolved": {
                    "pib_nazivnyi": context.get("pib_nazivnyi") or "",
                    "pib_rodovyi": context.get("pib_rodovyi") or "",
                    "zvanie": context.get("zvanie") or "",
                    "sluzhba_type": context.get("sluzhba_type") or "",
                    "birth_date": form_birth,
                },
                "filename": download_name,
                "file_base64": base64.b64encode(file_bytes).decode("ascii"),
            })

        response = make_response(send_file(
            io.BytesIO(file_bytes), as_attachment=True,
            download_name=download_name,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ))
        response.set_cookie('fileDownload', 'true', max_age=20)

        return response

    return render_template('medical_characteristic.html')


def _service_signatory_template_defaults():
    """Підписанти службової: вибір з тих самих списків, що й для рапортів."""
    service = load_service_signatories()
    vlk = load_vlk_signatories()
    return {
        "pidpysant_1_zvannya": service.get("pidpysant_1_zvannya") or vlk.get("tvo_kombrig_zvannya", ""),
        "pidpysant_1_pib": service.get("pidpysant_1_pib") or vlk.get("tvo_kombrig_pib", ""),
        "pidpysant_2_zvannya": service.get("pidpysant_2_zvannya") or vlk.get("tvo_kombat_zvannya", ""),
        "pidpysant_2_pib": service.get("pidpysant_2_pib") or vlk.get("tvo_kombat_pib", ""),
        "kombrig_options": vlk.get("kombrig_options", []),
        "kombat_options": vlk.get("kombat_options", []),
    }


@app.route('/service-characteristic', methods=['GET', 'POST'])
def service_characteristic():
    """Генерація службової характеристики."""
    signatory_defaults = _service_signatory_template_defaults()
    if request.method == 'POST':
        pib_nazivnyi = _normalize_spaces(request.form.get('pib_nazivnyi', ''))
        zvanie_input = _normalize_spaces(request.form.get('zvanie', ''))
        prizvyshche_input = _normalize_spaces(request.form.get('prizvyshche', ''))
        imya_input = _normalize_spaces(request.form.get('imya', ''))
        po_batkovi_input = _normalize_spaces(request.form.get('po_batkovi', ''))
        posada_input = _normalize_spaces(request.form.get('zaimana_posada', ''))
        komisariat_input = request.form.get('komisariat', '')
        data_pryzovu_input = request.form.get('data_pryzovu', '')
        osvita_type_input = request.form.get('osvita_type', '')
        navchalnyy_zaklad_input = request.form.get('navchalnyy_zaklad', '')
        misto_zakladu_input = request.form.get('misto_navchalnogo_zakladu', '')
        rik_zakinchennya_input = request.form.get('rik_zakinchennya', '')

        pidpysant_1_zvannya = _normalize_spaces(request.form.get('pidpysant_1_zvannya', ''))
        pidpysant_1_pib = _normalize_spaces(request.form.get('pidpysant_1_pib', ''))
        pidpysant_2_zvannya = _normalize_spaces(request.form.get('pidpysant_2_zvannya', ''))
        pidpysant_2_pib = _normalize_spaces(request.form.get('pidpysant_2_pib', ''))

        if not pib_nazivnyi:
            flash("Прізвище Ім'я по-батькові є обов'язковим полем", "error")
            return render_template('service_characteristic.html', signatory_defaults=signatory_defaults)
        if not pidpysant_1_zvannya or not pidpysant_1_pib:
            flash("Заповніть звання і ПІБ для 1-го підписанта", "error")
            return render_template('service_characteristic.html', signatory_defaults=signatory_defaults)
        if not pidpysant_2_zvannya or not pidpysant_2_pib:
            flash("Заповніть звання і ПІБ для 2-го підписанта", "error")
            return render_template('service_characteristic.html', signatory_defaults=signatory_defaults)

        komisariat_ta_data_prizovu = _format_komisariat_parts(komisariat_input, data_pryzovu_input)
        if not komisariat_ta_data_prizovu:
            flash(
                "Заповніть військовий комісаріат і коректну дату призову у форматі дд.мм.рррр",
                "error",
            )
            return render_template('service_characteristic.html', signatory_defaults=signatory_defaults)

        osvita = _format_osvita_parts(
            osvita_type_input,
            navchalnyy_zaklad_input,
            misto_zakladu_input,
            rik_zakinchennya_input,
        )
        if not osvita:
            flash(
                "Заповніть освіту: вид, навчальний заклад, місто та рік закінчення (4 цифри)",
                "error",
            )
            return render_template('service_characteristic.html', signatory_defaults=signatory_defaults)

        zvanie = zvanie_input
        prizvyshche = prizvyshche_input
        imya = imya_input
        po_batkovi = po_batkovi_input
        zaimana_posada = posada_input

        # Для службової характеристики пріоритет даних за штатом з base.xlsx.
        base_person = _lookup_person_in_base_excel(pib_nazivnyi)
        if base_person.get("zvanie"):
            zvanie = base_person["zvanie"]
        if base_person.get("zaimana_posada"):
            zaimana_posada = base_person["zaimana_posada"]
        prizvyshche = prizvyshche or base_person.get("prizvyshche", "")
        imya = imya or base_person.get("imya", "")
        po_batkovi = po_batkovi or base_person.get("po_batkovi", "")

        try:
            treatments_df = load_treatments_data()
        except Exception as e:
            logger.error(f"Помилка при завантаженні даних: {e}")
            flash(f"Помилка при завантаженні даних: {e}", "error")
            return render_template('service_characteristic.html', signatory_defaults=signatory_defaults)

        pib_clean = re.sub(r'\s+', ' ', pib_nazivnyi).strip().lower()
        soldier_records = treatments_df[treatments_df['ПІБ_чисте'] == pib_clean]
        if not soldier_records.empty:
            first_record = soldier_records.iloc[0]
            zvanie = zvanie or _excel_cell_str(first_record.get('Військове звання'))
            prizvyshche = prizvyshche or _excel_cell_str(first_record.get('Прізвище'))
            imya = imya or _excel_cell_str(first_record.get("Ім'я"))
            po_batkovi = po_batkovi or _excel_cell_str(first_record.get('По батькові'))
            if not zaimana_posada:
                zaimana_posada = _extract_position_from_row(first_record)

        if not (prizvyshche and imya and po_batkovi):
            p_sur, p_first, p_pat = _split_pib(pib_nazivnyi)
            prizvyshche = prizvyshche or p_sur
            imya = imya or p_first
            po_batkovi = po_batkovi or p_pat

        if not zvanie:
            flash("Не вдалося визначити військове звання. Заповніть поле вручну.", "error")
            return render_template('service_characteristic.html', signatory_defaults=signatory_defaults)
        if not (prizvyshche and imya and po_batkovi):
            flash("Не вдалося визначити Прізвище / Ім'я / По батькові. Уточніть ПІБ або заповніть поля вручну.", "error")
            return render_template('service_characteristic.html', signatory_defaults=signatory_defaults)

        if not zaimana_posada:
            zaimana_posada = "У розпорядженні командира військової частини"
        zaimana_posada = _ensure_final_period(zaimana_posada)
        zaimana_posada_z_velykoi = _title_first_letter(zaimana_posada.rstrip('.'))

        pib_gen_line = build_pib_rodovyi_for_document(
            f"{prizvyshche} {imya} {po_batkovi}",
            "",
        )
        pib_gen_parts = pib_gen_line.split(' ')
        gen_surname = _surname_first_capital(pib_gen_parts[0] if pib_gen_parts else prizvyshche)
        dat_surname = _surname_to_dative(prizvyshche)
        fraza_zvannya_ta_pib_rodovyi = _normalize_spaces(
            f"{format_rank_genitive(zvanie)} {_format_initials_name(gen_surname, imya, po_batkovi)}"
        )
        fraza_zvannya_ta_pib_davalyi = _normalize_spaces(
            f"{format_rank_dative(zvanie)} {_format_initials_name(dat_surname, imya, po_batkovi)}"
        )

        context = {
            'zvanie': zvanie,
            'prizvyshche': prizvyshche.upper(),
            'imya': imya,
            'po_batkovi': po_batkovi,
            'zaimana_posada': zaimana_posada,
            'zaimana_posada_z_velykoi': zaimana_posada_z_velykoi,
            'komisariat_ta_data_prizovu': komisariat_ta_data_prizovu,
            'osvita': osvita,
            'fraza_zvannya_ta_pib_rodovyi': fraza_zvannya_ta_pib_rodovyi,
            'fraza_zvannya_ta_pib_davalyi': fraza_zvannya_ta_pib_davalyi,
            'pidpysant_1_zvannya': pidpysant_1_zvannya,
            'pidpysant_1_pib': pidpysant_1_pib,
            'pidpysant_2_zvannya': pidpysant_2_zvannya,
            'pidpysant_2_pib': pidpysant_2_pib,
            'imya_ta_prizvyshche_pidpys': f"{imya} {prizvyshche.upper()}",
            'data_stvorennya': datetime.now().strftime("%d.%m.%Y"),
        }

        try:
            save_service_signatories({
                'pidpysant_1_zvannya': pidpysant_1_zvannya,
                'pidpysant_1_pib': pidpysant_1_pib,
                'pidpysant_2_zvannya': pidpysant_2_zvannya,
                'pidpysant_2_pib': pidpysant_2_pib,
            })
            doc = DocxTemplate(SERVICE_CHARACTERISTIC_TEMPLATE)
            doc.render(context)
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
        except Exception as e:
            logger.error("Помилка при рендері службової характеристики: %s", e)
            flash(f"Помилка генерації DOCX: {e}", "error")
            return render_template('service_characteristic.html', signatory_defaults=signatory_defaults)

        _loader_preview_delay()
        response = make_response(send_file(
            file_stream, as_attachment=True,
            download_name=f'Службова_характеристика_{imya}_{prizvyshche.upper()}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ))
        response.set_cookie('fileDownload', 'true', max_age=20)
        return response

    return render_template('service_characteristic.html', signatory_defaults=signatory_defaults)


def _build_report_body_text(
    report_type: str,
    zvanie: str,
    sluzhba_type: str,
    pib_nazivnyi: str,
    misce_napravlennya: str = "",
    data_gospitalizatsii: str = "",
    rehab_days: str = "21",
    rehab_mrc: str = "Пуща-Водиця",
) -> str:
    """Формує текст тіла рапорту залежно від типу."""
    if report_type == "vlk":
        return (
            "Прошу Вас про надання мені законної можливості проходження військово-лікарської комісії "
            "з метою встановлення придатності/непридатності до військової служби згідно Наказу МВС України "
            "від 14.08.2008 № 402 (зі змінами та додатками), з метою визначення придатності (непридатності) "
            "до подальшого проходження військової служби."
        )
    if report_type == "material_aid":
        return (
            "Прошу Вашого клопотання перед вищим командуванням про надання мені матеріальної допомоги "
            "для вирішення соціально-побутових питань за 2024 рік відповідно до наказу МВС України "
            "від 15.03.2018 № 200 «Про затвердження Інструкції про порядок виплати грошового забезпечення "
            "та одноразової грошової допомоги при звільненні військовослужбовцям Національної гвардії України "
            "та іншим особам»."
        )
    if report_type == "ozdorovlennya":
        return (
            "Прошу Вашого клопотання перед командиром військової частини 3029 про надання мені допомоги "
            "для оздоровлення за 2026 рік, згідно ПКМУ № 704 від 30.08.2017."
        )

    pib_title = _normalize_spaces(
        " ".join(_surname_first_capital(part) for part in pib_nazivnyi.split(" ") if part)
    )
    rank = _normalize_spaces(zvanie).lower()
    service = _normalize_spaces(sluzhba_type)
    clinic = _normalize_spaces(misce_napravlennya)

    if report_type == "hospitalization":
        return (
            f"Я, {rank} {service}, {pib_title}, прошу Вашого клопотання перед командиром військової частини "
            f"про надання мені дозволу на госпіталізацію в {clinic}. "
            f"Госпіталізація запланована на {_normalize_spaces(data_gospitalizatsii)}."
        )
    if report_type == "consultation":
        return (
            f"Я, {rank} {service}, {pib_title}, прошу Вашого клопотання перед командиром військової частини "
            f"про надання мені дозволу на консультацію в {clinic}."
        )
    if report_type == "rehabilitation":
        days = _normalize_spaces(rehab_days) or "21"
        mrc = _normalize_spaces(rehab_mrc) or "Пуща-Водиця"
        return (
            f"Прошу Вашого дозволу на проходження медичної реабілітації у МРЦ "
            f"«{mrc}» терміном на {days} календарний день. "
            "Претензій до медичної служби військової частини 3029 не маю."
        )
    return ""


# Рапорти з вільним текстом додатка (без «Копію консультативного висновку»).
_REPORT_TYPES_FREE_DODATOK = frozenset({"material_aid", "ozdorovlennya"})
_REPORT_TYPES_VACATION = frozenset({"vacation"})
_REPORT_TYPES_ALL = frozenset({
    "vlk", "hospitalization", "consultation", "material_aid", "ozdorovlennya", "vacation",
    "rehabilitation",
})


def _build_report_dodatok_line(
    report_type: str,
    *,
    dodatok: str = "",
    oglyad_nuber: str = "",
    likar: str = "",
    hospital: str = "",
) -> str:
    """Рядок після «До рапорту додаю:» (з нумерацією «1.»)."""
    if report_type in _REPORT_TYPES_FREE_DODATOK:
        body = _normalize_spaces(dodatok)
    else:
        # Лише для ВЛК / госпіталізації / консультації.
        parts = ["Копію консультативного висновку"]
        if oglyad_nuber:
            parts.append(f"№{oglyad_nuber}")
        if likar:
            parts.append(likar)
        if hospital:
            parts.append(hospital)
        body = _normalize_spaces(" ".join(parts))
    if not body:
        return ""
    return f"1.\t{body}"


def _set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = text


def _load_vlk_report_template() -> DocxTemplate:
    """Завантажує шаблон рапорту; рядок додатка = лише {{ dodatok_line }}."""
    marker = "{{ dodatok_line }}"
    docx = DocxDocument(VLK_REPORT_TEMPLATE)
    changed = False
    for paragraph in docx.paragraphs:
        text = paragraph.text or ""
        needs_fix = (
            "dodatok_line" in text
            or "oglyad_nuber" in text
            or "консультативного висновку" in text
        )
        if not needs_fix:
            continue
        if text.strip() == marker:
            continue
        _set_paragraph_text(paragraph, marker)
        changed = True
    stream = io.BytesIO()
    docx.save(stream)
    stream.seek(0)
    if changed:
        try:
            with open(VLK_REPORT_TEMPLATE, "wb") as fh:
                fh.write(stream.getvalue())
            stream.seek(0)
        except OSError:
            # Файл може бути відкритий у Word — працюємо з копією в пам’яті.
            stream.seek(0)
    return DocxTemplate(stream)


@app.route('/vlk-report', methods=['GET', 'POST'])
def vlk_report():
    """Генерація рапортів (медичні / матдопомога / оздоровлення / відпустка)."""
    signatory_defaults = load_service_signatories()
    vlk_signatory_defaults = load_vlk_signatories()
    defaults = {
        "report_type": "vlk",
        "sluzhba_type": "за мобілізацією",
        "misce_napravlennya": "",
        "oglyad_nuber": "",
        "likar": "",
        "hospital": "",
        "dodatok": "",
        "data_gospitalizatsii": "",
        "vacantion_start_day": "",
        "vacantion_adress": "",
        "vacantion_phone": "",
        "rehab_days": "21",
        "rehab_mrc": "Пуща-Водиця",
        "pib_nazivnyi": "",
        "tvo_kombrig_zvannya": vlk_signatory_defaults.get("tvo_kombrig_zvannya", ""),
        "tvo_kombrig_pib": vlk_signatory_defaults.get("tvo_kombrig_pib", ""),
        "tvo_kombat_zvannya": vlk_signatory_defaults.get("tvo_kombat_zvannya", "") or signatory_defaults.get("pidpysant_2_zvannya", ""),
        "tvo_kombat_pib": vlk_signatory_defaults.get("tvo_kombat_pib", "") or signatory_defaults.get("pidpysant_2_pib", ""),
        "kombrig_options": vlk_signatory_defaults.get("kombrig_options", []),
        "kombat_options": vlk_signatory_defaults.get("kombat_options", []),
    }

    if request.method == 'POST':
        report_type = _normalize_spaces(request.form.get("report_type", "vlk"))
        if report_type not in _REPORT_TYPES_ALL:
            report_type = "vlk"

        template_path = (
            VACATION_REPORT_TEMPLATE
            if report_type in _REPORT_TYPES_VACATION
            else VLK_REPORT_TEMPLATE
        )
        template_label = (
            "templates/vacantion_report.docx"
            if report_type in _REPORT_TYPES_VACATION
            else "templates/vlk_report.docx"
        )
        if not os.path.isfile(template_path):
            flash(f"Не знайдено шаблон {template_label}", "error")
            return render_template('vlk_report.html', defaults=defaults)

        sluzhba_type = _normalize_spaces(request.form.get("sluzhba_type", "за мобілізацією"))
        if sluzhba_type not in ("за мобілізацією", "за контрактом"):
            sluzhba_type = "за мобілізацією"

        payload = {
            "report_type": report_type,
            "sluzhba_type": sluzhba_type,
            "misce_napravlennya": _normalize_spaces(request.form.get("misce_napravlennya", "")),
            "oglyad_nuber": _normalize_spaces(request.form.get("oglyad_nuber", "")),
            "likar": _normalize_spaces(request.form.get("likar", "")),
            "hospital": _normalize_spaces(request.form.get("hospital", "")),
            "dodatok": _normalize_spaces(request.form.get("dodatok", "")),
            "data_gospitalizatsii": _normalize_spaces(request.form.get("data_gospitalizatsii", "")),
            "vacantion_start_day": _normalize_spaces(request.form.get("vacantion_start_day", "")),
            "vacantion_adress": _normalize_spaces(request.form.get("vacantion_adress", "")),
            "vacantion_phone": _normalize_spaces(request.form.get("vacantion_phone", "")),
            "rehab_days": _normalize_spaces(request.form.get("rehab_days", "21")) or "21",
            "rehab_mrc": _normalize_spaces(request.form.get("rehab_mrc", "Пуща-Водиця")) or "Пуща-Водиця",
            "pib_nazivnyi": _normalize_spaces(request.form.get("pib_nazivnyi", "")),
            "tvo_kombrig_zvannya": _normalize_spaces(request.form.get("tvo_kombrig_zvannya", "")),
            "tvo_kombrig_pib": _normalize_spaces(request.form.get("tvo_kombrig_pib", "")),
            "tvo_kombat_zvannya": _normalize_spaces(request.form.get("tvo_kombat_zvannya", "")),
            "tvo_kombat_pib": _normalize_spaces(request.form.get("tvo_kombat_pib", "")),
        }
        defaults.update(payload)

        required_ok = all([
            payload["pib_nazivnyi"],
            payload["tvo_kombrig_zvannya"],
            payload["tvo_kombrig_pib"],
            payload["tvo_kombat_zvannya"],
            payload["tvo_kombat_pib"],
        ])
        if report_type in _REPORT_TYPES_VACATION:
            if not all([
                payload["vacantion_start_day"],
                payload["vacantion_adress"],
                payload["vacantion_phone"],
            ]):
                required_ok = False
        elif report_type in _REPORT_TYPES_FREE_DODATOK:
            if not payload["dodatok"]:
                required_ok = False
        elif report_type == "rehabilitation":
            if not payload["likar"] or not payload["hospital"]:
                required_ok = False
            if not payload["rehab_days"]:
                required_ok = False
        else:
            if not payload["likar"] or not payload["hospital"]:
                required_ok = False
        if report_type in ("hospitalization", "consultation") and not payload["misce_napravlennya"]:
            required_ok = False
        if report_type == "hospitalization" and not payload["data_gospitalizatsii"]:
            required_ok = False
        if not required_ok:
            flash("Заповніть усі обов'язкові поля рапорту.", "error")
            return render_template('vlk_report.html', defaults=defaults)

        person = _lookup_person_in_base_excel(payload["pib_nazivnyi"])
        if not person.get("zvanie") or not person.get("zaimana_posada"):
            try:
                treatments_df = load_treatments_data()
                pib_clean = re.sub(r'\s+', ' ', payload["pib_nazivnyi"]).strip().lower()
                soldier_records = treatments_df[treatments_df['ПІБ_чисте'] == pib_clean]
                if not soldier_records.empty:
                    row = soldier_records.iloc[0]
                    person["zvanie"] = person.get("zvanie") or _excel_cell_str(row.get('Військове звання'))
                    person["zaimana_posada"] = person.get("zaimana_posada") or _extract_position_from_row(row)
                    person["prizvyshche"] = person.get("prizvyshche") or _excel_cell_str(row.get('Прізвище'))
                    person["imya"] = person.get("imya") or _excel_cell_str(row.get("Ім'я"))
                    person["po_batkovi"] = person.get("po_batkovi") or _excel_cell_str(row.get('По батькові'))
            except Exception:
                pass

        zvanie = person.get("zvanie", "")
        zaimana_posada = person.get("zaimana_posada", "")
        sur = person.get("prizvyshche", "")
        first = person.get("imya", "")
        pat = person.get("po_batkovi", "")

        if not zvanie or not zaimana_posada:
            flash("Не вдалося автоматично підтягнути звання/посаду/ПІБ з base.xlsx. Уточніть ПІБ.", "error")
            return render_template('vlk_report.html', defaults=defaults)

        if not (sur and first):
            p_sur, p_first, p_pat = _split_pib(payload["pib_nazivnyi"])
            sur = sur or p_sur
            first = first or p_first
            pat = pat or p_pat

        pib_full_nazivnyi = _normalize_spaces(f"{sur} {first} {pat}")
        pib_gen_line = build_pib_rodovyi_for_document(pib_full_nazivnyi, "")
        pib_gen_title = _normalize_spaces(
            " ".join(_surname_first_capital(part) for part in pib_gen_line.split(" ") if part)
        )
        fraza_zvannya_ta_pib_rodovyi = _normalize_spaces(
            f"{format_rank_genitive(zvanie)} {pib_gen_title}"
        )
        tvo_kombrig_zvannya_dav = format_rank_dative(payload["tvo_kombrig_zvannya"])
        tvo_kombrig_pib_dav = _uppercase_last_word(_pib_to_dative(payload["tvo_kombrig_pib"]))
        tvo_kombat_zvannya_dav = format_rank_dative(payload["tvo_kombat_zvannya"])
        tvo_kombat_pib_dav = _uppercase_last_word(_pib_to_dative(payload["tvo_kombat_pib"]))
        tvo_kombrig_zvannya_short = _rank_short(payload["tvo_kombrig_zvannya"])
        tvo_kombrig_pib_short = _pib_short_with_initials(payload["tvo_kombrig_pib"])
        tekst_raportu = _build_report_body_text(
            report_type=report_type,
            zvanie=zvanie,
            sluzhba_type=sluzhba_type,
            pib_nazivnyi=pib_full_nazivnyi,
            misce_napravlennya=payload["misce_napravlennya"],
            data_gospitalizatsii=payload["data_gospitalizatsii"],
            rehab_days=payload["rehab_days"],
            rehab_mrc=payload["rehab_mrc"],
        )
        dodatok_line = _build_report_dodatok_line(
            report_type,
            dodatok=payload["dodatok"],
            oglyad_nuber=payload["oglyad_nuber"],
            likar=payload["likar"],
            hospital=payload["hospital"],
        )
        context = {
            "oglyad_nuber": payload["oglyad_nuber"],
            "likar": payload["likar"],
            "hospital": payload["hospital"],
            "dodatok": payload["dodatok"],
            "dodatok_line": dodatok_line,
            "tekst_raportu": tekst_raportu,
            "vacantion_start_day": payload["vacantion_start_day"],
            "vacantion_adress": payload["vacantion_adress"],
            "vacantion_phone": payload["vacantion_phone"],
            "zaimana_posada_z_velykoi": _title_first_letter(_ensure_final_period(zaimana_posada).rstrip(".")),
            "zvanie": zvanie,
            "imya_ta_prizvyshche_pidpys": _normalize_spaces(f"{first} {sur.upper()}"),
            "tvo_kombrig_zvannya": payload["tvo_kombrig_zvannya"],
            "tvo_kombrig_pib": payload["tvo_kombrig_pib"],
            "tvo_kombrig_zvannya_dav": tvo_kombrig_zvannya_dav,
            "tvo_kombrig_pib_dav": tvo_kombrig_pib_dav,
            "fraza_zvannya_ta_pib_rodovyi": fraza_zvannya_ta_pib_rodovyi,
            "tvo_kombat_zvannya": payload["tvo_kombat_zvannya"],
            "tvo_kombat_pib": payload["tvo_kombat_pib"],
            "tvo_kombat_zvannya_dav": tvo_kombat_zvannya_dav,
            "tvo_kombat_pib_dav": tvo_kombat_pib_dav,
            "tvo_kombrig_zvannya_short": tvo_kombrig_zvannya_short,
            "tvo_kombrig_pib_short": tvo_kombrig_pib_short,
            "data_stvorennya": datetime.now().strftime("%d.%m.%Y"),
            # Сумісність із шаблонами, де частина плейсхолдерів введена ВЕЛИКИМИ.
            "TVO_KOMBRIG_PIB_DAV": tvo_kombrig_pib_dav,
            "TVO_KOMBAT_PIB_DAV": tvo_kombat_pib_dav,
        }

        try:
            save_vlk_signatories({
                "tvo_kombrig_zvannya": payload["tvo_kombrig_zvannya"],
                "tvo_kombrig_pib": payload["tvo_kombrig_pib"],
                "tvo_kombat_zvannya": payload["tvo_kombat_zvannya"],
                "tvo_kombat_pib": payload["tvo_kombat_pib"],
            })
            # Підтримка шаблону з потенційними описками у назвах плейсхолдерів.
            context["tvo_kombat_ zvannya"] = payload["tvo_kombat_zvannya"]
            context["tvo_kombat_ zvannya_dav"] = tvo_kombat_zvannya_dav
            if report_type in _REPORT_TYPES_VACATION:
                doc = DocxTemplate(VACATION_REPORT_TEMPLATE)
            else:
                doc = _load_vlk_report_template()

            doc.render(context)
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
        except Exception as e:
            logger.error("Помилка при генерації рапорту: %s", e)
            flash(f"Помилка генерації DOCX: {e}", "error")
            return render_template('vlk_report.html', defaults=defaults)

        file_prefix = {
            "vlk": "Рапорт_на_ВЛК",
            "hospitalization": "Рапорт_на_госпіталізацію",
            "consultation": "Рапорт_на_консультацію",
            "material_aid": "Рапорт_на_матеріальну_допомогу",
            "ozdorovlennya": "Рапорт_на_оздоровлення",
            "vacation": "Рапорт_на_відпустку",
            "rehabilitation": "Рапорт_на_реабілітацію",
        }.get(report_type, "Рапорт")
        pib_for_file = _normalize_spaces(f"{first} {sur}".strip()) or pib_full_nazivnyi or sur
        safe_name = _safe_download_stem(pib_for_file, "report")
        _loader_preview_delay()
        response = make_response(send_file(
            file_stream, as_attachment=True,
            download_name=f'{file_prefix}_{safe_name}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ))
        response.set_cookie('fileDownload', 'true', max_age=20)
        return response

    return render_template('vlk_report.html', defaults=defaults)


# ——— Амбулаторний журнал (SQLite) ———

OUTPATIENT_COLUMNS = [
    "id",
    "date",
    "pib",
    "rank_unit",
    "diagnosis",
    "referral_to",
    "lpz",
    "exam_result",
    "duty_exempt_days",
    "visit_type",
    "mkx10",
    "notes",
    "treatment_id",
    "care_type",
    "wa_event",
    "from_lpz",
    "leave_start",
    "leave_end",
    "leave_days",
]

# У шаблоні книги клітинки злиті: індекси перших унікальних tc у рядку.
_OUTPATIENT_LEFT_COLS = (0, 2, 3, 4, 6)   # №, дата, ПІБ, звання, діагноз
_OUTPATIENT_RIGHT_COLS = (0, 4, 6, 7, 8)  # направлення, звільнення, І/ІІ, МКХ, примітки
_OUTPATIENT_MAX_ROWS = 43

_OUTPATIENT_SEED = [
    {
        "date": "01.08.2026",
        "pib": "Коваленко Андрій Сергійович",
        "rank_unit": "солдат, 1 рота",
        "diagnosis": "ГРВІ",
        "referral_to": "терапевта",
        "exam_result": "призначено симптоматичне лікування",
        "duty_exempt_days": "1",
        "visit_type": "І",
        "mkx10": "J06.9",
        "notes": "",
    },
    {
        "date": "02.08.2026",
        "pib": "Шевченко Олег Миколайович",
        "rank_unit": "старший солдат, 2 рота",
        "diagnosis": "Люмбаго",
        "referral_to": "хірурга",
        "exam_result": "рекомендовано консультацію",
        "duty_exempt_days": "2",
        "visit_type": "І",
        "mkx10": "M54.5",
        "notes": "",
    },
    {
        "date": "03.08.2026",
        "pib": "Бондаренко Ігор Васильович",
        "rank_unit": "молодший сержант, розвідка",
        "diagnosis": "Отит зовнішній",
        "referral_to": "ЛОРа",
        "exam_result": "направлено до ТМО",
        "duty_exempt_days": "",
        "visit_type": "І",
        "mkx10": "H60.9",
        "notes": "",
    },
    {
        "date": "04.08.2026",
        "pib": "Мельник Дмитро Олександрович",
        "rank_unit": "сержант, мінометна батарея",
        "diagnosis": "Гіпертонічна хвороба, криз",
        "referral_to": "терапевта",
        "exam_result": "зниження тиску, спостереження",
        "duty_exempt_days": "1",
        "visit_type": "ІІ",
        "mkx10": "I10",
        "notes": "повторне звернення",
    },
    {
        "date": "05.08.2026",
        "pib": "Ткаченко Віталій Петрович",
        "rank_unit": "солдат, 3 рота",
        "diagnosis": "Ушиб м'яких тканин гомілки",
        "referral_to": "хірурга",
        "exam_result": "пов'язка, холод",
        "duty_exempt_days": "3",
        "visit_type": "І",
        "mkx10": "S80.0",
        "notes": "",
    },
    {
        "date": "05.08.2026",
        "pib": "Кравченко Сергій Іванович",
        "rank_unit": "старший сержант, штаб",
        "diagnosis": "Кон'юнктивіт",
        "referral_to": "офтальмолога",
        "exam_result": "призначено краплі",
        "duty_exempt_days": "",
        "visit_type": "І",
        "mkx10": "H10.9",
        "notes": "",
    },
    {
        "date": "06.08.2026",
        "pib": "Петренко Максим Андрійович",
        "rank_unit": "солдат, інженерна",
        "diagnosis": "Ангіна катаральна",
        "referral_to": "терапевта",
        "exam_result": "антибактеріальна терапія",
        "duty_exempt_days": "2",
        "visit_type": "І",
        "mkx10": "J03.9",
        "notes": "",
    },
    {
        "date": "06.08.2026",
        "pib": "Сидоренко Юрій Володимирович",
        "rank_unit": "капітан, 2 БОП",
        "diagnosis": "Гастрит загострення",
        "referral_to": "гастроентеролога",
        "exam_result": "дієта, обстеження",
        "duty_exempt_days": "1",
        "visit_type": "І",
        "mkx10": "K29.1",
        "notes": "",
    },
    {
        "date": "07.08.2026",
        "pib": "Гончар Павло Романович",
        "rank_unit": "солдат, охорона",
        "diagnosis": "Алергічний дерматит",
        "referral_to": "дерматолога",
        "exam_result": "антигістамінні",
        "duty_exempt_days": "",
        "visit_type": "І",
        "mkx10": "L23.9",
        "notes": "",
    },
    {
        "date": "08.08.2026",
        "pib": "Лисенко Артем Олегович",
        "rank_unit": "старший солдат, 1 рота",
        "diagnosis": "Головний біль напруги",
        "referral_to": "невролога",
        "exam_result": "спокій, знеболювальне",
        "duty_exempt_days": "1",
        "visit_type": "ІІ",
        "mkx10": "G44.2",
        "notes": "після фізо",
    },
]


def _outpatient_db_connect() -> sqlite3.Connection:
    os.makedirs(DATA_DIR, exist_ok=True)
    conn = sqlite3.connect(OUTPATIENT_JOURNAL_DB)
    conn.row_factory = sqlite3.Row
    return conn


def _outpatient_init_db() -> None:
    """Створює таблицю; міграція/тестові дані — лише при першому створенні файлу БД."""
    os.makedirs(DATA_DIR, exist_ok=True)
    first_create = not os.path.isfile(OUTPATIENT_JOURNAL_DB)
    conn = _outpatient_db_connect()
    try:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS outpatient_entries (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                date TEXT NOT NULL DEFAULT '',
                pib TEXT NOT NULL DEFAULT '',
                rank_unit TEXT NOT NULL DEFAULT '',
                diagnosis TEXT NOT NULL DEFAULT '',
                referral_to TEXT NOT NULL DEFAULT '',
                lpz TEXT NOT NULL DEFAULT '',
                exam_result TEXT NOT NULL DEFAULT '',
                duty_exempt_days TEXT NOT NULL DEFAULT '',
                visit_type TEXT NOT NULL DEFAULT '',
                mkx10 TEXT NOT NULL DEFAULT '',
                notes TEXT NOT NULL DEFAULT '',
                created_at TEXT NOT NULL DEFAULT (datetime('now','localtime'))
            )
            """
        )
        cols = {row[1] for row in conn.execute("PRAGMA table_info(outpatient_entries)").fetchall()}
        if "lpz" not in cols:
            conn.execute(
                "ALTER TABLE outpatient_entries ADD COLUMN lpz TEXT NOT NULL DEFAULT ''"
            )
        conn.commit()
        cards_db.ensure_card_schema(conn)
        # Не засівати знову після повного очищення журналу користувачем.
        if not first_create:
            return

        count = conn.execute("SELECT COUNT(*) AS c FROM outpatient_entries").fetchone()["c"]
        if count == 0 and os.path.isfile(OUTPATIENT_JOURNAL_FILE):
            try:
                df = pd.read_excel(OUTPATIENT_JOURNAL_FILE, dtype=str).fillna("")
                for _, row in df.iterrows():
                    conn.execute(
                        """
                        INSERT INTO outpatient_entries
                        (date, pib, rank_unit, diagnosis, referral_to, exam_result,
                         duty_exempt_days, visit_type, mkx10, notes)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """,
                        (
                            _normalize_spaces(str(row.get("date", ""))),
                            _normalize_spaces(str(row.get("pib", ""))),
                            _normalize_spaces(str(row.get("rank_unit", ""))),
                            _normalize_spaces(str(row.get("diagnosis", ""))),
                            _normalize_spaces(str(row.get("referral_to", ""))),
                            _normalize_spaces(str(row.get("exam_result", ""))),
                            _normalize_spaces(str(row.get("duty_exempt_days", ""))),
                            _normalize_spaces(str(row.get("visit_type", ""))),
                            _normalize_spaces(str(row.get("mkx10", ""))),
                            _normalize_spaces(str(row.get("notes", ""))),
                        ),
                    )
                conn.commit()
                count = conn.execute("SELECT COUNT(*) AS c FROM outpatient_entries").fetchone()["c"]
                logger.info("Мігровано %s записів журналу з Excel у SQLite", count)
            except Exception as e:
                logger.warning("Міграція outpatient_journal.xlsx не вдалася: %s", e)
        count = conn.execute("SELECT COUNT(*) AS c FROM outpatient_entries").fetchone()["c"]
        if count == 0:
            for rec in _OUTPATIENT_SEED:
                conn.execute(
                    """
                    INSERT INTO outpatient_entries
                    (date, pib, rank_unit, diagnosis, referral_to, exam_result,
                     duty_exempt_days, visit_type, mkx10, notes)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        rec["date"],
                        rec["pib"],
                        rec["rank_unit"],
                        rec["diagnosis"],
                        rec["referral_to"],
                        rec["exam_result"],
                        rec["duty_exempt_days"],
                        rec["visit_type"],
                        rec["mkx10"],
                        rec["notes"],
                    ),
                )
            conn.commit()
            logger.info("Додано %s тестових записів амбулаторного журналу", len(_OUTPATIENT_SEED))
    finally:
        conn.close()


_OUTPATIENT_UNIT_SHTAB = "Штаб"
_OUTPATIENT_UNIT_ROTA_1 = "1-ша рота оперативного призначення (на бронетранспортерах)"
_OUTPATIENT_UNIT_ROTA_2 = "2-га рота оперативного призначення (на бронетранспортерах)"
_OUTPATIENT_UNIT_ROTA_3 = "3-тя рота оперативного призначення (на бронетранспортерах)"
_OUTPATIENT_UNIT_MORTAR_120 = "Мінометна батарея (120 мм міномети)"
_OUTPATIENT_UNIT_MORTAR_60 = "Мінометна батарея (60 мм (82 мм) міномети)"
_OUTPATIENT_UNIT_FIRE_SUPPORT = "Рота вогневої підтримки"
_OUTPATIENT_UNIT_COMMS = "Взвод зв'язку"
_OUTPATIENT_UNIT_RECON = "Розвідувальний взвод спеціального призначення"
_OUTPATIENT_UNIT_ENGINEER = "Інженерно-саперне відділення"
_OUTPATIENT_UNIT_UAV = "Взвод безпілотних комплексів спеціального призначення"
_OUTPATIENT_UNIT_LOGISTICS = "Взвод матеріально-технічного забезпечення"
_OUTPATIENT_UNIT_MAINT = "Взвод технічного обслуговування озброєння та військової техніки"
_OUTPATIENT_UNIT_MED = "Медичний пункт"

_OUTPATIENT_CANONICAL_UNITS = (
    _OUTPATIENT_UNIT_SHTAB,
    _OUTPATIENT_UNIT_ROTA_1,
    _OUTPATIENT_UNIT_ROTA_2,
    _OUTPATIENT_UNIT_ROTA_3,
    _OUTPATIENT_UNIT_MORTAR_120,
    _OUTPATIENT_UNIT_MORTAR_60,
    _OUTPATIENT_UNIT_FIRE_SUPPORT,
    _OUTPATIENT_UNIT_COMMS,
    _OUTPATIENT_UNIT_RECON,
    _OUTPATIENT_UNIT_ENGINEER,
    _OUTPATIENT_UNIT_UAV,
    _OUTPATIENT_UNIT_LOGISTICS,
    _OUTPATIENT_UNIT_MAINT,
    _OUTPATIENT_UNIT_MED,
)

_OUTPATIENT_ROTA_COMPANY_RE = re.compile(
    r"(?P<num>\d+)\s*[-‑]?\s*(?:ша|га|тя|ої|ї)\s+роти\b",
    re.IGNORECASE,
)

# Специфічні підрозділи (порядок важливий — спочатку вужчі патерни).
_OUTPATIENT_UNIT_PATTERNS = (
    (re.compile(r"медичн\w*\s+пункт", re.IGNORECASE), _OUTPATIENT_UNIT_MED),
    (re.compile(r"мінометн\w*\s+батаре.*120\s*мм|120\s*мм.*міномет", re.IGNORECASE), _OUTPATIENT_UNIT_MORTAR_120),
    (re.compile(r"мінометн\w*\s+батаре.*(?:60|82)\s*мм|(?:60|82)\s*мм.*міномет", re.IGNORECASE), _OUTPATIENT_UNIT_MORTAR_60),
    (re.compile(r"мінометн\w*\s+батаре", re.IGNORECASE), _OUTPATIENT_UNIT_MORTAR_120),
    (re.compile(r"рота\s+вогневої\s+підтримк|вогневої\s+підтримк", re.IGNORECASE), _OUTPATIENT_UNIT_FIRE_SUPPORT),
    (re.compile(r"взвод\s+зв.?язку|\bзв.?язку\b", re.IGNORECASE), _OUTPATIENT_UNIT_COMMS),
    (re.compile(r"розвідувальн\w*\s+взвод|розвідк\w*\s+(?:взвод|відділен)", re.IGNORECASE), _OUTPATIENT_UNIT_RECON),
    (re.compile(r"інженерно[-\s]?саперн|інженерн\w*\s+відділен", re.IGNORECASE), _OUTPATIENT_UNIT_ENGINEER),
    (re.compile(r"безпілотн|бпк|бпла", re.IGNORECASE), _OUTPATIENT_UNIT_UAV),
    (re.compile(r"матеріально[-\s]?технічн\w*\s+забезпечен|\bмтз\b", re.IGNORECASE), _OUTPATIENT_UNIT_LOGISTICS),
    (re.compile(r"технічного\s+обслуговування|ремонтн\w*\s+взвод|\bтовт\b", re.IGNORECASE), _OUTPATIENT_UNIT_MAINT),
    (re.compile(r"\bштаб", re.IGNORECASE), _OUTPATIENT_UNIT_SHTAB),
)

_OUTPATIENT_ROTA_BY_NUM = {
    1: _OUTPATIENT_UNIT_ROTA_1,
    2: _OUTPATIENT_UNIT_ROTA_2,
    3: _OUTPATIENT_UNIT_ROTA_3,
}


def _parse_outpatient_unit(rank_unit: str) -> str:
    """
    Нормалізує підрозділ із рядка «звання, посада…» до канонічної назви.
    Приклади:
      «…2-ої роти оперативного призначення…» → 2-га рота ОП (на БТР)
      «…медичного пункту…» → Медичний пункт
    """
    text = _normalize_spaces(rank_unit)
    if not text:
        return ""
    text_cf = text.casefold()
    for label in _OUTPATIENT_CANONICAL_UNITS:
        if label.casefold() in text_cf:
            return label

    # «1-ої / 2-ої / 3-ої роти» (або 1-ша / 2-га / 3-тя роти)
    rota_hits = list(_OUTPATIENT_ROTA_COMPANY_RE.finditer(text))
    if rota_hits:
        num = int(rota_hits[-1].group("num"))
        if num in _OUTPATIENT_ROTA_BY_NUM:
            return _OUTPATIENT_ROTA_BY_NUM[num]

    for rx, label in _OUTPATIENT_UNIT_PATTERNS:
        if rx.search(text):
            return label
    return ""


def list_outpatient_units() -> list:
    """Повний фіксований список підрозділів батальйону для фільтра."""
    return list(_OUTPATIENT_CANONICAL_UNITS)


def list_outpatient_entries(
    order_desc: bool = True,
    pib_query: str = "",
    unit: str = "",
) -> list:
    _outpatient_init_db()
    order = "DESC" if order_desc else "ASC"
    conn = _outpatient_db_connect()
    try:
        rows = conn.execute(
            f"SELECT * FROM outpatient_entries ORDER BY id {order}"
        ).fetchall()
        records = [dict(r) for r in rows]
    finally:
        conn.close()
    # SQLite NOCASE/LOWER не працює для кирилиці — фільтр у Python.
    q = _normalize_spaces(pib_query).casefold()
    if q:
        records = [
            r for r in records
            if q in _normalize_spaces(r.get("pib", "")).casefold()
        ]
    unit_q = _normalize_spaces(unit)
    if unit_q:
        records = [
            r for r in records
            if _parse_outpatient_unit(r.get("rank_unit", "")) == unit_q
        ]
    return records


def list_outpatient_entries_by_date(date_str: str) -> list:
    """Записи за дату у форматі дд.мм.рррр (порядок зростання id)."""
    day = _normalize_spaces(date_str)
    if not day:
        return []
    _outpatient_init_db()
    conn = _outpatient_db_connect()
    try:
        rows = conn.execute(
            """
            SELECT * FROM outpatient_entries
            WHERE date = ?
            ORDER BY id ASC
            """,
            (day,),
        ).fetchall()
        return [dict(r) for r in rows]
    finally:
        conn.close()


def _iso_to_outpatient_date(iso_date: str) -> str:
    """YYYY-MM-DD → дд.мм.рррр."""
    raw = _normalize_spaces(iso_date)
    if not raw:
        return ""
    try:
        return datetime.strptime(raw, "%Y-%m-%d").strftime("%d.%m.%Y")
    except ValueError:
        if re.fullmatch(r"\d{2}\.\d{2}\.\d{4}", raw):
            return raw
        return ""


def search_outpatient_pib(query: str, limit: int = 15) -> list:
    """Автокомпліт ПІБ за записами журналу (кирилиця без урахування регістру)."""
    q = _normalize_spaces(query).casefold()
    if len(q) < 1:
        return []
    _outpatient_init_db()
    conn = _outpatient_db_connect()
    try:
        rows = conn.execute(
            """
            SELECT DISTINCT pib FROM outpatient_entries
            WHERE pib IS NOT NULL AND TRIM(pib) != ''
            ORDER BY pib
            """
        ).fetchall()
    finally:
        conn.close()
    hits = []
    for r in rows:
        pib = _normalize_spaces(r["pib"])
        if not pib:
            continue
        if q in pib.casefold():
            hits.append(pib)
        if len(hits) >= int(limit):
            break
    return hits


def get_outpatient_entry(entry_id) -> dict | None:
    _outpatient_init_db()
    try:
        eid = int(entry_id)
    except (TypeError, ValueError):
        return None
    conn = _outpatient_db_connect()
    try:
        row = conn.execute(
            "SELECT * FROM outpatient_entries WHERE id = ?", (eid,)
        ).fetchone()
        return dict(row) if row else None
    finally:
        conn.close()


def insert_outpatient_entry(payload: dict) -> int:
    _outpatient_init_db()
    conn = _outpatient_db_connect()
    try:
        tid = payload.get("treatment_id")
        try:
            tid_val = int(tid) if tid not in (None, "") else None
        except (TypeError, ValueError):
            tid_val = None
        cur = conn.execute(
            """
            INSERT INTO outpatient_entries
            (date, pib, rank_unit, diagnosis, referral_to, lpz, exam_result,
             duty_exempt_days, visit_type, mkx10, notes, treatment_id,
             care_type, from_lpz, leave_start, leave_end, leave_days, wa_event)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                payload.get("date", ""),
                payload.get("pib", ""),
                payload.get("rank_unit", ""),
                payload.get("diagnosis", ""),
                payload.get("referral_to", ""),
                payload.get("lpz", ""),
                payload.get("exam_result", ""),
                payload.get("duty_exempt_days", ""),
                payload.get("visit_type", ""),
                payload.get("mkx10", ""),
                payload.get("notes", ""),
                tid_val,
                payload.get("care_type", ""),
                payload.get("from_lpz", ""),
                payload.get("leave_start", ""),
                payload.get("leave_end", ""),
                payload.get("leave_days", ""),
                payload.get("wa_event", ""),
            ),
        )
        conn.commit()
        return int(cur.lastrowid)
    finally:
        conn.close()


def update_outpatient_entry(entry_id, payload: dict) -> bool:
    _outpatient_init_db()
    try:
        eid = int(entry_id)
    except (TypeError, ValueError):
        return False
    conn = _outpatient_db_connect()
    try:
        tid = payload.get("treatment_id")
        try:
            tid_val = int(tid) if tid not in (None, "") else None
        except (TypeError, ValueError):
            tid_val = None
        cur = conn.execute(
            """
            UPDATE outpatient_entries SET
                date = ?, pib = ?, rank_unit = ?, diagnosis = ?,
                referral_to = ?, lpz = ?, exam_result = ?, duty_exempt_days = ?,
                visit_type = ?, mkx10 = ?, notes = ?, treatment_id = ?,
                care_type = ?, from_lpz = ?, leave_start = ?, leave_end = ?, leave_days = ?,
                wa_event = ?
            WHERE id = ?
            """,
            (
                payload.get("date", ""),
                payload.get("pib", ""),
                payload.get("rank_unit", ""),
                payload.get("diagnosis", ""),
                payload.get("referral_to", ""),
                payload.get("lpz", ""),
                payload.get("exam_result", ""),
                payload.get("duty_exempt_days", ""),
                payload.get("visit_type", ""),
                payload.get("mkx10", ""),
                payload.get("notes", ""),
                tid_val,
                payload.get("care_type", ""),
                payload.get("from_lpz", ""),
                payload.get("leave_start", ""),
                payload.get("leave_end", ""),
                payload.get("leave_days", ""),
                payload.get("wa_event", ""),
                eid,
            ),
        )
        conn.commit()
        return cur.rowcount > 0
    finally:
        conn.close()


def delete_outpatient_entry(entry_id) -> bool:
    _outpatient_init_db()
    try:
        eid = int(entry_id)
    except (TypeError, ValueError):
        return False
    conn = _outpatient_db_connect()
    try:
        cur = conn.execute("DELETE FROM outpatient_entries WHERE id = ?", (eid,))
        conn.commit()
        return cur.rowcount > 0
    finally:
        conn.close()


def _outpatient_today() -> str:
    return datetime.now().strftime("%d.%m.%Y")


def _outpatient_row_dict(form) -> dict:
    leave_start = _normalize_spaces(form.get("leave_start", ""))
    leave_end = _normalize_spaces(form.get("leave_end", ""))
    leave_days = _normalize_spaces(form.get("leave_days", ""))
    visit_date = _normalize_spaces(form.get("date", "")) or _outpatient_today()
    # автопідрахунок днів з дат, якщо не вказано
    if leave_start and leave_end and not leave_days:
        start_dt = cards_db.parse_ui_date(leave_start)
        end_dt = cards_db.parse_ui_date(leave_end)
        if start_dt and end_dt:
            leave_days = str(cards_db.calendar_days_inclusive(start_dt, end_dt))
    elif leave_start and leave_days and not leave_end:
        start_dt = cards_db.parse_ui_date(leave_start)
        try:
            days_n = int(leave_days)
        except ValueError:
            days_n = 0
        if start_dt and days_n > 0:
            leave_end = cards_db.format_ui_date(cards_db.add_days_inclusive(start_dt, days_n))
    care_type = _normalize_spaces(form.get("care_type", ""))
    wa_event = _normalize_spaces(form.get("wa_event", ""))
    from_lpz = _normalize_spaces(form.get("from_lpz", ""))
    if wa_event not in (
        cards_db.WA_EVENT_OPEN,
        cards_db.WA_EVENT_EXTEND,
        cards_db.WA_EVENT_DISCHARGE,
    ):
        wa_event = ""
    if care_type not in cards_db.CARE_WA_EVENTS:
        # ВЛК без селекту події — завжди «початок»
        if care_type == "vlk":
            wa_event = cards_db.WA_EVENT_OPEN
        else:
            wa_event = ""
    elif wa_event and wa_event not in {k for k, _ in cards_db.CARE_WA_EVENTS.get(care_type, [])}:
        wa_event = ""
    # Консультація / направлення — без періоду «з/по»
    if care_type in ("consultation", "referral"):
        leave_start = ""
        leave_end = ""
        leave_days = ""
        from_lpz = ""
    elif care_type in ("outpatient", "day_hospital"):
        # одна ключова дата залежно від події
        if wa_event == cards_db.WA_EVENT_OPEN:
            leave_start = leave_start or visit_date
            if not leave_end:
                leave_end = leave_start
        elif wa_event == cards_db.WA_EVENT_EXTEND:
            leave_end = leave_end or visit_date
        elif wa_event == cards_db.WA_EVENT_DISCHARGE:
            leave_end = leave_end or visit_date
        from_lpz = ""
    elif care_type == "inpatient":
        if wa_event == cards_db.WA_EVENT_OPEN:
            leave_start = leave_start or visit_date
            leave_end = ""
            leave_days = ""
            from_lpz = ""
        elif wa_event == cards_db.WA_EVENT_EXTEND:
            # переведений — дата переводу в date / leave_start
            leave_start = leave_start or visit_date
            leave_end = ""
            leave_days = ""
        elif wa_event == cards_db.WA_EVENT_DISCHARGE:
            from_lpz = ""
            if leave_start and leave_end and not leave_days:
                start_dt = cards_db.parse_ui_date(leave_start)
                end_dt = cards_db.parse_ui_date(leave_end)
                if start_dt and end_dt:
                    leave_days = str(cards_db.calendar_days_inclusive(start_dt, end_dt))
    elif care_type == "vlk":
        leave_start = leave_start or visit_date
        leave_end = ""
        leave_days = ""
        from_lpz = ""
        wa_event = cards_db.WA_EVENT_OPEN
    return {
        "date": visit_date,
        "pib": _normalize_spaces(form.get("pib", "")),
        "rank_unit": _normalize_spaces(form.get("rank_unit", "")),
        "diagnosis": _normalize_spaces(form.get("diagnosis", "")),
        "referral_to": _normalize_spaces(form.get("referral_to", "")),
        "lpz": _normalize_spaces(form.get("lpz", "")),
        "exam_result": _normalize_spaces(form.get("exam_result", "")),
        "duty_exempt_days": leave_days,
        "visit_type": "",
        "mkx10": _normalize_spaces(form.get("mkx10", "")),
        "notes": _normalize_spaces(form.get("notes", "")),
        "treatment_id": _normalize_spaces(form.get("treatment_id", "")),
        "new_treatment_title": _normalize_spaces(form.get("new_treatment_title", "")),
        "care_type": care_type,
        "wa_event": wa_event,
        "from_lpz": from_lpz,
        "leave_start": leave_start,
        "leave_end": leave_end,
        "leave_days": leave_days,
    }


def _set_docx_cell_text(cell, text: str) -> None:
    cell.text = str(text or "")


def _outpatient_page_count(total: int) -> int:
    if total <= 0:
        return 0
    return (total + _OUTPATIENT_MAX_ROWS - 1) // _OUTPATIENT_MAX_ROWS


def _fill_outpatient_journal_docx(records: list, page: int = 1) -> io.BytesIO:
    """
    Заповнює одну сторінку книги обліку (до 43 записів).
    page — номер сторінки з 1; № з/п у бланку = id запису.
    """
    if not os.path.isfile(OUTPATIENT_JOURNAL_TEMPLATE):
        raise FileNotFoundError("Не знайдено templates/record_keeping_for_outpatients.docx")
    page = max(1, int(page or 1))
    start = (page - 1) * _OUTPATIENT_MAX_ROWS
    batch = records[start: start + _OUTPATIENT_MAX_ROWS]
    if not batch:
        raise ValueError(f"На сторінці {page} немає записів")

    doc = DocxDocument(OUTPATIENT_JOURNAL_TEMPLATE)
    if len(doc.tables) < 2:
        raise ValueError("У шаблоні журналу немає очікуваної таблиці")
    table = doc.tables[1]
    for i, rec in enumerate(batch, start=1):
        left_ri = i
        right_ri = 44 + i
        if left_ri >= len(table.rows) or right_ri >= len(table.rows):
            break
        left = table.rows[left_ri]
        right = table.rows[right_ri]
        num = rec.get("id", i)
        left_vals = [
            str(num),
            rec.get("date", ""),
            rec.get("pib", ""),
            rec.get("rank_unit", ""),
            rec.get("diagnosis", ""),
        ]
        notes_parts = [rec.get("exam_result", ""), rec.get("notes", "")]
        notes = _normalize_spaces(" ".join(p for p in notes_parts if p))
        where_parts = [rec.get("referral_to", ""), rec.get("lpz", "")]
        where_to = _normalize_spaces(" → ".join(p for p in where_parts if p))
        right_vals = [
            where_to,
            rec.get("duty_exempt_days", ""),
            rec.get("visit_type", ""),
            rec.get("mkx10", ""),
            notes,
        ]
        for col_i, val in zip(_OUTPATIENT_LEFT_COLS, left_vals):
            if col_i < len(left.cells):
                _set_docx_cell_text(left.cells[col_i], val)
        for col_i, val in zip(_OUTPATIENT_RIGHT_COLS, right_vals):
            if col_i < len(right.cells):
                _set_docx_cell_text(right.cells[col_i], val)
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _push_media_to_dropbox(treatment_id: int, media: dict, data: bytes = None) -> str:
    """
    Копіює медіа в Dropbox:
    /{root}/patients/{ПІБ}/{лікування}/файл
    Повертає шлях або "" якщо Dropbox не налаштовано.
    """
    from utils.dropbox_sync import (
        dropbox_configured,
        patient_media_dropbox_path,
        upload_bytes_to_dropbox,
    )

    if not dropbox_configured():
        return ""
    treatment = cards_db.get_treatment(int(treatment_id))
    if not treatment:
        raise ValueError("Лікування не знайдено для Dropbox")
    blob = data
    if blob is None:
        path = cards_db.media_disk_path(media)
        with open(path, "rb") as fh:
            blob = fh.read()
    dropbox_path = patient_media_dropbox_path(
        pib=treatment.get("pib") or "patient",
        treatment_title=treatment.get("title") or "treatment",
        treatment_id=int(treatment_id),
        filename=media.get("original_name") or media.get("filename") or "file",
    )
    return upload_bytes_to_dropbox(blob, dropbox_path) or ""


def _save_uploaded_media_list(
    treatment_id: int,
    storages,
    *,
    visit_id: int = None,
) -> tuple[int, int, str]:
    """
    Зберігає список FileStorage у медіа + Dropbox.
    Повертає (saved_count, dropbox_count, dropbox_error).
    """
    saved = 0
    synced = 0
    dropbox_error = ""
    for storage in storages or []:
        if not storage or not storage.filename:
            continue
        try:
            media = cards_db.save_media_file(
                int(treatment_id),
                storage,
                visit_id=visit_id,
            )
            saved += 1
            try:
                if _push_media_to_dropbox(
                    int(treatment_id),
                    media,
                    data=media.get("bytes"),
                ):
                    synced += 1
            except Exception as e:
                dropbox_error = dropbox_error or str(e)
                logger.warning("Dropbox sync media #%s: %s", media.get("id"), e)
        except Exception as e:
            logger.warning("Не збережено медіа %s: %s", storage.filename, e)
    return saved, synced, dropbox_error


@app.route("/api/dropbox_status", methods=["GET"])
def api_dropbox_status():
    """Діагностика підключення до Dropbox."""
    from utils.dropbox_sync import check_connection

    return jsonify(check_connection())


@app.route("/api/dropbox_resync", methods=["POST"])
def api_dropbox_resync():
    """Повторно заливає в Dropbox усі локальні медіа (overwrite)."""
    from utils.dropbox_sync import check_connection

    status = check_connection()
    if not status.get("ok"):
        return jsonify({
            "ok": False,
            "error": status.get("error") or "Dropbox недоступний",
        }), 400

    uploaded = 0
    skipped = 0
    failed = 0
    first_error = ""
    for media in cards_db.list_all_media():
        try:
            disk_path = cards_db.media_disk_path(media)
            if not os.path.isfile(disk_path):
                skipped += 1
                continue
            if _push_media_to_dropbox(int(media["treatment_id"]), media):
                uploaded += 1
            else:
                skipped += 1
        except Exception as e:
            failed += 1
            first_error = first_error or str(e)
            logger.warning("Dropbox resync media #%s: %s", media.get("id"), e)
    return jsonify({
        "ok": True,
        "uploaded": uploaded,
        "skipped": skipped,
        "failed": failed,
        "error": first_error,
    })


@app.route("/api/outpatient_pib", methods=["GET"])
def api_outpatient_pib():
    """Автокомпліт ПІБ за записами амбулаторного журналу."""
    q = request.args.get("q", "").strip()
    if len(q) < 1:
        return jsonify({"results": []})
    names = search_outpatient_pib(q, limit=15)
    return jsonify({
        "results": [{"value": name, "label": name} for name in names],
    })


@app.route("/api/extract_visit_text", methods=["POST"])
def api_extract_visit_text():
    """Зчитує діагноз і рекомендації з фото/PDF через Gemini; за наявності treatment_id — зберігає в медіа."""
    from utils.gemini_extract import (
        ALLOWED_MIME,
        extract_diagnosis_from_parts,
        gemini_configured,
    )

    if not gemini_configured():
        return jsonify({
            "ok": False,
            "error": "Не задано GEMINI_API_KEY у файлі .env",
        }), 503
    files = request.files.getlist("files") or request.files.getlist("file")
    parts = []
    file_payloads = []  # (filename, bytes, mime) для збереження в медіа
    for storage in files:
        if not storage or not storage.filename:
            continue
        mime = (storage.mimetype or "").lower().strip()
        if mime not in ALLOWED_MIME:
            name = (storage.filename or "").lower()
            if name.endswith((".jpg", ".jpeg")):
                mime = "image/jpeg"
            elif name.endswith(".png"):
                mime = "image/png"
            elif name.endswith(".webp"):
                mime = "image/webp"
            elif name.endswith(".gif"):
                mime = "image/gif"
            elif name.endswith(".pdf"):
                mime = "application/pdf"
            elif name.endswith((".heic", ".heif")):
                mime = "image/heic"
        blob = storage.read()
        if not blob:
            continue
        parts.append((blob, mime))
        file_payloads.append((storage.filename, blob, mime))
        if len(parts) >= 5:
            break
    if not parts:
        return jsonify({"ok": False, "error": "Додайте фото або PDF висновку."}), 400
    try:
        extracted = extract_diagnosis_from_parts(parts)
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    except Exception as e:
        logger.exception("Gemini extract")
        return jsonify({"ok": False, "error": str(e)}), 502

    media_saved = 0
    dropbox_synced = 0
    dropbox_error = ""
    tid_raw = _normalize_spaces(request.form.get("treatment_id", ""))
    visit_raw = _normalize_spaces(request.form.get("visit_id", ""))
    try:
        tid = int(tid_raw) if tid_raw and tid_raw != "new" else None
    except (TypeError, ValueError):
        tid = None
    try:
        visit_id = int(visit_raw) if visit_raw else None
    except (TypeError, ValueError):
        visit_id = None
    if tid:
        for filename, blob, mime in file_payloads:
            try:
                media = cards_db.save_media_bytes(
                    tid,
                    blob,
                    filename,
                    visit_id=visit_id,
                    mime_hint=mime,
                )
                media_saved += 1
                try:
                    if _push_media_to_dropbox(tid, media, data=blob):
                        dropbox_synced += 1
                except Exception as e:
                    dropbox_error = dropbox_error or str(e)
                    logger.warning("Dropbox після Gemini: %s", e)
            except Exception as e:
                logger.warning("Медіа після Gemini (%s): %s", filename, e)

    return jsonify({
        "ok": True,
        "diagnosis": extracted.get("diagnosis") or "",
        "recommendations": extracted.get("recommendations") or "",
        "media_saved": media_saved,
        "dropbox_synced": dropbox_synced,
        "dropbox_error": dropbox_error,
    })


@app.route("/outpatient-journal", methods=["GET", "POST"])
def outpatient_journal():
    """Ведення амбулаторного журналу (SQLite)."""
    _outpatient_init_db()
    edit_id = request.args.get("edit", "").strip()
    edit_row = get_outpatient_entry(edit_id) if edit_id else None

    if request.method == "POST":
        action = _normalize_spaces(request.form.get("action", "save"))

        if action == "delete":
            entry_id = request.form.get("id", "").strip()
            if delete_outpatient_entry(entry_id):
                flash("Запис видалено.", "success")
            else:
                flash("Запис не знайдено.", "error")
            return redirect(url_for("outpatient_journal"))

        payload = _outpatient_row_dict(request.form)
        if not payload["pib"]:
            flash("Вкажіть ПІБ.", "error")
            return redirect(url_for("outpatient_journal"))

        entry_id = request.form.get("id", "").strip()
        try:
            tid = cards_db.resolve_treatment_for_visit(
                pib=payload["pib"],
                rank_unit=payload["rank_unit"],
                treatment_id=payload.get("treatment_id", ""),
                new_treatment_title=payload.get("new_treatment_title", ""),
                visit_date=payload.get("date") or _outpatient_today(),
                diagnosis=payload.get("diagnosis", ""),
            )
        except Exception as e:
            logger.error("Лікування для звернення: %s", e)
            flash(f"Помилка прив'язки лікування: {e}", "error")
            return redirect(url_for("outpatient_journal"))
        payload["treatment_id"] = tid if tid is not None else ""

        # Картка пацієнта + дані з Excel
        try:
            pid = cards_db.get_or_create_patient(payload["pib"], payload.get("rank_unit") or "")
            _sync_patient_card(pid, payload["pib"])
        except Exception as e:
            logger.warning("Синхронізація картки пацієнта: %s", e)

        if entry_id:
            existing = get_outpatient_entry(entry_id)
            if existing:
                # Не затирати поля, які Tom Select міг не надіслати порожніми
                care_now = payload.get("care_type") or existing.get("care_type") or ""
                wa_now = payload.get("wa_event") or existing.get("wa_event") or ""
                for key in (
                    "pib", "rank_unit", "diagnosis", "referral_to", "lpz",
                    "exam_result", "notes", "care_type", "wa_event", "from_lpz", "date",
                    "leave_start", "leave_end", "leave_days", "mkx10",
                ):
                    if not payload.get(key) and existing.get(key):
                        # навмисно порожні поля за сценарієм дії — не відновлювати
                        if key in ("leave_start", "leave_end", "leave_days") and care_now in (
                            "consultation", "referral",
                        ):
                            continue
                        if key in ("leave_end", "leave_days") and care_now == "vlk":
                            continue
                        if key in ("leave_end", "leave_days") and care_now == "inpatient" and wa_now in (
                            cards_db.WA_EVENT_OPEN,
                            cards_db.WA_EVENT_EXTEND,
                        ):
                            continue
                        if key == "from_lpz" and not (
                            care_now == "inpatient" and wa_now == cards_db.WA_EVENT_EXTEND
                        ):
                            continue
                        payload[key] = existing.get(key)
                if tid is None and existing.get("treatment_id"):
                    payload["treatment_id"] = existing.get("treatment_id")
            if update_outpatient_entry(entry_id, payload):
                visit_id = int(entry_id)
                flash("Запис оновлено.", "success")
            else:
                flash("Запис для редагування не знайдено.", "error")
                return redirect(url_for("outpatient_journal"))
        else:
            visit_id = insert_outpatient_entry(payload)
            flash("Запис додано.", "success")

        # Фото з Gemini / прикріплені до звернення → медіа лікування + Dropbox
        try:
            tid_media = int(payload["treatment_id"]) if payload.get("treatment_id") else None
        except (TypeError, ValueError):
            tid_media = None
        if tid_media and visit_id:
            media_files = (
                request.files.getlist("gemini_media")
                or request.files.getlist("media_files")
            )
            saved_n, drop_n, drop_err = _save_uploaded_media_list(
                tid_media,
                media_files,
                visit_id=int(visit_id),
            )
            if saved_n:
                msg = f"Додано медіа: {saved_n}."
                if drop_n:
                    msg += f" У Dropbox: {drop_n}."
                flash(msg, "success")
            if drop_err:
                flash(f"Dropbox не синхронізовано: {drop_err}", "error")

        # Явне «Закрив / Виписаний» у зверненні → завершити лікування
        if (
            payload.get("wa_event") == cards_db.WA_EVENT_DISCHARGE
            and payload.get("treatment_id")
        ):
            try:
                cards_db.set_treatment_status(int(payload["treatment_id"]), "closed")
            except (TypeError, ValueError):
                pass

        # Період живе в полях звернення (leave_*); нагадування — з leave_end
        tid = payload.get("treatment_id")
        try:
            tid_i = int(tid) if tid not in (None, "") else None
        except (TypeError, ValueError):
            tid_i = None
        return_tid = _normalize_spaces(request.form.get("return_treatment", ""))
        try:
            return_tid_i = int(return_tid) if return_tid else None
        except (TypeError, ValueError):
            return_tid_i = None
        target_tid = tid_i or return_tid_i
        if target_tid:
            return redirect(url_for(
                "patient_treatment_detail",
                treatment_id=target_tid,
                wa=visit_id,
            ))
        return redirect(url_for("outpatient_journal"))

    pib_query = _normalize_spaces(request.args.get("q", ""))
    unit_filter = _normalize_spaces(request.args.get("unit", ""))
    show_new = request.args.get("new", "").strip() in ("1", "true", "yes")
    extend_tid = request.args.get("extend", "").strip()
    extend_prefill = None
    if extend_tid and not edit_row:
        try:
            extend_prefill = cards_db.suggest_extend_prefill(int(extend_tid))
        except (TypeError, ValueError):
            extend_prefill = None
        if extend_prefill:
            show_new = True
            if not pib_query:
                pib_query = extend_prefill.get("pib") or ""

    show_form = bool(edit_row) or show_new
    records = list_outpatient_entries(
        order_desc=True,
        pib_query=pib_query,
        unit=unit_filter,
    )
    pib_map = cards_db.patient_id_map_by_pib()
    for rec in records:
        rec["unit_label"] = _parse_outpatient_unit(rec.get("rank_unit", ""))
        ct = rec.get("care_type") or ""
        rec["care_type_label"] = cards_db.RESTRICTION_KINDS.get(ct, ct)
        we = (rec.get("wa_event") or "").strip()
        rec["wa_event_label"] = ""
        for key, label in cards_db.CARE_WA_EVENTS.get(ct, []):
            if key == we:
                rec["wa_event_label"] = label
                break
        if not rec["wa_event_label"] and we:
            rec["wa_event_label"] = cards_db.WA_EVENT_LABELS.get(we, we)
        rec["exam_result_items"] = cards_db.parse_recommendation_items(
            rec.get("exam_result") or ""
        )
        rec["referral_to_label"] = cards_db.specialist_to_nominative(
            rec.get("referral_to") or ""
        ) or (rec.get("referral_to") or "")
        pib_key = _normalize_spaces(rec.get("pib") or "").casefold()
        rec["patient_id"] = pib_map.get(pib_key)
    all_asc = list_outpatient_entries(order_desc=False)
    journal_count = len(all_asc)
    unit_options = list_outpatient_units()
    edit_defaults = {c: "" for c in OUTPATIENT_COLUMNS}
    if edit_row:
        edit_defaults = {c: str(edit_row.get(c) or "") for c in OUTPATIENT_COLUMNS}
    elif extend_prefill:
        for k, v in extend_prefill.items():
            if k in edit_defaults or k in (
                "treatment_id", "pib", "rank_unit", "diagnosis", "referral_to",
                "lpz", "exam_result", "care_type", "wa_event", "leave_start", "leave_end",
                "leave_days", "notes",
              ):
                edit_defaults[k] = str(v or "")

    return render_template(
        "outpatient_journal.html",
        records=records,
        edit=edit_defaults,
        show_form=show_form,
        journal_count=journal_count,
        pib_query=pib_query,
        unit_filter=unit_filter,
        unit_options=unit_options,
        today_iso=datetime.now().strftime("%Y-%m-%d"),
        today_ui=_outpatient_today(),
        care_types=cards_db.RESTRICTION_KINDS,
        care_wa_events=cards_db.CARE_WA_EVENTS,
        extend_mode=bool(extend_prefill),
        return_treatment=request.args.get("return_treatment", "").strip(),
        reminders=cards_db.list_upcoming_reminders(),
        reminder_within=REMINDER_WITHIN_DAYS,
    )


@app.route("/api/patient_treatments", methods=["GET"])
def api_patient_treatments():
    """Список лікувань пацієнта для селекту в журналі."""
    pib = _normalize_spaces(request.args.get("pib", ""))
    active_only = request.args.get("active_only", "1") != "0"
    if not pib:
        return jsonify({"results": []})
    _outpatient_init_db()
    items = cards_db.list_treatments_for_pib(pib, active_only=active_only)
    return jsonify({
        "results": [
            {
                "id": t["id"],
                "title": t["title"],
                "status": t["status"],
                "label": t["title"] + (" (завершене)" if t["status"] != "active" else ""),
            }
            for t in items
        ]
    })


@app.route("/patient-cards", methods=["GET", "POST"])
def patient_cards():
    """Картотека пацієнтів."""
    _outpatient_init_db()

    if request.method == "POST":
        action = _normalize_spaces(request.form.get("action", ""))
        if action == "sync_all_base":
            try:
                stats = _ensure_patients_from_base(enrich_treatments=True)
                flash(
                    f"Синхронізовано з кадрової бази: створено {stats['created']} карток, "
                    f"оновлено дані в {stats['updated']}. "
                    f"У базі {stats['total_base']} в/с, у картотеці зараз {stats['total_patients']}.",
                    "success",
                )
            except Exception as e:
                logger.exception("Синхронізація карток з base.xlsx")
                flash(f"Помилка синхронізації: {e}", "error")
            return redirect(url_for("patient_cards"))

    q = _normalize_spaces(request.args.get("q", ""))
    patients = cards_db.list_patients(q)
    return render_template(
        "patient_cards.html",
        patients=patients,
        pib_query=q,
    )


def _vlk_referral_basis_label(value: str) -> str:
    key = _normalize_spaces(value).casefold()
    for opt_key, label in cards_db.VLK_REFERRAL_BASIS_OPTIONS:
        if opt_key == key or label.casefold() == key:
            return label
    return ""


def _build_vlk_fabula_context(patient: dict, form) -> dict:
    """Контекст фабули на ВЛК з картки + форми."""
    pib = _normalize_spaces(patient.get("pib") or "")
    treatments_person = _lookup_person_in_treatments_excel(pib)
    diagnosis = _normalize_spaces(form.get("diagnosis", ""))
    if not diagnosis:
        diagnosis = cards_db.latest_diagnosis_for_patient(int(patient["id"]))
    if not diagnosis:
        diagnosis = treatments_person.get("diagnosis") or ""

    service_category = cards_db.normalize_service_category(
        form.get("service_category") or patient.get("service_category") or ""
    )
    if not service_category:
        service_category = treatments_person.get("service_category") or ""

    rank = _normalize_spaces(patient.get("rank") or "")
    position = _normalize_spaces(patient.get("position") or "")
    birth_date = _normalize_spaces(form.get("birth_date") or patient.get("birth_date") or "")
    if not birth_date:
        birth_date = treatments_person.get("birth_date") or ""
    phone = _normalize_spaces(form.get("phone") or patient.get("phone") or "")
    if not phone:
        phone = treatments_person.get("phone") or ""
    ipn = _normalize_spaces(form.get("ipn") or patient.get("ipn") or "")
    enlistment_date = _normalize_spaces(
        form.get("enlistment_date") or patient.get("enlistment_date") or ""
    )
    komisariat = _normalize_spaces(form.get("komisariat") or patient.get("komisariat") or "")
    vlk_place = _normalize_spaces(form.get("vlk_place", ""))
    referral_basis = _vlk_referral_basis_label(form.get("referral_basis", ""))
    rank_and_category = cards_db.format_rank_and_category(rank, service_category)

    return {
        "vlk_place": vlk_place,
        "rank": rank,
        "service_category": service_category,
        "rank_and_category": rank_and_category,
        "position": position,
        "pib": pib,
        "birth_date": birth_date,
        "enlistment_date": enlistment_date,
        "komisariat": komisariat,
        "diagnosis": diagnosis,
        "ipn": ipn,
        "phone": phone,
        "referral_basis": referral_basis,
        "data_stvorennya": datetime.now().strftime("%d.%m.%Y"),
    }


def _format_vlk_fabula_text(ctx: dict) -> str:
    lines = [
        "ФАБУЛА НА ВЛК",
        "",
        f"1. Місце проходження ВЛК: {ctx.get('vlk_place') or '—'}",
        f"2. Звання та категорія: {ctx.get('rank_and_category') or '—'}",
        f"3. Повна посада: {ctx.get('position') or '—'}",
        f"4. ПІБ: {ctx.get('pib') or '—'}",
        f"5. Дата народження: {ctx.get('birth_date') or '—'}",
        f"6. Дата призову на службу: {ctx.get('enlistment_date') or '—'}",
        f"7. Яким військкоматом призваний: {ctx.get('komisariat') or '—'}",
        f"8. Діагноз: {ctx.get('diagnosis') or '—'}",
        f"9. ІПН: {ctx.get('ipn') or '—'}",
        f"10. Номер телефону: {ctx.get('phone') or '—'}",
        f"11. Підстава направлення: {ctx.get('referral_basis') or '—'}",
    ]
    return "\n".join(lines)


def _ensure_vlk_fabula_template() -> str:
    """Створює templates/vlk_fabula.docx, якщо ще немає."""
    path = VLK_FABULA_TEMPLATE
    if os.path.isfile(path):
        return path
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc = DocxDocument()
    title = doc.add_paragraph()
    run = title.add_run("ФАБУЛА НА ВЛК")
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    fields = [
        ("1. Місце проходження ВЛК:", "{{ vlk_place }}"),
        ("2. Звання та категорія:", "{{ rank_and_category }}"),
        ("3. Повна посада:", "{{ position }}"),
        ("4. ПІБ:", "{{ pib }}"),
        ("5. Дата народження:", "{{ birth_date }}"),
        ("6. Дата призову на службу:", "{{ enlistment_date }}"),
        ("7. Яким військкоматом призваний:", "{{ komisariat }}"),
        ("8. Діагноз:", "{{ diagnosis }}"),
        ("9. ІПН:", "{{ ipn }}"),
        ("10. Номер телефону:", "{{ phone }}"),
        ("11. Підстава направлення:", "{{ referral_basis }}"),
    ]
    for label, placeholder in fields:
        p = doc.add_paragraph()
        p.add_run(f"{label} ").bold = True
        p.add_run(placeholder)
    doc.add_paragraph("")
    foot = doc.add_paragraph("{{ data_stvorennya }}")
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save(path)
    return path


@app.route("/patient-cards/<int:patient_id>", methods=["GET", "POST"])
def patient_card_detail(patient_id: int):
    _outpatient_init_db()
    patient = cards_db.get_patient(patient_id)
    if not patient:
        flash("Пацієнта не знайдено.", "error")
        return redirect(url_for("patient_cards"))

    if request.method == "POST":
        action = _normalize_spaces(request.form.get("action", "save"))
        if action == "sync_base":
            changed = _sync_patient_card(
                patient_id, patient.get("pib") or "", overwrite=True
            )
            if changed:
                flash("Дані оновлено з кадрової бази та бази лікувань.", "success")
            else:
                flash(
                    "За цим ПІБ не знайдено даних у base.xlsx / treatments для оновлення.",
                    "error",
                )
            return redirect(url_for("patient_card_detail", patient_id=patient_id))

        if action == "save":
            cards_db.update_patient(
                patient_id,
                birth_date=_normalize_spaces(request.form.get("birth_date", "")),
                phone=_normalize_spaces(request.form.get("phone", "")),
                ipn=_normalize_spaces(request.form.get("ipn", "")),
                service_category=_normalize_spaces(request.form.get("service_category", "")),
                enlistment_date=_normalize_spaces(request.form.get("enlistment_date", "")),
                komisariat=_normalize_spaces(request.form.get("komisariat", "")),
            )
            flash("Картку збережено.", "success")
            return redirect(url_for("patient_card_detail", patient_id=patient_id))

        if action == "vlk_fabula":
            # зберегти ручні поля картки перед генерацією
            cards_db.update_patient(
                patient_id,
                birth_date=_normalize_spaces(request.form.get("birth_date", ""))
                or (patient.get("birth_date") or ""),
                phone=_normalize_spaces(request.form.get("phone", ""))
                or (patient.get("phone") or ""),
                ipn=_normalize_spaces(request.form.get("ipn", ""))
                or (patient.get("ipn") or ""),
                service_category=_normalize_spaces(request.form.get("service_category", ""))
                or (patient.get("service_category") or ""),
                enlistment_date=_normalize_spaces(request.form.get("enlistment_date", "")),
                komisariat=_normalize_spaces(request.form.get("komisariat", "")),
            )
            patient = cards_db.get_patient(patient_id)
            ctx = _build_vlk_fabula_context(patient, request.form)
            if not ctx.get("vlk_place"):
                flash("Вкажіть місце проходження ВЛК.", "error")
                return redirect(url_for("patient_card_detail", patient_id=patient_id))
            if not ctx.get("referral_basis"):
                flash("Оберіть підставу направлення.", "error")
                return redirect(url_for("patient_card_detail", patient_id=patient_id))
            try:
                template_path = _ensure_vlk_fabula_template()
                doc = DocxTemplate(template_path)
                doc.render(ctx)
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                pib_for_file = _normalize_spaces(ctx.get("pib") or "")
                pib_for_file = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "", pib_for_file).strip(" .")
                filename = f"Фабула на ВЛК - {pib_for_file or 'пацієнт'}.docx"
                return send_file(
                    buf,
                    as_attachment=True,
                    download_name=filename,
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception as e:
                logger.exception("Фабула на ВЛК")
                flash(f"Помилка створення фабули: {e}", "error")
                return redirect(url_for("patient_card_detail", patient_id=patient_id))

        return redirect(url_for("patient_card_detail", patient_id=patient_id))

    # м'який sync при відкритті (лише порожні поля)
    _sync_patient_card(patient_id, patient.get("pib") or "", overwrite=False)
    patient = cards_db.get_patient(patient_id)
    treatments = cards_db.list_treatments_for_patient(patient_id)
    unit_line = cards_db.format_wa_unit_line(
        patient.get("unit_short") or "",
        canonical_unit=_parse_outpatient_unit(patient.get("rank_unit") or ""),
    )
    diagnosis_default = cards_db.latest_diagnosis_for_patient(patient_id)
    if not diagnosis_default:
        diagnosis_default = _lookup_person_in_treatments_excel(
            patient.get("pib") or ""
        ).get("diagnosis") or ""
    return render_template(
        "patient_card_detail.html",
        patient=patient,
        treatments=treatments,
        unit_line=unit_line,
        service_category_options=cards_db.SERVICE_CATEGORY_OPTIONS,
        referral_basis_options=cards_db.VLK_REFERRAL_BASIS_OPTIONS,
        diagnosis_default=diagnosis_default,
        rank_and_category=cards_db.format_rank_and_category(
            patient.get("rank") or "",
            patient.get("service_category") or "",
        ),
    )


@app.route("/patient-cards/treatment/<int:treatment_id>", methods=["GET", "POST"])
def patient_treatment_detail(treatment_id: int):
    _outpatient_init_db()
    treatment = cards_db.get_treatment(treatment_id)
    if not treatment:
        flash("Лікування не знайдено.", "error")
        return redirect(url_for("patient_cards"))

    if request.method == "POST":
        action = _normalize_spaces(request.form.get("action", ""))
        if action == "close":
            cards_db.set_treatment_status(treatment_id, "closed")
            flash("Лікування завершено.", "success")
            return redirect(url_for(
                "patient_treatment_detail",
                treatment_id=treatment_id,
                wa_event="discharge",
            ))
        if action == "reopen":
            cards_db.set_treatment_status(treatment_id, "active")
            flash("Лікування знову активне.", "success")
        elif action == "upload_media":
            storages = [
                f
                for f in (
                    request.files.getlist("media_files")
                    or request.files.getlist("media_file")
                )
                if f and f.filename
            ]
            if not storages:
                flash("Файл не вибрано.", "error")
            else:
                saved_n, drop_n, drop_err = _save_uploaded_media_list(
                    treatment_id,
                    storages,
                )
                if saved_n:
                    msg = f"Завантажено файлів: {saved_n}."
                    if drop_n:
                        msg += f" У Dropbox: {drop_n}."
                    flash(msg, "success")
                failed_n = len(storages) - saved_n
                if failed_n > 0:
                    flash(f"Не збережено файлів: {failed_n}.", "error")
                if drop_err:
                    flash(f"Dropbox не синхронізовано: {drop_err}", "error")
        elif action == "delete_media":
            mid = request.form.get("media_id", "")
            try:
                if cards_db.delete_media(int(mid)):
                    flash("Файл видалено.", "success")
                else:
                    flash("Файл не знайдено.", "error")
            except Exception:
                flash("Не вдалося видалити файл.", "error")
        elif action == "delete_visit":
            visit_id = request.form.get("visit_id", "").strip()
            entry = get_outpatient_entry(visit_id)
            if (
                entry
                and str(entry.get("treatment_id") or "") == str(treatment_id)
                and delete_outpatient_entry(visit_id)
            ):
                flash("Звернення видалено.", "success")
            else:
                flash("Не вдалося видалити звернення.", "error")
        return redirect(url_for("patient_treatment_detail", treatment_id=treatment_id))

    treatment = cards_db.get_treatment(treatment_id)
    visits = cards_db.list_visits_for_treatment(treatment_id)
    media = cards_db.list_media(treatment_id)
    days_count = cards_db.treatment_day_count(treatment)
    totals = cards_db.visit_care_totals(treatment_id)

    wa_text = ""
    wa_event = request.args.get("wa_event", "").strip()
    wa_id = request.args.get("wa", "").strip()
    if wa_event == "discharge" or (not wa_id and treatment.get("status") == "closed"):
        patient = _enrich_patient_for_whatsapp(
            cards_db.get_patient(int(treatment["patient_id"]))
        )
        last = visits[-1] if visits else {}
        care, leave_start, leave_end = cards_db.discharge_period(treatment, visits)
        # для ВЛК дата завершення — день закриття лікування, місце — з останнього звернення
        if care == "vlk":
            leave_end = (
                _normalize_spaces(treatment.get("closed_on") or "")
                or leave_end
                or last.get("date")
                or cards_db.today_ui()
            )
        wa_text = cards_db.build_whatsapp_message(
            patient or {"pib": treatment.get("pib")},
            action=cards_db.wa_action_phrase(
                cards_db.WA_EVENT_DISCHARGE,
                care_type=care,
                lpz=last.get("lpz") or "",
                date=last.get("date") or "",
                leave_start=leave_start,
                leave_end=leave_end,
                specialist=last.get("referral_to") or "",
            ),
            diagnosis=last.get("diagnosis") or treatment.get("title") or "",
            recommendations=last.get("exam_result") or last.get("notes") or "",
            canonical_unit=_parse_outpatient_unit(
                (patient or {}).get("rank_unit") or treatment.get("patient_rank_unit") or ""
            ),
        )
    elif wa_id:
        entry = get_outpatient_entry(wa_id)
        if entry:
            wa_text = _whatsapp_for_visit(entry)
    elif visits:
        # за замовчуванням — текст останнього звернення
        wa_text = _whatsapp_for_visit(visits[-1])

    return render_template(
        "patient_treatment_detail.html",
        treatment=treatment,
        visits=visits,
        media=media,
        days_count=days_count,
        care_totals=totals,
        care_types=cards_db.RESTRICTION_KINDS,
        today_ui=_outpatient_today(),
        wa_text=wa_text,
    )


@app.route("/patient-cards/media/<int:media_id>", methods=["GET"])
def patient_media_file(media_id: int):
    _outpatient_init_db()
    media = cards_db.get_media(media_id)
    if not media:
        flash("Файл не знайдено.", "error")
        return redirect(url_for("patient_cards"))
    path = cards_db.media_disk_path(media)
    if not os.path.isfile(path):
        flash("Файл відсутній на диску.", "error")
        return redirect(url_for("patient_treatment_detail", treatment_id=media["treatment_id"]))
    return send_file(
        path,
        mimetype=media.get("mime") or "application/octet-stream",
        as_attachment=False,
        download_name=media.get("original_name") or media.get("filename"),
    )


def _payments_month_year_from_request():
    """Місяць/рік відомості з query/form; за замовчуванням — поточний."""
    from utils.payments_unpaid import MONTH_EN_TO_NUM, MONTH_UA, list_payment_files

    now = datetime.now()
    year_raw = request.values.get("year") or str(now.year)
    month_raw = request.values.get("month") or str(now.month)
    try:
        year = int(year_raw)
    except (TypeError, ValueError):
        year = now.year
    month = None
    if isinstance(month_raw, str) and month_raw.strip().isalpha():
        month = MONTH_EN_TO_NUM.get(month_raw.strip().lower())
    if month is None:
        try:
            month = int(month_raw)
        except (TypeError, ValueError):
            month = now.month
    if month < 1 or month > 12:
        month = now.month
    # якщо є payment-файли — за замовчуванням останній місяць з файлів при першому заході без params
    if "year" not in request.values and "month" not in request.values:
        files = list_payment_files()
        if files:
            last = files[-1]
            year, month = last.year, last.month
    return year, month, MONTH_UA.get(month, str(month))


@app.route("/payments-vedomost", methods=["GET", "POST"])
def payments_vedomost():
    """Прев’ю неоплачених бойових стаціонарів / реабілітації / відпусток + DOCX."""
    from utils.payments_unpaid import (
        TYPE_UA,
        aggregate_vedomost_rows,
        find_unpaid_segments,
        list_payment_files,
        load_combat_patterns,
        load_payments_treatments_df,
        render_payments_docx,
        unique_matched_diagnoses,
    )

    os.makedirs(PAYMENTS_DIR, exist_ok=True)
    year, month, month_ua = _payments_month_year_from_request()
    payment_files = list_payment_files()
    patterns = [
        {"id": pid, "label": label, "regex": cre.pattern}
        for pid, label, cre in load_combat_patterns(force_reload=True)
    ]

    action = _normalize_spaces(request.values.get("action", ""))
    download = action == "download" or request.values.get("download") == "1"

    treatments_df = load_payments_treatments_df(year)
    treatments_source = f"treatments_{year}.xlsx"
    segments = []
    rows = []
    error = None
    if treatments_df is None or getattr(treatments_df, "empty", True):
        error = (
            f"Немає файлу {treatments_source} у data/. "
            "Завантажте актуальну базу лікувань на сторінці «Бази даних»."
        )
    else:
        try:
            segments = find_unpaid_segments(treatments_df, year=year, month=month)
            rows = aggregate_vedomost_rows(segments)
        except Exception as e:
            logger.exception("payments-vedomost")
            error = str(e)

    if download and not error:
        if not os.path.isfile(PAYMENTS_TEMPLATE):
            flash("Немає шаблону payments_template.docx", "error")
            return redirect(url_for("payments_vedomost", year=year, month=month))
        os.makedirs(TEMP_DIR, exist_ok=True)
        out_name = f"Vidomist_neoplachenykh_{month:02d}_{year}.docx"
        out_path = os.path.join(TEMP_DIR, out_name)
        try:
            render_payments_docx(
                PAYMENTS_TEMPLATE,
                out_path,
                rows,
                year=year,
                month=month,
            )
        except Exception as e:
            logger.exception("payments docx")
            flash(f"Помилка генерації Word: {e}", "error")
            return redirect(url_for("payments_vedomost", year=year, month=month))
        return send_file(
            out_path,
            as_attachment=True,
            download_name=out_name,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    preview_rows = []
    for seg in segments:
        preview_rows.append(
            {
                "pib": seg.pib,
                "care_type": TYPE_UA.get(seg.care_type, seg.care_type),
                "start": seg.start.strftime("%d.%m.%Y"),
                "end": seg.end.strftime("%d.%m.%Y"),
                "days": seg.days,
                "pattern": seg.pattern_label,
                "snippet": seg.diagnosis_snippet,
                "unit": seg.unit,
                "rank": seg.rank,
                "missing_discharge": seg.missing_discharge,
            }
        )

    return render_template(
        "payments_vedomost.html",
        year=year,
        month=month,
        month_ua=month_ua,
        payment_files=payment_files,
        patterns=patterns,
        preview_rows=preview_rows,
        vedomost_rows=rows,
        matched_diagnoses=unique_matched_diagnoses(segments),
        error=error,
        payments_dir=PAYMENTS_DIR,
        history_start=PAYMENTS_HISTORY_START,
        unit_filter=PAYMENTS_UNIT_FILTER,
        carry_day=PAYMENTS_PREV_MONTH_CARRY_DAY,
        treatments_source=treatments_source,
    )


def _warmup_treatments_cache():
    """Фоновий прогрів кешу після старту — перший пошук не чекає читання великих Excel."""
    try:
        with app.app_context():
            load_treatments_data()
            load_base_personnel_data()
            load_rkch_personnel_index()
        logger.info("Прогрів кешу Excel завершено — пошук готовий.")
    except Exception as e:
        logger.warning(
            "Прогрів кешу Excel не вдався (дані завантажаться при першому запиті): %s",
            e,
        )


if __name__ == '__main__':
    threading.Thread(target=_warmup_treatments_cache, daemon=True).start()
    app.run(debug=DEBUG)